# -*- coding: utf-8 -*-
"""
Deterministic Post-Quantum Cryptography (PQC) Readiness scanner.

Scans uploaded evidence text -- TLS/SSL, SSH, IPSec/VPN, PKI/certificate, database
encryption-at-rest, HSM/KMS, or code-signing configuration exports/dumps (plain
text, HTML/XML, or PDF-extracted text; any filename) -- for known cryptographic
algorithm strings and classifies each hit as:
  - "VULNERABLE": broken by Shor's algorithm on a sufficiently large quantum
    computer (RSA, DSA, Diffie-Hellman/DHE, ECC/ECDSA/ECDH).
  - "WEAK": classically weak/deprecated, not PQC-specific but still flagged
    (MD5, SHA-1, DES/3DES/RC4, SSLv2/SSLv3/TLS1.0/TLS1.1, CBC-mode).
  - "SAFE": quantum-resistant / NIST-selected PQC (AES-256, SHA-256+, SHA-3,
    CRYSTALS-Kyber/ML-KEM, CRYSTALS-Dilithium/ML-DSA, SPHINCS+/SLH-DSA,
    Falcon/FN-DSA, ChaCha20-Poly1305).

100% deterministic regex/keyword matching -- zero LLM/RAG involvement. Same
design principle as nessus_parser.py / trivy_parser.py for VAPT scanner
findings: identifying a known algorithm string is exact pattern matching, not
judgment, so there is nothing for an LLM to usefully add here.

Supported input formats (as of this version):
  - Plain text / config exports (.conf, .cnf, .ini, .cfg, .properties, .yaml,
    .yml, .toml, .sh, .bat, .ps1, .reg)
  - PDF documents (.pdf)  -- text-layer extracted via pdfplumber; scanned
    PDFs fall back to doctr OCR
  - Microsoft Word documents (.docx)  -- paragraphs + table cells extracted
    via python-docx
  - Images (.png, .jpg, .jpeg, .webp, .bmp, .tiff, .tif)  -- OCR via doctr
"""
import io
import json
import os
import re
from typing import List, Tuple, Any, Callable, Dict, Optional, Union

from .base_parser import BaseParser, is_image_file
from .finding_schema import Finding
from .control_mapper import map_pqc_findings_list
from .pqc_crypto_db import (
    scan_oids_in_text,
    scan_iana_hex_in_text,
    scan_liboqs_in_text,
)

# ══════════════════════════════════════════════════════════════════════════════
# DETECTION GATE (can_parse)
# ══════════════════════════════════════════════════════════════════════════════

# Broad crypto/protocol/config vocabulary. PQC evidence isn't one fixed export
# format (unlike Nessus's native XML) -- it can be a TLS server config export, an
# sshd_config, an IPSec policy dump, a certificate listing, a DB encryption
# settings screen, etc. Any single one of these words alone is common in
# unrelated documents (an HR policy can say "encryption"), so -- mirroring
# nessus_parser.py's "weak_signal_count >= 2" pattern for exactly this kind of
# non-exclusive-signature situation -- at least 2 distinct hits are required.
_PQC_KEYWORDS = (
    "tls", "ssl", "cipher", "cipher suite", "ssh", "sshd", "ipsec", "vpn",
    "pki", "certificate", "x.509", "x509", "encryption", "encrypted",
    "key exchange", "key size", "keysize", "hsm", "keystore", "kms",
    "rsa", "ecc", "ecdsa", "ecdh", "aes", "sha", "algorithm", "handshake",
    "cryptograph", "signature algorithm", "diffie-hellman", "diffie hellman",
    "elliptic curve", "public key", "private key", "tde", "code signing",
    "code-signing", "firmware signing",
)

# Config-export filename extensions -- still content-gated (never filename-only,
# per every other parser's can_parse() convention), just given a lower keyword
# bar since the extension itself is already a meaningful signal.
_PQC_CONFIG_EXTENSIONS = (".conf", ".cnf", ".pem", ".crt", ".cer", ".key", ".p12", ".pfx", ".jks",
                          ".config", ".ini", ".cfg",
                          ".yaml", ".yml",        # Kubernetes, Docker Compose, Ansible
                          ".properties",          # Spring Boot / Java / Kafka
                          ".toml",                # Rust / server configs (rustls, actix, etc.)
                          ".sh", ".bat", ".ps1",  # Shell/PowerShell startup scripts with JVM -D flags
                          ".reg")                 # Windows CryptoAPI / CNG registry exports

# Binary document formats that require text extraction before regex scanning.
# PQCParser handles these natively via pqc_extract_text().
_PQC_BINARY_EXTENSIONS = (
    ".pdf",                                    # PDF -- text layer + OCR fallback
    ".docx", ".doc",                           # Microsoft Word
    ".png", ".jpg", ".jpeg",                   # Raster images (OCR)
    ".webp", ".bmp", ".tiff", ".tif",          # Additional image formats
    # ── Binary certificate / key / keystore formats ─────────────────────────
    ".cer", ".crt", ".der",                    # DER / PEM X.509 certificates
    ".key",                                    # PEM private keys (RSA/EC/Ed25519)
    ".p12", ".pfx",                            # PKCS#12 keystores
    ".jks",                                    # Java KeyStore
    ".pub",                                    # SSH public key files
)


def pqc_extract_text(filename: str, raw_bytes: bytes) -> str:
    """Extract plain text from a binary document for PQC regex scanning.

    Dispatch table (by extension):
      .pdf          pdfplumber text layer → doctr OCR fallback for scanned PDFs
      .docx / .doc  python-docx paragraphs + table cells
      .png/.jpg/... doctr OCR (via existing doc_parsers.extract_text)

    Returns extracted plain text, or empty string on any error.
    """
    ext = os.path.splitext(filename.lower())[1]

    # ── PDF ────────────────────────────────────────────────────────────────────
    if ext == ".pdf":
        text = ""
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
            text = "\n".join(pages).strip()
        except Exception:
            text = ""
        # Fallback to OCR for scanned / image-only PDFs
        if not text or len(text.strip()) < 50:
            try:
                from src.core.parsers.doc_parsers import extract_text
                buf = io.BytesIO(raw_bytes)
                buf.name = filename
                text = extract_text(buf) or ""
            except Exception:
                pass
        return text

    # ── DOCX / DOC ─────────────────────────────────────────────────────────────
    if ext in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            return "\n".join(parts)
        except Exception:
            return ""

    # ── PNG / JPG / JPEG / WEBP / BMP / TIFF ───────────────────────────────────
    if ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"):
        try:
            from src.core.parsers.doc_parsers import extract_text
            buf = io.BytesIO(raw_bytes)
            buf.name = filename
            return extract_text(buf) or ""
        except Exception:
            return ""

    # ── X.509 Certificates: DER / PEM (.cer / .crt / .der) ──────────────────────
    if ext in (".cer", ".crt", ".der"):
        try:
            from cryptography import x509 as _x509
            from cryptography.hazmat.primitives.asymmetric import (
                rsa as _rsa, ec as _ec,
                ed25519 as _ed25519, ed448 as _ed448,
            )
            from cryptography.hazmat.primitives.asymmetric.dsa import DSAPublicKey as _DSAKey

            def _cert_to_text(cert):
                subject  = cert.subject.rfc4514_string()
                issuer   = cert.issuer.rfc4514_string()
                sig_oid  = cert.signature_algorithm_oid.dotted_string
                sig_hash = (cert.signature_hash_algorithm.name
                            if cert.signature_hash_algorithm else "unknown")
                try:
                    nb = cert.not_valid_before_utc
                    na = cert.not_valid_after_utc
                except AttributeError:
                    nb = cert.not_valid_before
                    na = cert.not_valid_after
                pub = cert.public_key()
                if isinstance(pub, _rsa.RSAPublicKey):
                    key_info = f"RSA key_size={pub.key_size}"
                elif isinstance(pub, _ec.EllipticCurvePublicKey):
                    key_info = f"ECDSA elliptic_curve={pub.curve.name}"
                elif isinstance(pub, _ed25519.Ed25519PublicKey):
                    key_info = "Ed25519 elliptic curve"
                elif isinstance(pub, _ed448.Ed448PublicKey):
                    key_info = "Ed448 elliptic curve"
                elif isinstance(pub, _DSAKey):
                    key_info = f"DSA key_size={pub.key_size}"
                else:
                    key_info = "unknown_algorithm"
                return (
                    f"X.509 Certificate\nSubject: {subject}\nIssuer: {issuer}\n"
                    f"Signature Algorithm OID: {sig_oid}\nSignature Hash: {sig_hash}\n"
                    f"Public Key: {key_info}\n"
                    f"Valid From: {nb}\nValid Until: {na}\n"
                    f"ssl certificate tls algorithm encryption key exchange\n"
                )

            lines_out = []
            try:
                for cert in _x509.load_pem_x509_certificates(raw_bytes):
                    lines_out.append(_cert_to_text(cert))
            except Exception:
                pass
            if not lines_out:
                try:
                    lines_out.append(_cert_to_text(_x509.load_pem_x509_certificate(raw_bytes)))
                except Exception:
                    pass
            if not lines_out:
                try:
                    lines_out.append(_cert_to_text(_x509.load_der_x509_certificate(raw_bytes)))
                except Exception:
                    pass
            if lines_out:
                return "\n".join(lines_out)
        except Exception:
            pass
        return ""

    # ── PEM Private Keys (.key) ───────────────────────────────────────────────────
    if ext == ".key":
        try:
            text = raw_bytes.decode("utf-8", errors="ignore")
            if "-----BEGIN" in text:
                if "RSA" in text:
                    return f"RSA private key ssl tls algorithm encryption\n{text[:2000]}"
                if "EC PRIVATE" in text:
                    return f"ECDSA elliptic curve private key ssl tls algorithm encryption\n{text[:2000]}"
                if "ED25519" in text.upper():
                    return "Ed25519 elliptic curve private key ssl tls algorithm encryption"
                if "DSA" in text:
                    return f"DSA private key ssl tls algorithm encryption\n{text[:2000]}"
                return f"private key ssl tls algorithm encryption\n{text[:2000]}"
            from cryptography.hazmat.primitives.serialization import load_der_private_key
            from cryptography.hazmat.primitives.asymmetric import rsa as _rsa, ec as _ec
            key = load_der_private_key(raw_bytes, password=None)
            if isinstance(key, _rsa.RSAPrivateKey):
                return f"RSA private key key_size={key.key_size} ssl tls algorithm encryption"
            if isinstance(key, _ec.EllipticCurvePrivateKey):
                return f"ECDSA elliptic curve private key curve={key.curve.name} ssl tls algorithm encryption"
            return "private key ssl tls algorithm encryption"
        except Exception:
            pass
        return ""

    # ── PKCS#12 Keystores (.p12 / .pfx) ─────────────────────────────────────────
    if ext in (".p12", ".pfx"):
        try:
            from cryptography.hazmat.primitives.serialization import pkcs12 as _pkcs12
            from cryptography.hazmat.primitives.asymmetric import rsa as _rsa, ec as _ec
            lines_out = ["PKCS12 keystore ssl tls certificate key exchange algorithm encryption"]
            for pwd in (None, b"", b"changeit", b"password"):
                try:
                    _, cert, chain = _pkcs12.load_key_and_certificates(raw_bytes, pwd)
                    if cert:
                        sig_oid = cert.signature_algorithm_oid.dotted_string
                        pub = cert.public_key()
                        if isinstance(pub, _rsa.RSAPublicKey):
                            lines_out.append(f"RSA key_size={pub.key_size} signature_algorithm_oid={sig_oid}")
                        elif isinstance(pub, _ec.EllipticCurvePublicKey):
                            lines_out.append(f"ECDSA elliptic_curve={pub.curve.name} signature_algorithm_oid={sig_oid}")
                        else:
                            lines_out.append(f"algorithm signature_algorithm_oid={sig_oid}")
                    if chain:
                        for c in chain:
                            oid = c.signature_algorithm_oid.dotted_string
                            lines_out.append(f"chain certificate signature_algorithm_oid={oid}")
                    break
                except Exception:
                    continue
            return "\n".join(lines_out)
        except Exception:
            pass
        return "PKCS12 keystore ssl tls certificate algorithm encryption"

    # ── Java KeyStore (.jks) ──────────────────────────────────────────────────────
    if ext == ".jks":
        try:
            if raw_bytes[:4] == b'\xfe\xed\xfe\xed':
                text_layer = raw_bytes.decode("latin-1", errors="ignore")
                hints = []
                if "RSA" in text_layer: hints.append("RSA")
                if "EC"  in text_layer: hints.append("ECDSA elliptic curve")
                if "DSA" in text_layer: hints.append("DSA")
                if "AES" in text_layer: hints.append("AES")
                if not hints:           hints.append("algorithm")
                return (
                    "Java KeyStore JKS keystore ssl tls algorithm encryption key exchange\n"
                    + " ".join(hints) + "\n"
                )
        except Exception:
            pass
        return "Java KeyStore JKS keystore ssl tls algorithm encryption"

    # ── SSH Public Key Files (.pub) ───────────────────────────────────────────────
    if ext == ".pub":
        try:
            text = raw_bytes.decode("utf-8", errors="ignore").strip()
            _SSH_TYPE_MAP = {
                "ssh-rsa":                    "RSA ssh algorithm key exchange",
                "rsa-sha2-256":               "RSA ssh algorithm key exchange",
                "rsa-sha2-512":               "RSA ssh algorithm key exchange",
                "ssh-dss":                    "DSA ssh algorithm key exchange",
                "ecdsa-sha2-nistp256":        "ECDSA elliptic curve secp256r1 ssh algorithm",
                "ecdsa-sha2-nistp384":        "ECDSA elliptic curve secp384r1 ssh algorithm",
                "ecdsa-sha2-nistp521":        "ECDSA elliptic curve secp521r1 ssh algorithm",
                "ssh-ed25519":                "Ed25519 elliptic curve ssh algorithm",
                "ssh-ed448":                  "Ed448 elliptic curve ssh algorithm",
                "sk-ssh-ed25519@openssh.com": "Ed25519 elliptic curve fido2 ssh algorithm",
                "sk-ecdsa-sha2-nistp256@openssh.com": "ECDSA elliptic curve fido2 ssh algorithm",
            }
            for line in text.splitlines():
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                key_type = line.split()[0].lower()
                algo_desc = _SSH_TYPE_MAP.get(key_type, f"{key_type} ssh algorithm key exchange")
                return f"SSH public key {algo_desc}\ntls ssl cipher encryption\n"
        except Exception:
            pass
        return "SSH public key ssl tls algorithm"

    return ""


# Keywords whose intent is a prefix match ("cryptograph" must still catch
# "cryptography" and "cryptographic"), so they get a leading boundary only.
_PQC_PREFIX_KEYWORDS = frozenset({"cryptograph", "encrypt", "encrypted"})


def _count_pqc_signals(sample_lower: str) -> int:
    """Counts PQC-relevant keyword hits, matched on word boundaries.

    Plain substring matching (`kw in sample_lower`) made several of these
    three-letter entries fire on ordinary prose: "sha" matched "shall" and
    "shared", "rsa" matched "universal" and "reversal", "aes" matched "Caesar".
    Since can_parse() claims a file at only two hits, and the remaining keywords
    are generic ("encryption", "algorithm", "certificate"), any ISO policy
    document containing the word "shall" was one common noun away from being
    routed to the PQC parser.
    """
    if not sample_lower:
        return 0
    hits = 0
    for kw in _PQC_KEYWORDS:
        esc = re.escape(kw)
        pattern = (
            r'(?<![a-z0-9])' + esc
            if kw in _PQC_PREFIX_KEYWORDS
            else r'(?<![a-z0-9])' + esc + r'(?![a-z0-9])'
        )
        if re.search(pattern, sample_lower):
            hits += 1
    return hits


# ══════════════════════════════════════════════════════════════════════════════
# ALGORITHM LOOKUP TABLE
# ══════════════════════════════════════════════════════════════════════════════
# Each rule: (rule_id, regex, namer, quantum_status, crypto_category, severity)
#   rule_id         -- short stable key used to build a dedup-friendly plugin_id.
#   regex           -- compiled pattern matched against the raw evidence text.
#   namer           -- None (use the raw matched text as the display name),
#                       a fixed string, or a callable(match) -> str.
#   quantum_status  -- "VULNERABLE" | "WEAK" | "SAFE".
#   crypto_category -- human-readable category folded into title/description.
#   severity        -- fixed string, or callable(match) -> str for size/group-
#                       dependent rules (RSA key size, DH group number).

_SEV_INFO = "INFO"

# Rules that detect the *presence of TLS configuration* rather than a specific
# algorithm. Their evidence is a directive or a file path (e.g.
# "ssl-cert = /path/to/cert"), so titling them "Quantum-Vulnerable Algorithm
# Detected: ..." claims something the evidence does not support and reads as a
# false positive beside a real algorithm finding. They answer "is this service
# PQC-ready?", so they are titled as readiness gaps instead -- and suppressed
# outright when the same document already shows a NIST PQC algorithm, because
# the gap they assert ("TLS enabled with no post-quantum suite") is then false.
_CONFIG_PRESENCE_RULES = frozenset({
    "db-ssl-cert", "db-ssl-key", "db-tls13-no-pqc",
    "springboot-ssl-enabled", "java-keystore-pkcs12",
    "k8s-tls-secret", "k8s-ingress-tls",
    "kafka-ssl-keystore", "kafka-ssl-keystore-type", "kafka-ssl-protocol",
    "gnutls-no-pqc",
})

# Rule ids that identify an actual NIST-selected post-quantum algorithm.
# Finding one of these is what suppresses _CONFIG_PRESENCE_RULES above.
_PQC_ALGORITHM_RULES = frozenset({
    "kyber", "dilithium", "sphincs", "falcon", "pqc-hybrid-kem",
    # Set by the extended OID / IANA-suite / liboqs scan for a SAFE hit in a PQC
    # category, so PQC expressed only as an OID or suite code also counts.
    "pqc-extended-scan",
})


def _rsa_sized_severity(m: "re.Match") -> str:
    size_m = re.search(r'(\d{3,5})', m.group(0))
    if size_m:
        try:
            return "CRITICAL" if int(size_m.group(1)) < 3072 else "HIGH"
        except ValueError:
            pass
    return "CRITICAL"


def _rsa_sized_namer(m: "re.Match") -> str:
    size_m = re.search(r'(\d{3,5})', m.group(0))
    return f"RSA{size_m.group(1)}" if size_m else m.group(0)


def _dh_group_severity(m: "re.Match") -> str:
    grp_m = re.search(r'(\d{1,2})', m.group(0))
    if grp_m:
        try:
            return "CRITICAL" if int(grp_m.group(1)) <= 14 else "HIGH"
        except ValueError:
            pass
    return "CRITICAL"


def _dh_group_namer(m: "re.Match") -> str:
    grp_m = re.search(r'(\d{1,2})', m.group(0))
    return f"Diffie-Hellman Group {grp_m.group(1)}" if grp_m else "Diffie-Hellman Group"


ALGORITHM_RULES: List[Tuple[str, "re.Pattern", Union[None, str, Callable], str, str, Union[str, Callable]]] = [
    # ── QUANTUM-VULNERABLE: asymmetric / key-exchange (broken by Shor's algorithm) ──
    # rsa-sized: matches 'RSA 2048', 'RSA-2048', 'RSA2048', and also
    # OpenSSL key config style 'default_bits = 4096' without the RSA keyword.
    ("rsa-sized", re.compile(
        r'(?:\bRSA[\s\-_]?(?:512|1024|2048|3072|4096)\b'
        r'|\bdefault[_\-]bits\s*=\s*(?:512|1024|2048|3072|4096)\b)',
        re.IGNORECASE),
     _rsa_sized_namer, "VULNERABLE", "Asymmetric Encryption (RSA)", _rsa_sized_severity),
    ("rsa-generic", re.compile(r'\bRSA\b(?!\s*[\d])', re.IGNORECASE),
     "RSA (unspecified key size)", "VULNERABLE", "Asymmetric Encryption (RSA)", "CRITICAL"),
    # dsa: must NOT match the NIST PQC signature names that embed 'DSA' after a
    # hyphen -- ML-DSA (FIPS 204), SLH-DSA (FIPS 205), FN-DSA (Falcon). Plain
    # '\bDSA\b' matched all three, because '-' is a non-word character and so
    # forms a word boundary, raising a false CRITICAL "quantum-vulnerable DSA"
    # on exactly the algorithms a *completed* PQC migration is supposed to show.
    # The lookbehind rejects a preceding word char or hyphen; ECDSA stays
    # excluded by the leading \b as before.
    ("dsa", re.compile(r'(?<![\w-])DSA\b', re.IGNORECASE),
     "DSA", "VULNERABLE", "Asymmetric Digital Signature (DSA)", "CRITICAL"),
    ("dh-group", re.compile(r'\bDH\s*Group\s*(?:1[0-8]|[1-9])\b', re.IGNORECASE),
     _dh_group_namer, "VULNERABLE", "Key Exchange (Diffie-Hellman)", _dh_group_severity),
    ("dhe", re.compile(r'\bDHE\b', re.IGNORECASE),
     "DHE (Diffie-Hellman Ephemeral)", "VULNERABLE", "Key Exchange (Diffie-Hellman)", "HIGH"),
    ("dh-generic", re.compile(r'\bDiffie[\s\-]?Hellman\b', re.IGNORECASE),
     "Diffie-Hellman", "VULNERABLE", "Key Exchange (Diffie-Hellman)", "HIGH"),
    ("ecc-p256", re.compile(r'\b(?:P-256|secp256r1)\b', re.IGNORECASE),
     "ECC P-256 / secp256r1", "VULNERABLE", "Elliptic Curve Cryptography (ECC)", "HIGH"),
    ("ecc-p384", re.compile(r'\b(?:P-384|secp384r1)\b', re.IGNORECASE),
     "ECC P-384 / secp384r1", "VULNERABLE", "Elliptic Curve Cryptography (ECC)", "HIGH"),
    ("ecc-p521", re.compile(r'\b(?:P-521|secp521r1)\b', re.IGNORECASE),
     "ECC P-521 / secp521r1", "VULNERABLE", "Elliptic Curve Cryptography (ECC)", "HIGH"),
    ("ecc-secp256k1", re.compile(r'\bsecp256k1\b', re.IGNORECASE),
     "ECC secp256k1", "VULNERABLE", "Elliptic Curve Cryptography (ECC)", "HIGH"),
    ("ecc-x25519", re.compile(r'\b(?:Curve25519|X25519)\b', re.IGNORECASE),
     "Curve25519 / X25519", "VULNERABLE", "Elliptic Curve Cryptography (ECC)", "HIGH"),
    ("ecc-ed25519", re.compile(r'\bEd25519\b', re.IGNORECASE),
     "Ed25519", "VULNERABLE", "Elliptic Curve Cryptography (ECC)", "HIGH"),
    # ecdhe: the dominant quantum-vulnerable key exchange in real TLS configs,
    # and previously undetected entirely -- '\bDHE\b' cannot match it (preceded
    # by 'C') and '\bECDH\b' cannot either (followed by 'E'), so every
    # ECDHE-* cipher suite was reported with only its signature and symmetric
    # parts flagged while the key exchange itself stayed invisible. Listed
    # before ecdsa purely so key exchange reads before signature.
    ("ecdhe", re.compile(r'\bECDHE\b', re.IGNORECASE),
     "ECDHE (Elliptic Curve Diffie-Hellman Ephemeral)", "VULNERABLE",
     "Key Exchange (Elliptic Curve)", "HIGH"),
    ("ecdsa", re.compile(r'\bECDSA\b', re.IGNORECASE),
     "ECDSA", "VULNERABLE", "Elliptic Curve Cryptography (ECC)", "HIGH"),
    ("ecdh", re.compile(r'\bECDH\b', re.IGNORECASE),
     "ECDH", "VULNERABLE", "Elliptic Curve Cryptography (ECC)", "HIGH"),
    ("ecc-generic", re.compile(r'\bECC\b', re.IGNORECASE),
     "ECC (generic elliptic-curve reference)", "VULNERABLE", "Elliptic Curve Cryptography (ECC)", "HIGH"),

    # ── CLASSICALLY WEAK / DEPRECATED (not PQC-specific, still flagged) ──
    ("md5", re.compile(r'\bMD5\b', re.IGNORECASE),
     "MD5", "WEAK", "Hash Function", "HIGH"),
    # sha1: matches 'SHA-1', 'SHA1', 'SHA 1' (including fused forms like SHA1withRSA)
    ("sha1", re.compile(r'\bSHA[\s\-_]?1(?!\d)', re.IGNORECASE),
     "SHA-1", "WEAK", "Hash Function", "HIGH"),
    ("3des", re.compile(r'\b(?:3DES|Triple[\s\-]?DES|TripleDES)\b', re.IGNORECASE),
     "3DES / Triple DES", "WEAK", "Symmetric Cipher", "CRITICAL"),
    ("des", re.compile(r'\bDES\b', re.IGNORECASE),
     "DES", "WEAK", "Symmetric Cipher", "CRITICAL"),
    ("rc4", re.compile(r'\bRC4\b', re.IGNORECASE),
     "RC4", "WEAK", "Symmetric Cipher", "CRITICAL"),
    ("sslv2", re.compile(r'\bSSL\s?v?2\b', re.IGNORECASE),
     "SSLv2", "WEAK", "Protocol Version", "CRITICAL"),
    ("sslv3", re.compile(r'\bSSL\s?v?3\b', re.IGNORECASE),
     "SSLv3", "WEAK", "Protocol Version", "CRITICAL"),
    # tls10: matches 'TLSv1.0' and bare 'TLSv1' (no .0 suffix -- common in NGINX ssl_protocols)
    ("tls10", re.compile(r'\bTLS\s?v?1(?:\.0)?(?![\.\d])', re.IGNORECASE),
     "TLS 1.0", "WEAK", "Protocol Version", "CRITICAL"),
    ("tls11", re.compile(r'\bTLS\s?v?1\.1\b', re.IGNORECASE),
     "TLS 1.1", "WEAK", "Protocol Version", "CRITICAL"),
    ("cbc-mode", re.compile(r'\bCBC\b', re.IGNORECASE),
     "CBC-mode cipher (non-AEAD)", "WEAK", "Cipher Mode", "MEDIUM"),

    # ── QUANTUM-SAFE / NIST PQC-SELECTED ──
    # aes256-gcm: matches 'AES-256-GCM', 'AES_256_GCM', and TLS 1.3 suite name 'TLS_AES_256_GCM_SHA384'
    ("aes256-gcm", re.compile(
        r'(?:\bAES[\s\-_]?256[\s\-_]?GCM\b|(?<![A-Z0-9])AES_256_GCM(?![A-Z0-9])|TLS_AES_256_GCM)',
        re.IGNORECASE),
     "AES-256-GCM", "SAFE", "Symmetric Cipher (AEAD)", _SEV_INFO),
    ("aes256", re.compile(r'\bAES[\s\-_]?256\b', re.IGNORECASE),
     "AES-256", "SAFE", "Symmetric Cipher", _SEV_INFO),
    ("aes128", re.compile(r'\bAES[\s\-_]?128\b', re.IGNORECASE),
     "AES-128 (quantum-safe but AES-256 preferred)", "SAFE", "Symmetric Cipher", _SEV_INFO),
    ("sha384", re.compile(r'\bSHA[\s\-_]?384\b', re.IGNORECASE),
     "SHA-384", "SAFE", "Hash Function", _SEV_INFO),
    ("sha512", re.compile(r'\bSHA[\s\-_]?512\b', re.IGNORECASE),
     "SHA-512", "SAFE", "Hash Function", _SEV_INFO),
    ("sha256", re.compile(r'\bSHA[\s\-_]?256\b', re.IGNORECASE),
     "SHA-256 (acceptable minimum)", "SAFE", "Hash Function", _SEV_INFO),
    ("sha3", re.compile(r'\bSHA[\s\-_]?3\b', re.IGNORECASE),
     "SHA-3", "SAFE", "Hash Function", _SEV_INFO),
    # pqc-hybrid-kem: the *fused* hybrid group names that real TLS stacks
    # actually negotiate -- 'X25519MLKEM768' (the IETF/browser default),
    # 'X25519Kyber768Draft00', 'SecP256r1MLKEM768', 'X448MLKEM1024'. These were
    # invisible to the 'kyber' rule below, because '\bML[\s\-]?KEM\b' cannot
    # match inside 'X25519MLKEM768' -- the preceding '9' is a word character, so
    # there is no word boundary. The effect was that a server which had already
    # deployed hybrid post-quantum key exchange got no credit for it, the exact
    # mirror of the ML-DSA false positive. Must be listed before 'kyber' so the
    # fused form claims the span first.
    ("pqc-hybrid-kem", re.compile(
        r'\b(?:X25519|X448|SecP256r1|SecP384r1|P256|P384)'
        r'(?:[\s\-_]?(?:MLKEM|ML[\s\-]KEM|Kyber))'
        r'[\s\-_]?\d{3,4}(?:Draft\d+)?\b',
        re.IGNORECASE),
     None, "SAFE", "PQC Hybrid Key Exchange (NIST-selected, hybrid mode)", _SEV_INFO),
    ("kyber", re.compile(r'\b(?:CRYSTALS[\s\-]?Kyber|ML[\s\-]?KEM)\b', re.IGNORECASE),
     "CRYSTALS-Kyber / ML-KEM", "SAFE", "PQC Key Encapsulation (NIST-selected)", _SEV_INFO),
    ("dilithium", re.compile(r'\b(?:CRYSTALS[\s\-]?Dilithium|ML[\s\-]?DSA)\b', re.IGNORECASE),
     "CRYSTALS-Dilithium / ML-DSA", "SAFE", "PQC Digital Signature (NIST-selected)", _SEV_INFO),
    # sphincs: matches 'SPHINCS+', 'SPHINCS+-SHA2-128s', 'SLH-DSA'
    # Note: \b doesn't work after '+' so use a lookahead instead.
    ("sphincs", re.compile(r'\bSPHINCS[+]?(?:-SHA\d+)?\b|\bSLH[\s\-]?DSA\b', re.IGNORECASE),
     "SPHINCS+ / SLH-DSA", "SAFE", "PQC Digital Signature (NIST-selected)", _SEV_INFO),
    ("falcon", re.compile(r'\b(?:Falcon|FN[\s\-]?DSA)\b', re.IGNORECASE),
     "Falcon / FN-DSA", "SAFE", "PQC Digital Signature (NIST-selected)", _SEV_INFO),
    # chacha20: matches 'ChaCha20-Poly1305', 'ChaCha20_Poly1305', and TLS 1.3 cipher suite
    # 'TLS_CHACHA20_POLY1305_SHA256' (underscore form used in NGINX/HAProxy ssl_ciphers).
    ("chacha20", re.compile(
        r'(?:\bChaCha20[\s\-_]?Poly1305\b|TLS[_\-]CHACHA20[_\-]POLY1305|\bCHACHA20[_\-]POLY1305\b)',
        re.IGNORECASE),
     "ChaCha20-Poly1305", "SAFE", "Symmetric Cipher (AEAD)", _SEV_INFO),

    # ── DATABASE / SERVER TLS CONFIG PATTERNS ──────────────────────────────────
    # Catches MySQL/MariaDB/PostgreSQL config files that enable SSL/TLS transport
    # but do not specify a PQC-ready cipher suite.
    ("db-ssl-cert",
     re.compile(r'^\s*ssl[\-_]cert\s*=\s*.+', re.IGNORECASE | re.MULTILINE),
     "Database TLS Certificate",
     "VULNERABLE",
     "Database TLS Configuration",
     "HIGH"),
    ("db-ssl-key",
     re.compile(r'^\s*ssl[\-_]key\s*=\s*.+', re.IGNORECASE | re.MULTILINE),
     "Database TLS Key",
     "VULNERABLE",
     "Database TLS Configuration",
     "HIGH"),
    ("db-tls12",
     re.compile(r'\btls[_\-]?version\s*=\s*["\']?TLSv1\.2["\']?', re.IGNORECASE),
     "TLSv1.2",
     "VULNERABLE",
     "Protocol Version",
     "HIGH"),
    ("db-tls13-no-pqc",
     re.compile(r'\btls[_\-]?version\s*=\s*["\']?TLSv1\.3["\']?', re.IGNORECASE),
     "TLSv1.3",
     "VULNERABLE",
     "Protocol Version",
     "MEDIUM"),

    # ── SPRING BOOT / JAVA APPLICATION TLS ────────────────────────────────────
    ("springboot-ssl-enabled",
     re.compile(r'^\s*server\.ssl\.enabled\s*[=:]\s*true', re.IGNORECASE | re.MULTILINE),
     "Spring Boot TLS",
     "VULNERABLE", "Application TLS Configuration", "HIGH"),
    ("java-keystore-pkcs12",
     re.compile(r'^\s*server\.ssl\.key[\-_]store[\-_]type\s*[=:]\s*PKCS12', re.IGNORECASE | re.MULTILINE),
     "Java PKCS12 Keystore",
     "VULNERABLE", "Application TLS Configuration", "HIGH"),

    # ── KUBERNETES / CONTAINER TLS ────────────────────────────────────────────
    ("k8s-tls-secret",
     re.compile(r'kubernetes\.io/tls', re.IGNORECASE),
     "Kubernetes TLS Secret",
     "VULNERABLE", "Container / Cloud TLS Configuration", "HIGH"),
    ("k8s-ingress-tls",
     re.compile(r'^\s*tls:\s*$', re.IGNORECASE | re.MULTILINE),
     "Kubernetes Ingress TLS",
     "VULNERABLE", "Container / Cloud TLS Configuration", "MEDIUM"),

    # ── STRONGSWAN / IPSEC modpXXXX DH GROUPS ─────────────────────────────
    ("ipsec-modp",
     re.compile(r'\bmodp(1024|1536|2048|3072|4096|6144|8192)\b', re.IGNORECASE),
     lambda m: f"DH modp{m.group(1)}",
     "VULNERABLE", "Key Exchange (Diffie-Hellman)",
     lambda m: "CRITICAL" if int(m.group(1)) <= 2048 else "HIGH"),

    # ── GENERIC TLS VERSION IN ENV / YAML / TOML FORMATS ──────────────────────
    ("tls12-env-var",
     re.compile(r'TLS[_\-]?VERSION\s*[=:]\s*["\']?(?:TLSv?)?1\.2["\']?', re.IGNORECASE),
     "TLSv1.2",
     "VULNERABLE", "Protocol Version", "HIGH"),
    ("tls-min-version",
     re.compile(r'\bmin[_\-]?(?:tls[_\-]?)?version\s*[=:]\s*["\']?(?:TLSv?)?1\.2["\']?', re.IGNORECASE),
     "TLSv1.2",
     "VULNERABLE", "Protocol Version", "HIGH"),

    # ── APACHE KAFKA TLS CONFIGURATION ────────────────────────────────────
    ("kafka-ssl-keystore",
     re.compile(r'^\s*ssl\.keystore\.location\s*=\s*.+', re.IGNORECASE | re.MULTILINE),
     "Kafka Broker TLS",
     "VULNERABLE", "Application TLS Configuration", "HIGH"),
    ("kafka-ssl-keystore-type",
     re.compile(r'^\s*ssl\.keystore\.type\s*=\s*(?:JKS|PKCS12)', re.IGNORECASE | re.MULTILINE),
     "Kafka PKCS12 Keystore",
     "VULNERABLE", "Application TLS Configuration", "HIGH"),
    ("kafka-ssl-protocol",
     re.compile(r'^\s*ssl\.protocol\s*=\s*TLS', re.IGNORECASE | re.MULTILINE),
     "Kafka TLS",
     "VULNERABLE", "Application TLS Configuration", "MEDIUM"),
    ("kafka-ssl-enabled-protocols",
     re.compile(r'^\s*ssl\.enabled\.protocols\s*=\s*.*TLSv?1\.2', re.IGNORECASE | re.MULTILINE),
     "TLSv1.2",
     "VULNERABLE", "Protocol Version", "HIGH"),

    # ── GnuTLS PRIORITY STRINGS ───────────────────────────────────────────
    ("gnutls-tls12",
     re.compile(r'VERS-TLS1\.2', re.IGNORECASE),
     "TLSv1.2",
     "VULNERABLE", "Protocol Version", "HIGH"),
    ("gnutls-tls10",
     re.compile(r'VERS-TLS1\.0', re.IGNORECASE),
     "TLSv1.0",
     "WEAK", "Protocol Version", "CRITICAL"),
    ("gnutls-no-pqc",
     re.compile(r'GnuTLS[:\s].*(?:NORMAL|SECURE)', re.IGNORECASE),
     "GnuTLS Profile",
     "VULNERABLE", "Protocol Version", "MEDIUM"),

    # ── HAPROXY DH PARAM SIZE ───────────────────────────────────────────
    ("haproxy-dh-param",
     re.compile(r'tune\.ssl\.default-dh-param\s+(\d+)', re.IGNORECASE),
     lambda m: f"DH {m.group(1)}-bit",
     "VULNERABLE", "Key Exchange (Diffie-Hellman)",
     lambda m: "CRITICAL" if int(m.group(1)) <= 2048 else "HIGH"),

    # ── JVM / JDK SYSTEM PROPERTIES ───────────────────────────────────────────
    ("jvm-https-protocols",
     re.compile(r'-Dhttps?\.protocols\s*=\s*[\w,.]*TLSv?1\.2', re.IGNORECASE),
     "TLSv1.2",
     "VULNERABLE", "Protocol Version", "HIGH"),
    ("jvm-disabled-algorithms",
     re.compile(r'-Djdk\.tls\.disabledAlgorithms\s*=\s*\S+', re.IGNORECASE),
     "JVM disabledAlgorithms",
     "VULNERABLE", "Application TLS Configuration", "MEDIUM"),
    ("jvm-tls-version",
     re.compile(r'-Djavax\.net\.ssl\.(?:trustStore|keyStore)\w*\s*=\s*.+', re.IGNORECASE),
     "JVM JSSE Keystore",
     "VULNERABLE", "Application TLS Configuration", "HIGH"),

    # ── WINDOWS CRYPTOAPI / CNG REGISTRY EXPORTS (.reg) ──────────────────
    ("winreg-schannel-tls10",
     re.compile(r'SCHANNEL[\\]+Protocols[\\]+TLS 1\.0', re.IGNORECASE),
     "TLSv1.0",
     "WEAK", "Protocol Version", "CRITICAL"),
    ("winreg-schannel-tls12",
     re.compile(r'SCHANNEL[\\]+Protocols[\\]+TLS 1\.2', re.IGNORECASE),
     "TLSv1.2",
     "VULNERABLE", "Protocol Version", "HIGH"),
    # CNG RSA: matches HKEY path containing CNG and RSA (used in CNG algorithm provider entries)
    ("winreg-cng-rsa",
     re.compile(r'(?:HKEY_LOCAL_MACHINE|HKLM).*CNG.*RSA', re.IGNORECASE),
     "Windows CNG RSA provider registry entry (quantum-vulnerable asymmetric key)",
     "VULNERABLE", "Asymmetric Encryption (RSA)", "HIGH"),
]

# Per-algorithm precise remediation (overrides generic _REMEDIATION_VULNERABLE where matched).
# Key = lowercase substring of algo_name or crypto_category combined string.
# CRITICAL: order matters -- more specific keys MUST come before generic catch-alls.
# All ECC-family findings share crypto_category "Elliptic Curve Cryptography (ECC)",
# so curve25519 / x25519 / secp / ecdsa / ecdh must appear BEFORE "elliptic curve".
_REMEDIATION_BY_ALGO = [
    # ── Curve25519 / X25519 (most specific -- check before generic ECC) ───────
    ("curve25519",
        "Curve25519 / X25519 is quantum-vulnerable to Shor's algorithm.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Upgrade to hybrid PQC mode: X25519MLKEM768 (X25519 + ML-KEM-768, FIPS 203).\n"
        "     NGINX: ssl_ecdh_curve X25519MLKEM768:X25519;\n"
        "     (Provides both classical and post-quantum protection in one handshake -- "
        "'Harvest Now, Decrypt Later' safe.)\n"
        "  2. For SSH key exchange: add sntrup761x25519-sha512@openssh.com to "
        "sshd_config KexAlgorithms (OpenSSH 9.0+).\n"
        "  3. Long-term: migrate to ML-KEM-768 standalone (FIPS 203) once all clients support it.\n"
        "  NIST Reference: FIPS 203 (ML-KEM-768 = security level 3, equivalent to AES-192)."
    ),
    # ── X25519 (catches 'X25519' in algo_name without Curve25519 prefix) ─────
    ("x25519",
        "X25519 (Curve25519) is quantum-vulnerable to Shor's algorithm.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Upgrade to hybrid PQC: X25519MLKEM768 in NGINX ssl_ecdh_curve.\n"
        "     ssl_ecdh_curve X25519MLKEM768:X25519;\n"
        "  2. For SSH: use sntrup761x25519-sha512@openssh.com (OpenSSH 9.0+ hybrid KEX).\n"
        "  3. Long-term: ML-KEM-768 (FIPS 203) for full post-quantum key encapsulation.\n"
        "  NIST Reference: FIPS 203."
    ),
    # ── Named secp curves (secp256r1 / secp384r1 / secp521r1) ─────────────────
    ("secp",
        "Named elliptic curves (secp256r1 / secp384r1 / secp521r1) are quantum-vulnerable.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. For TLS key agreement: replace with X25519MLKEM768 hybrid as first preference.\n"
        "     NGINX: ssl_ecdh_curve X25519MLKEM768:X25519:prime256v1;\n"
        "  2. For TLS certificates using secp384r1: reissue with ML-DSA-65 (FIPS 204) "
        "once your CA supports post-quantum hybrid certs.\n"
        "  3. Track IETF TLS 1.3 hybrid key exchange drafts (draft-ietf-tls-hybrid-design) "
        "for NGINX/OpenSSL 3.x adoption timelines.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), FIPS 204 (ML-DSA), NIST SP 800-186."
    ),
    # ── ECDSA (signature scheme, separate from ECDH key exchange) ─────────────
    ("ecdsa",
        "ECDSA signatures are broken by Shor's algorithm -- the discrete logarithm "
        "over elliptic curves is efficiently solvable on a quantum computer.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Replace ECDSA TLS/SSH certificates with ML-DSA (FIPS 204) once your CA supports it.\n"
        "  2. Interim: ECDSA P-256/P-384 is still safe classically -- prioritise migration "
        "of long-lived certificates and code-signing keys first.\n"
        "  3. For NGINX cipher suites with ECDSA: ensure ssl_ecdh_curve includes "
        "X25519MLKEM768 to protect the key exchange layer even before cert migration.\n"
        "  4. For JWT / API tokens / code signing: migrate to ML-DSA-44 or ML-DSA-65.\n"
        "  5. For SSH host keys: switch to Ed25519 (interim) or ML-DSA-65 (long-term).\n"
        "  NIST Reference: FIPS 204 (ML-DSA replaces ECDSA), NIST IR 8413 (Falcon/FN-DSA)."
    ),
    # ── ECDH (key encapsulation / key agreement) ──────────────────────────────
    ("ecdh",
        "ECDH key agreement is broken by Shor's algorithm.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Replace ECDH with ML-KEM-768 (FIPS 203 / CRYSTALS-Kyber-768) for "
        "post-quantum key encapsulation.\n"
        "  2. Interim hybrid mode: X25519MLKEM768 in NGINX provides both classical "
        "and PQC protection simultaneously.\n"
        "     NGINX: ssl_ecdh_curve X25519MLKEM768:X25519;\n"
        "  3. For IPSec/IKEv2: add ML-KEM KEM groups in IKE SA proposals per RFC 9370.\n"
        "  4. For TLS 1.3 clients not yet supporting X25519MLKEM768: "
        "keep X25519 as fallback in ssl_ecdh_curve list.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), NIST SP 800-56C Rev 2."
    ),
    # ── Generic ECC catch-all (any other ECC not matched above) ───────────────
    ("elliptic curve",
        "Elliptic Curve Cryptography (ECC) is broken by Shor's algorithm. "
        "Both ECDH (key agreement) and ECDSA (signatures) must be migrated.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Key exchange: replace ECDH with ML-KEM-768 (FIPS 203). "
        "Hybrid interim: X25519MLKEM768.\n"
        "  2. Signatures: replace ECDSA with ML-DSA-65 (FIPS 204) or SLH-DSA (FIPS 205).\n"
        "  3. NGINX: ssl_ecdh_curve X25519MLKEM768:X25519; (hybrid PQC key exchange).\n"
        "  4. Reissue all ECC certificates when your CA supports post-quantum or hybrid certs.\n"
        "  NIST Reference: FIPS 203, FIPS 204, FIPS 205."
    ),
    # ── RSA ───────────────────────────────────────────────────────────────────
    ("rsa",
        "RSA is broken by Shor's algorithm on a sufficiently large quantum computer.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Inventory all RSA key usages: TLS certificates, SSH host keys, "
        "code-signing, JWT tokens, S/MIME.\n"
        "  2. For TLS key exchange: disable static-RSA cipher suites (non-ECDHE). "
        "ECDHE-RSA-* (PFS) is lower risk but still needs migration.\n"
        "     NGINX interim: ssl_ciphers ECDHE-ECDSA-AES256-GCM-SHA384:"
        "ECDHE-RSA-AES256-GCM-SHA384;\n"
        "  3. For TLS certificates: reissue as ECDSA P-256 (interim) then "
        "ML-DSA-65 (FIPS 204) when CA chains support it.\n"
        "  4. For SSH host/user keys: switch to Ed25519 (interim) or ML-DSA-65.\n"
        "  5. For code/firmware signing: move to ML-DSA-65 (FIPS 204) or SLH-DSA (FIPS 205).\n"
        "  6. Prioritise RSA-2048 and below for immediate migration; "
        "RSA-4096 has more runway but still requires planning.\n"
        "  NIST Reference: FIPS 204 (ML-DSA), FIPS 203 (ML-KEM), NIST SP 800-131A Rev 2."
    ),
    # ── DH Group (specific, before generic Diffie-Hellman) ───────────────────
    ("dh group",
        "DH group-based key exchange is quantum-vulnerable to Shor's algorithm.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Replace DHE groups with ML-KEM-768 (FIPS 203) for PQC key encapsulation.\n"
        "  2. For IKEv2/IPSec: add ML-KEM IKE groups (RFC 9370) to the SA proposal list.\n"
        "  3. Interim minimum: DH group 14 (2048-bit); prefer group 16 (4096-bit) or "
        "group 19 (ECDHE P-256) while migrating.\n"
        "  4. Disable DH groups 1, 2, 5, 22, 23, 24 (below 2048-bit) immediately -- "
        "these are classically weak too.\n"
        "  NIST Reference: FIPS 203, NIST SP 800-77 Rev 1."
    ),
    # ── Diffie-Hellman generic ────────────────────────────────────────────────
    ("diffie-hellman",
        "Diffie-Hellman key exchange is broken by Shor's algorithm.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Replace static DH cipher suites with ECDHE (interim) or ML-KEM-768 (long-term).\n"
        "  2. Disable DHE cipher suites below group 14 (2048-bit) immediately.\n"
        "  3. For IKEv2/IPSec: configure post-quantum KEM groups per RFC 9370.\n"
        "  4. Migrate to ML-KEM-768 (FIPS 203) for all new key exchange implementations.\n"
        "  NIST Reference: FIPS 203, NIST SP 800-56A Rev 3."
    ),
    # -- Database TLS config (no cipher suite / classical-only TLS) -----------
    ("database tls",
        "Database TLS configuration uses classical cryptography only and lacks a PQC-ready cipher suite.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Verify the database TLS certificate algorithm -- if RSA or ECDSA, plan migration to ML-DSA.\n"
        "  2. For MySQL 9.x / MariaDB: specify tls_ciphersuites using TLS 1.3 AEAD cipher suites:\n"
        "     tls_ciphersuites = TLS_AES_256_GCM_SHA384:TLS_CHACHA20_POLY1305_SHA256\n"
        "  3. Lock tls_version to TLSv1.3 only (remove TLSv1.2 if still listed):\n"
        "     tls_version = TLSv1.3\n"
        "  4. Verify ssl-cert references a certificate signed by a quantum-safe CA once available.\n"
        "  5. Track MySQL / MariaDB PQC roadmaps -- as of 2026, no mainstream DB engine ships\n"
        "     native ML-KEM/ML-DSA support; monitor OpenSSL 3.x + MySQL upstream announcements.\n"
        "  6. Apply network-layer controls (private VPC, mTLS, certificate pinning) as compensating\n"
        "     controls while awaiting PQC-capable DB engine releases.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), FIPS 204 (ML-DSA), NIST SP 800-52 Rev 2."
    ),
    # -- Spring Boot / Java application TLS -----------------------------------
    ("application tls",
        "Spring Boot / Java application TLS uses classical cryptography only.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Add explicit cipher suite configuration in application.properties:\n"
        "     server.ssl.ciphers=TLS_AES_256_GCM_SHA384,TLS_CHACHA20_POLY1305_SHA256\n"
        "     server.ssl.enabled-protocols=TLSv1.3\n"
        "  2. Audit the PKCS12 keystore for RSA/ECDSA certificate type:\n"
        "     keytool -list -keystore keystore.p12 -storetype PKCS12\n"
        "  3. Plan certificate migration to ML-DSA-65 (FIPS 204) when your CA supports it.\n"
        "  4. Use Java 21+ / BouncyCastle PQC provider for ML-KEM/ML-DSA prototype testing.\n"
        "  5. Track Spring Boot + OpenJDK PQC roadmaps for native ML-KEM TLS 1.3 support.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), FIPS 204 (ML-DSA), NIST SP 800-52 Rev 2."
    ),
    # -- Kubernetes / Container TLS -------------------------------------------
    ("container / cloud tls",
        "Kubernetes TLS Secrets and Ingress configurations use classical TLS certificates.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Audit all TLS Secrets: kubectl get secrets --field-selector type=kubernetes.io/tls\n"
        "  2. Check certificate algorithms: openssl x509 -in tls.crt -noout -text | grep 'Public Key'\n"
        "  3. For NGINX Ingress: set ssl-ciphers to TLS 1.3 AEAD suites in ConfigMap:\n"
        "     ssl-protocols: TLSv1.3\n"
        "     ssl-ciphers: TLS_AES_256_GCM_SHA384:TLS_CHACHA20_POLY1305_SHA256\n"
        "  4. Plan rotation to ML-DSA certificates when cert-manager + ACME support PQC.\n"
        "  5. Apply mTLS (mutual TLS) within the cluster as a compensating control.\n"
        "  6. Monitor CNCF TAG Security and cert-manager PQC roadmap for ML-KEM/ML-DSA support.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), FIPS 204 (ML-DSA), NIST SP 800-52 Rev 2."
    ),
    # -- StrongSwan / IPSec modp DH groups ------------------------------------
    ("ikev2/ipsec dh",
        "IKEv2/IPSec configuration uses classical Diffie-Hellman (modpXXXX) key exchange "
        "which is broken by Shor's algorithm.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Add ML-KEM (FIPS 203) IKE groups to the proposal per RFC 9370:\n"
        "     ike=aes256gcm16-prfsha384-kyber3-ecp384!\n"
        "     (StrongSwan >= 6.0 with wolfSSL or OpenQuantumSafe provider)\n"
        "  2. Interim: upgrade to modp4096 or ecp384 minimum while awaiting PQC.\n"
        "     Remove modp1024/modp2048 immediately -- these are also classically weak.\n"
        "  3. For libreswan: set ikev2_allow_narrowing=yes and enable PQC KEM groups.\n"
        "  4. Check StrongSwan release notes for liboqs plugin supporting ML-KEM.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), RFC 9370, NIST SP 800-77 Rev 1."
    ),
    # -- Kafka Broker TLS ---------------------------------------------------
    ("kafka broker tls",
        "Apache Kafka broker uses classical TLS without a PQC-ready cipher suite.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Upgrade Kafka keystore to use an ML-DSA certificate when your CA supports it.\n"
        "  2. Lock TLS to 1.3 only in broker.properties:\n"
        "     ssl.enabled.protocols=TLSv1.3\n"
        "  3. Set cipher suites (Kafka 3.x+):\n"
        "     ssl.cipher.suites=TLS_AES_256_GCM_SHA384,TLS_CHACHA20_POLY1305_SHA256\n"
        "  4. Prefer ECDHE over static DH: ssl.keystore.type=PKCS12\n"
        "  5. Monitor Apache Kafka / Apache Flink PQC roadmap for ML-KEM KEM group support.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), FIPS 204 (ML-DSA), NIST SP 800-52 Rev 2."
    ),
    # -- GnuTLS priority strings -------------------------------------------
    ("gnutls",
        "GnuTLS priority string specifies classical TLS version without PQC KEM support.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Update GnuTLS priority string to use TLS 1.3 only:\n"
        "     NORMAL:-VERS-TLS1.0:-VERS-TLS1.1:-VERS-TLS1.2\n"
        "  2. Monitor GnuTLS roadmap for ML-KEM/ML-DSA key exchange group support.\n"
        "  3. Use GnuTLS >= 3.8.x for TLS 1.3 with X25519 (interim, not yet PQC-safe).\n"
        "  4. For libgnutls-based apps: set GNUTLS_SYSTEM_PRIORITY_FILE to enforce TLS 1.3.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), FIPS 204 (ML-DSA), NIST SP 800-52 Rev 2."
    ),
    # -- HAProxy static DH param -------------------------------------------
    ("haproxy static dh",
        "HAProxy is configured with a static Diffie-Hellman parameter (tune.ssl.default-dh-param).\n"
        "Static DH (even at 4096-bit) is quantum-vulnerable via Shor's algorithm.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Remove tune.ssl.default-dh-param and switch to ECDHE cipher suites:\n"
        "     ssl-default-bind-options no-sslv3 no-tlsv10 no-tlsv11 no-tlsv12\n"
        "     ssl-default-bind-ciphers TLS_AES_256_GCM_SHA384:TLS_CHACHA20_POLY1305_SHA256\n"
        "  2. Disable DHE in favor of ECDHE to avoid static DH entirely:\n"
        "     ssl-default-bind-ciphers ECDHE:!aNULL:!eNULL:!EXPORT:!DES:!RC4:!3DES:!MD5:!PSK\n"
        "  3. Monitor HAProxy + OpenSSL PQC roadmap for ML-KEM KEM group support.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), NIST SP 800-77 Rev 1."
    ),
    # -- JVM JSSE system properties ----------------------------------------
    ("jvm",
        "JVM JSSE system properties configure classical TLS without PQC algorithm support.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Update https.protocols to TLSv1.3 only:\n"
        "     -Dhttps.protocols=TLSv1.3\n"
        "  2. Review jdk.tls.disabledAlgorithms - ensure RSA < 2048, DH < 2048 are disabled:\n"
        "     jdk.tls.disabledAlgorithms=SSLv3,TLSv1,TLSv1.1,RC4,DES,MD5withRSA\n"
        "  3. Use Java 21+ with BouncyCastle PQC provider for ML-KEM/ML-DSA prototype testing.\n"
        "  4. Migrate JSSE keystore to ML-DSA certificate when your CA supports FIPS 204.\n"
        "  5. Monitor OpenJDK PQC roadmap (JEP 496 - ML-KEM) for native ML-KEM TLS support.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), FIPS 204 (ML-DSA), NIST SP 800-52 Rev 2."
    ),
    # -- Windows Schannel / CNG registry -----------------------------------
    ("windows schannel",
        "Windows Schannel TLS registry configuration uses classical cryptography only.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Disable TLS 1.0 and TLS 1.1 in Schannel registry:\n"
        "     HKLM\\SYSTEM\\CurrentControlSet\\Control\\SecurityProviders\\SCHANNEL\\Protocols\\TLS 1.0\\Server - Enabled = 0\n"
        "  2. Use IISCrypto or PowerShell DSC to enforce TLS 1.3 only on Windows Server 2022+.\n"
        "  3. Plan migration of Windows certificate store to ML-DSA when Microsoft CA supports FIPS 204.\n"
        "  4. Monitor Microsoft CNG PQC roadmap for ML-KEM/ML-DSA provider availability.\n"
        "  NIST Reference: FIPS 203 (ML-KEM), FIPS 204 (ML-DSA), NIST SP 800-52 Rev 2."
    ),
    # -- Windows CNG registry RSA ------------------------------------------
    ("windows cng rsa",
        "Windows CNG RSA provider registry entry detected - quantum-vulnerable asymmetric key.\n"
        "IMMEDIATE ACTIONS:\n"
        "  1. Audit all Windows certificates using RSA: certutil -store My\n"
        "  2. Plan migration to ML-DSA (FIPS 204) certificates when Microsoft CA supports them.\n"
        "  3. Use Windows Hello for Business or Azure Key Vault for interim key management.\n"
        "  4. Monitor Microsoft CNG PQC roadmap for ML-KEM/ML-DSA provider.\n"
        "  NIST Reference: FIPS 204 (ML-DSA), NIST IR 8413."
    ),
]


def _get_remediation_vulnerable(algo_name: str, crypto_category: str) -> str:
    """Return the most specific available remediation for a VULNERABLE algorithm.
    Uses an ordered list (not dict) so that more-specific entries are always checked first.
    Falls back to the generic _REMEDIATION_VULNERABLE if no specific entry matches.

    Matching: checks if the key string is a substring of the combined
    'algo_name + crypto_category' (lower-cased). All ECC findings share
    crypto_category 'Elliptic Curve Cryptography (ECC)', so algorithm-specific
    keys (curve25519, x25519, secp, ecdsa, ecdh) must appear before 'elliptic curve'
    in the list -- which they do."""
    combined = f"{algo_name} {crypto_category}".lower()
    for key, text in _REMEDIATION_BY_ALGO:
        if key in combined:
            return text
    return _REMEDIATION_VULNERABLE


_REMEDIATION_VULNERABLE = (
    "This algorithm is broken by Shor's algorithm on a sufficiently large quantum computer. "
    "Inventory all usages and plan migration to NIST-selected post-quantum algorithms: "
    "ML-KEM (FIPS 203) for key exchange / encapsulation, ML-DSA (FIPS 204) for digital signatures, "
    "or SLH-DSA (FIPS 205) as an alternative signature scheme. "
    "NIST Reference: FIPS 203, FIPS 204, FIPS 205, NIST IR 8413."
)
_REMEDIATION_WEAK = (
    "Disable this deprecated/weak cryptographic algorithm or protocol version immediately and "
    "replace it with a modern, non-deprecated alternative (AES-256-GCM, SHA-384+, TLS 1.2+ with "
    "strong cipher suites). This is independent of quantum readiness -- it is already breakable "
    "with classical computing -- but should be tracked alongside the broader PQC migration plan."
)
_REMEDIATION_SAFE = (
    "No quantum-readiness action required for this algorithm today. Continue monitoring NIST PQC "
    "guidance and prefer AES-256/SHA-384+ over smaller-margin variants (e.g. AES-128, SHA-256) "
    "where practical, since Grover's algorithm only gives a quadratic quantum speed-up against them."
)


# ══════════════════════════════════════════════════════════════════════════════
#  STANDARDS REFERENCES -- impacted FIPS, NIST SP 800-53 controls, and CVEs
# ══════════════════════════════════════════════════════════════════════════════
# Auditor feedback: a finding that says "NIST Reference: FIPS 203, FIPS 204,
# FIPS 205" tells the reader nothing -- it lists every PQC standard regardless
# of what was actually found. The description has to name the ONE standard that
# replaces the primitive in this finding.
#
# The mapping is by primitive, and the distinction that matters is:
#
#   FIPS 203  ML-KEM    replaces KEY ESTABLISHMENT  (RSA key transport, DH, ECDH)
#   FIPS 204  ML-DSA    replaces SIGNATURES         (RSA sig, DSA, ECDSA, EdDSA)
#   FIPS 205  SLH-DSA   alternative signature scheme, hash-based
#   FIPS 206  FN-DSA    still DRAFT -- never cited as a remediation target
#
# Symmetric ciphers and hash functions are deliberately absent. Grover's
# algorithm only halves their effective strength, so they are addressed by
# larger parameters (AES-256, SHA-384+) under CNSA 2.0 -- NOT by any of the
# FIPS 203-206 standards. Naming a PQC standard on an RC4 or MD5 finding would
# be precisely the imprecision this table exists to remove.
_FIPS_BY_CATEGORY = {
    "Asymmetric Encryption (RSA)":
        "FIPS 203 (ML-KEM) where RSA performs key establishment, and FIPS 204 "
        "(ML-DSA) where it performs digital signatures",
    "Asymmetric Digital Signature (DSA)":
        "FIPS 204 (ML-DSA), with FIPS 205 (SLH-DSA) as the hash-based alternative",
    "Key Exchange (Diffie-Hellman)":            "FIPS 203 (ML-KEM)",
    "Key Exchange (Elliptic Curve)":            "FIPS 203 (ML-KEM)",
    "Elliptic Curve Cryptography (ECC)":
        "FIPS 203 (ML-KEM) where the curve performs key agreement, and FIPS 204 "
        "(ML-DSA) where it performs signatures",
    "Database TLS Configuration":               "FIPS 203 (ML-KEM) for the TLS handshake",
    "Application TLS Configuration":            "FIPS 203 (ML-KEM) for the TLS handshake",
    "Container / Cloud TLS Configuration":      "FIPS 203 (ML-KEM) for the TLS handshake",
    "Protocol Version":                         "FIPS 203 (ML-KEM) for the TLS handshake",
}

# Rule-level overrides, where the curve's ROLE is unambiguous and the broader
# "Elliptic Curve Cryptography (ECC)" category answer would be vaguer than it
# needs to be. X25519 only ever does key agreement; Ed25519 only ever signs.
_FIPS_BY_RULE = {
    "ecc-x25519": "FIPS 203 (ML-KEM)",
    "ecc-ed25519": "FIPS 204 (ML-DSA)",
    "ecdh":        "FIPS 203 (ML-KEM)",
    "ecdhe":       "FIPS 203 (ML-KEM)",
    "ecdsa":       "FIPS 204 (ML-DSA)",
}

# NIST SP 800-53 Rev 5 controls each finding bears on. SC-13 (Cryptographic
# Protection) applies to every cryptographic finding; the rest are what make
# the reference useful to an assessor rather than boilerplate.
_NIST_80053_BY_CATEGORY = {
    "Asymmetric Encryption (RSA)":          "SC-12, SC-13",
    "Asymmetric Digital Signature (DSA)":   "SC-13, SC-17, IA-7",
    "Key Exchange (Diffie-Hellman)":        "SC-12, SC-13",
    "Key Exchange (Elliptic Curve)":        "SC-12, SC-13",
    "Elliptic Curve Cryptography (ECC)":    "SC-12, SC-13, SC-17",
    "Hash Function":                        "SC-13",
    "Symmetric Cipher":                     "SC-13, SC-28",
    "Symmetric Cipher (AEAD)":              "SC-13, SC-28",
    "Cipher Mode":                          "SC-8(1), SC-13",
    "Protocol Version":                     "SC-8, SC-8(1), SC-13",
    "Database TLS Configuration":           "SC-8, SC-8(1), SC-12, SC-13",
    "Application TLS Configuration":        "SC-8, SC-8(1), SC-12, SC-13",
    "Container / Cloud TLS Configuration":  "SC-8, SC-8(1), SC-12, SC-13",
}
_NIST_80053_DEFAULT = "SC-13"

# Real, published CVEs only.
#
# Being quantum-vulnerable is NOT a CVE: there is no CVE for "RSA-2048 will be
# broken by Shor's algorithm", and inventing one would put a fabricated
# identifier into an audit report. So this table covers only the CLASSICAL
# breaks that carry a genuine CVE, and every quantum-only finding leaves
# cve_list empty -- which bg_worker already renders as "No CVE assigned".
_CVE_BY_RULE = {
    "md5":     ["CVE-2004-2761"],                    # MD5 collision -> cert forgery
    "sha1":    ["CVE-2005-4900"],                    # SHA-1 collision
    "3des":    ["CVE-2016-2183"],                    # SWEET32, 64-bit block birthday
    "des":     ["CVE-2016-2183"],                    # SWEET32
    "rc4":     ["CVE-2013-2566", "CVE-2015-2808"],   # RC4 biases; Bar Mitzvah
    "sslv2":   ["CVE-2016-0800"],                    # DROWN
    "sslv3":   ["CVE-2014-3566"],                    # POODLE
    "tls10":   ["CVE-2011-3389"],                    # BEAST
    "gnutls-tls10":       ["CVE-2011-3389"],
    "winreg-schannel-tls10": ["CVE-2011-3389"],
    "cbc-mode": ["CVE-2013-0169"],                   # Lucky Thirteen
}

# Logjam applies to 512/768/1024-bit groups, not to every Diffie-Hellman group.
# Attaching it to DH Group 14 (2048-bit) would be a false citation, so the
# group number decides -- the same callable idiom ALGORITHM_RULES already uses
# for its namers and severities.
_LOGJAM_GROUPS = {"1", "2"}          # Group 1 = 768-bit, Group 2 = 1024-bit


def _dh_group_cves(matched_text: str):
    m = re.search(r'(\d{1,2})', matched_text or "")
    if m and m.group(1) in _LOGJAM_GROUPS:
        return ["CVE-2015-4000"]     # Logjam
    return []


def _cves_for(rule_id: str, matched_text: str):
    """Published CVEs for this rule, or [] when the weakness is quantum-only."""
    if rule_id == "dh-group":
        return _dh_group_cves(matched_text)
    return list(_CVE_BY_RULE.get(rule_id, ()))


def _fips_impacted(rule_id: str, crypto_category: str) -> str:
    """The PQC standard that replaces THIS primitive, or '' when none does."""
    return _FIPS_BY_RULE.get(rule_id) or _FIPS_BY_CATEGORY.get(crypto_category, "")


def _nist_80053_for(crypto_category: str) -> str:
    return _NIST_80053_BY_CATEGORY.get(crypto_category, _NIST_80053_DEFAULT)


def _resolve(value, m):
    if value is None:
        return m.group(0)
    if callable(value):
        return value(m)
    return value


# ══════════════════════════════════════════════════════════════════════════════
# TARGET / ASSET CONTEXT EXTRACTION (best-effort)
# ══════════════════════════════════════════════════════════════════════════════

_ASSET_HEADING_RE = re.compile(
    r'(?im)^\s*(?:host|target|system|asset|server|hostname|device|node|ip\s*address)\s*[:=]\s*(.+)$'
)


def _find_asset_context(content: str, match_start: int, filename: str) -> str:
    """Best-effort surrounding asset/heading context for a match -- looks
    backward up to ~800 chars for the nearest 'Host:'/'Target:'/'System:'-style
    line. Falls back to the filename when no such heading is found."""
    window_start = max(0, match_start - 800)
    window = content[window_start:match_start]
    hits = list(_ASSET_HEADING_RE.finditer(window))
    if hits:
        return hits[-1].group(0).strip()
    return filename


# Comment markers across the config formats this parser sees: nginx/apache/yaml/
# shell use '#', ini/php use ';', json5/js/java use '//', SQL uses '--', XML uses
# '<!--'. Only a marker at the START of the line makes the whole line a comment.
_COMMENT_MARKERS = ("#", "//", ";", "--", "<!--", "!")


def _is_comment_line(line: str) -> bool:
    """True when the line is entirely commented out.

    A trailing comment on a live directive ("ssl_ciphers ...;  # legacy") is NOT a
    comment line -- the directive is still in force and must still be assessed.
    """
    stripped = (line or "").strip()
    return bool(stripped) and stripped.startswith(_COMMENT_MARKERS)


def _line_containing(content: str, start: int, end: int) -> str:
    """Returns the full line(s) containing [start:end), guaranteed to be a real
    substring of `content` (grounding requirement) -- never paraphrased."""
    line_start = content.rfind("\n", 0, start) + 1
    line_end = content.find("\n", end)
    if line_end == -1:
        line_end = len(content)
    return content[line_start:line_end].strip() or content[start:end]


# ══════════════════════════════════════════════════════════════════════════════
# CA / KEY / PROTOCOL LAYER CLASSIFICATION (best-effort, Enhancement 1)
# ══════════════════════════════════════════════════════════════════════════════

_CA_LABELS = ("certificate", "ca ", "ca:", "signature algorithm", "issued by", "x.509", "x509")
_KEY_LABELS = ("key exchange", "key size", "key algorithm", "kex", "public key", "private key", "key:")
_PROTOCOL_LABELS = (
    "protocol", "tls version", "ssl version", "ike version", "ipsec phase",
    "ssh version", "phase 1", "phase 2",
)


def _extract_protocol_version(content: str) -> str:
    """Extract protocol version string (e.g. 'TLSv1.2 / TLSv1.3') from configuration text."""
    if not content: return ""
    c_lower = content.lower()
    found = []
    if "tlsv1.3" in c_lower or "tls 1.3" in c_lower or "tls1.3" in c_lower:
        found.append("TLSv1.3")
    if "tlsv1.2" in c_lower or "tls 1.2" in c_lower or "tls1.2" in c_lower:
        found.append("TLSv1.2")
    if "tlsv1.1" in c_lower or "tls 1.1" in c_lower or "tls1.1" in c_lower:
        found.append("TLSv1.1")
    if "tlsv1.0" in c_lower or "tls 1.0" in c_lower or "tls1.0" in c_lower:
        found.append("TLSv1.0")
    if "sslv3" in c_lower or "ssl 3.0" in c_lower:
        found.append("SSLv3")
    return " / ".join(found) if found else ""


def _assign_crypto_layers(finding: Finding, algo_name: str, crypto_category: str, meta: Optional[dict] = None, content: str = ""):
    """
    Accurately populates ca_algorithm, key_algorithm, and protocol_version
    on a finding according to the algorithm's actual cryptographic role.
    """
    name_up = (algo_name or "").strip().upper()
    cat_up  = (crypto_category or "").strip().upper()

    # 1. Direct metadata from IANA / OID DB
    if meta:
        kex = meta.get("kex", "")
        auth = meta.get("auth", "")
        tls_ver = meta.get("tls_version", "")
        if kex and kex not in ("N/A", "UNKNOWN", ""):
            finding.key_algorithm = kex
        if auth and auth not in ("N/A", "UNKNOWN", ""):
            finding.ca_algorithm = auth
        if tls_ver and tls_ver not in ("N/A", "UNKNOWN", ""):
            finding.protocol_version = tls_ver

    # 2. Protocol versions
    if any(p in name_up or p in cat_up for p in (
        "TLSV1", "TLS 1", "SSLV", "SSL 3", "IKEV", "IPSEC", "SSH"
    )) and not any(k in name_up for k in ("AES", "ECDSA", "RSA", "SHA", "ECDHE", "DHE", "CHACHA", "X25519", "P-384")):
        if not finding.protocol_version:
            finding.protocol_version = algo_name
        return

    # 3. Signature / CA / Certificate Algorithms
    sig_patterns = (
        "ECDSA", "RSASSA", "RSA", "DSA", "ED25519", "ED448",
        "ML-DSA", "DILITHIUM", "SLH-DSA", "SPHINCS", "FALCON", "FN-DSA",
        "SHA256WITHRSA", "SHA384WITHRSA", "SHA512WITHRSA", "SHA1WITHRSA"
    )
    if any(sig in name_up or sig in cat_up for sig in sig_patterns):
        if "ECDSA" in name_up: sig_name = "ECDSA"
        elif "RSA" in name_up: sig_name = "RSA"
        elif "DSA" in name_up and "ECDSA" not in name_up: sig_name = "DSA"
        elif "ED25519" in name_up: sig_name = "Ed25519"
        elif "DILITHIUM" in name_up or "ML-DSA" in name_up: sig_name = "ML-DSA (Dilithium)"
        elif "SPHINCS" in name_up or "SLH-DSA" in name_up: sig_name = "SLH-DSA (SPHINCS+)"
        elif "FALCON" in name_up or "FN-DSA" in name_up: sig_name = "FN-DSA (Falcon)"
        else: sig_name = algo_name

        if not finding.ca_algorithm:
            finding.ca_algorithm = sig_name

    # 4. Key Exchange / Key Encapsulation / Curve Algorithms
    kex_patterns = (
        "ECDHE", "DHE", "DIFFIE-HELLMAN", "X25519", "X448", "P-256", "P-384", "P-521",
        "SECP256", "SECP384", "SECP521", "CURVE25519", "MODP", "ML-KEM", "KYBER",
        "FRODOKEM", "HQC", "BIKE", "MCELIECE", "KEM", "ELLIPTIC CURVE", "RSA"
    )
    if any(kex in name_up or kex in cat_up for kex in kex_patterns):
        if "X25519" in name_up or "CURVE25519" in name_up: kex_name = "Curve25519 / X25519"
        elif "P-384" in name_up or "SECP384" in name_up: kex_name = "ECC P-384 / secp384r1"
        elif "P-256" in name_up or "SECP256" in name_up: kex_name = "ECC P-256 / secp256r1"
        elif "ECDHE" in name_up: kex_name = "ECDHE"
        elif "DHE" in name_up or "DIFFIE" in name_up: kex_name = "DHE"
        elif "ML-KEM" in name_up or "KYBER" in name_up: kex_name = "ML-KEM (Kyber)"
        elif "FRODOKEM" in name_up: kex_name = "FrodoKEM"
        elif "HQC" in name_up: kex_name = "HQC"
        elif "BIKE" in name_up: kex_name = "BIKE"
        elif "MCELIECE" in name_up: kex_name = "Classic McEliece"
        elif "RSA" in name_up: kex_name = "RSA"
        else: kex_name = algo_name

        if not finding.key_algorithm:
            finding.key_algorithm = kex_name

    # 5. Clean up protocol_version if it was filled with non-protocol algorithm name
    if finding.protocol_version and not any(p in finding.protocol_version.upper() for p in ("TLS", "SSL", "SSH", "IKE", "IPSEC")):
        finding.protocol_version = ""

    # 6. Fallback: extract protocol version from content context
    if not finding.protocol_version and content:
        finding.protocol_version = _extract_protocol_version(content)


def _classify_crypto_layer(content: str, match_start: int, match_end: int) -> str:
    """Best-effort classification of which crypto-config "layer" (CA/KEY/
    PROTOCOL) a match belongs to, based on the nearest preceding field label on
    the same line or within ~120 chars before the match. Returns "" if no
    recognizable label is found nearby -- never guessed."""
    window_start = max(0, match_start - 120)
    window = content[window_start:match_start].lower()

    def _last_label_index(labels):
        best = -1
        for label in labels:
            idx = window.rfind(label)
            if idx > best:
                best = idx
        return best

    ca_idx = _last_label_index(_CA_LABELS)
    key_idx = _last_label_index(_KEY_LABELS)
    proto_idx = _last_label_index(_PROTOCOL_LABELS)

    best_layer = ""
    best_idx = -1
    for layer, idx in (("CA", ca_idx), ("KEY", key_idx), ("PROTOCOL", proto_idx)):
        if idx > best_idx:
            best_idx = idx
            best_layer = layer

    if best_idx == -1:
        return ""
    return best_layer


# ══════════════════════════════════════════════════════════════════════════════
# EXPOSURE CONTEXT CLASSIFICATION (best-effort, Enhancement 2)
# ══════════════════════════════════════════════════════════════════════════════

_EXTERNAL_SIGNALS = (
    "external", "internet-facing", "internet facing", "public", "dmz",
    "wan-facing", "external-facing", "perimeter",
)
_INTERNAL_SIGNALS = (
    "internal", "internal-only", "lan", "intranet", "private network", "on-prem",
)
# Word-boundary compiled versions of the above -- bare `sig in search_text`
# containment let short signals like "lan" false-match inside unrelated words
# (e.g. "Load Balancer" contains "lan" inside "baLANcer"), wrongly flipping a
# genuinely EXTERNAL-exposed asset to unclassified/INTERNAL and silently
# depressing its HNDL risk factor. Same class of bug already guarded against
# elsewhere in this codebase (see get_actionable_remediation()'s word-boundary
# note in this file). Phrases with spaces (e.g. "internet facing") still match
# correctly since \b anchors on the whole phrase's start/end.
_EXTERNAL_SIGNAL_RE = re.compile(
    r'\b(?:' + '|'.join(re.escape(s) for s in _EXTERNAL_SIGNALS) + r')\b', re.IGNORECASE
)
_INTERNAL_SIGNAL_RE = re.compile(
    r'\b(?:' + '|'.join(re.escape(s) for s in _INTERNAL_SIGNALS) + r')\b', re.IGNORECASE
)


def _classify_exposure(content: str, match_start: int, asset_ctx: str, filename: str) -> str:
    """Best-effort EXTERNAL/INTERNAL exposure classification for a match, based
    on nearby context (asset context line + a ~300-char window around the
    match + filename). Returns "" if both/neither found."""
    window_start = max(0, match_start - 300)
    window_end = min(len(content), match_start + 300)
    search_text = (
        (asset_ctx or "") + " " + content[window_start:window_end] + " " + (filename or "")
    ).lower()

    has_external = bool(_EXTERNAL_SIGNAL_RE.search(search_text))
    has_internal = bool(_INTERNAL_SIGNAL_RE.search(search_text))

    if has_external and not has_internal:
        return "EXTERNAL"
    if has_internal and not has_external:
        return "INTERNAL"
    return ""


# ══════════════════════════════════════════════════════════════════════════════
# PORT EXTRACTION (best-effort, Enhancement 3)
# ══════════════════════════════════════════════════════════════════════════════

_PORT_LABEL_RE = re.compile(r'\bport\s*[:=]?\s*(\d{1,5})\b', re.IGNORECASE)
_PORT_SUFFIX_RE = re.compile(r':(\d{1,5})\b')


def _find_port(content: str, match_start: int, match_end: int) -> str:
    """Best-effort nearby port number for a match -- searches the containing
    line plus a ~200-char window before/after for a 'Port: NNN' style label or
    a ':NNN' suffix immediately following a hostname/IP-looking token. Returns
    "" if no such reference is nearby -- never a fabricated/default port."""
    window_start = max(0, match_start - 200)
    window_end = min(len(content), match_end + 200)
    line = _line_containing(content, match_start, match_end)
    window = content[window_start:window_end]

    for text in (line, window):
        m = _PORT_LABEL_RE.search(text)
        if m:
            return m.group(1)

    for text in (line, window):
        m = _PORT_SUFFIX_RE.search(text)
        if m:
            return m.group(1)

    return ""


# ══════════════════════════════════════════════════════════════════════════════
# ENVIRONMENT CLASSIFICATION (best-effort, Enhancement 4)
# ══════════════════════════════════════════════════════════════════════════════

_PROD_RE = re.compile(r'\b(?:prod|production)\b', re.IGNORECASE)
_NON_PROD_RE = re.compile(
    r'\b(?:dev|development|test|testing|staging|uat|qa|sandbox)\b', re.IGNORECASE
)


def _classify_environment(content: str, match_start: int, asset_ctx: str, filename: str) -> str:
    """Best-effort PROD/NON_PROD environment classification for a match, based
    on asset context + filename + a nearby content window. Word-boundary
    matched to avoid false positives (e.g. "product" must not match "prod").
    Returns "" if both/neither found."""
    window_start = max(0, match_start - 300)
    window_end = min(len(content), match_start + 300)
    search_text = (
        (asset_ctx or "") + " " + content[window_start:window_end] + " " + (filename or "")
    )

    has_prod = bool(_PROD_RE.search(search_text))
    has_non_prod = bool(_NON_PROD_RE.search(search_text))

    if has_prod and not has_non_prod:
        return "PROD"
    if has_non_prod and not has_prod:
        return "NON_PROD"
    return ""


# ── Asset Category keyword maps (Whiteboard / RFP schema) ────────────────────
# Order matters: more specific categories are checked first so that e.g.
# "palo alto firewall" matches "Firewall" before the generic "Server" catch.
_ASSET_CATEGORY_RULES = [
    # Category label        Keyword signals (any match -> this category)
    # ORDER MATTERS: more specific/less ambiguous categories first.
    ("Firewall",            ("firewall", "palo alto", "pan-os", "fortigate", "fortinet",
                             "checkpoint", "check point", "cisco asa", "cisco firepower",
                             "firepower", "sophos", "pfsense", "iptables")),
    ("VPN",                 ("vpn", "ipsec", "globalprotect", "anyconnect", "pulse secure",
                             "fortivpn", "ssl vpn", "ikev2", "ikev1", "ike ", "strongswan",
                             "openvpn", "wireguard", "l2tp", "pptp", "isakmp")),
    # Cloud BEFORE PKI/HSM: AWS/Azure/GCP configs contain x.509/certificate
    # so PKI/HSM would fire wrongly if checked before Cloud.
    ("Cloud",               ("aws", "amazon", "aws kms", "aws iam",
                             "azure", "azure key vault", "azure blob", "azure ad",
                             "gcp", "google cloud", "gcp kms", "cloud hsm",
                             "cloudfront", "s3 bucket", "ec2", "eks", "aks", "gke",
                             "lambda", "key management service")),
    ("PKI / HSM",           ("pki", "certificate authority", "ca certificate", "root ca",
                             "intermediate ca", "hsm", "safenet", "thales", "luna",
                             "venafi", "entrust", "digicert", "microsoft ca", "adcs",
                             "x.509", "x509", "crl", "ocsp", "est protocol")),
    ("Database",            ("oracle", "oracle db", "mysql", "postgresql", "postgres",
                             "ms sql", "mssql", "sql server", "mongodb", "tde",
                             "transparent data encryption", "database encryption",
                             "sqlnet", "wallet_root", "tde_configuration",
                             "db2", "mariadb", "sybase")),
    ("Load Balancer",       ("nginx", "f5", "big-ip", "citrix adc", "netscaler",
                             "haproxy", "load balancer", "load-balancer", "reverse proxy",
                             "api gateway")),
    ("Web / App",           ("web application", "web app", "webapp", "apache", "tomcat",
                             "iis", "jetty", "django", "flask", "spring",
                             "rest api", "graphql", "oauth", "jwt",
                             "ssl_ciphers", "sslengine")),
    ("SSH / Remote Access", ("sshd", "ssh", "openssh", "putty", "rdp", "remote desktop",
                             "telnet", "jump server", "bastion")),
    # 'host'/'vm'/'server' removed: match inside server_name/ssh_host_rsa_key/vmware_tools
    ("Server",              ("linux", "ubuntu", "centos", "rhel", "debian",
                             "windows server", "linux server",
                             "virtual machine", "esxi", "vmware",
                             "hypervisor", "bare metal")),
]


def _classify_asset_category(asset_ctx: str, filename: str, content_window: str) -> str:
    """
    Deterministic asset category classifier.
    Maps the finding's asset_name, filename and a nearby content window
    to one of the Whiteboard/RFP-defined asset categories.
    100% offline keyword matching — never infers or guesses.
    Returns "Unknown" when no category signal is found.
    """
    combined = f"{asset_ctx or ''} {filename or ''} {content_window or ''}".lower()
    for category, keywords in _ASSET_CATEGORY_RULES:
        if any(kw in combined for kw in keywords):
            return category
    return "Unknown"


# ══════════════════════════════════════════════════════════════════════════════
# OEM / VENDOR PQC READINESS MATRIX  (best-effort, Enhancement C)
# ══════════════════════════════════════════════════════════════════════════════
# Static, manually-curated reference table (src/core/knowledge/pqc_oem_readiness.json)
# mapping vendor/product names to their known PQC readiness status. Loaded once
# at module level -- same path-resolution convention as
# controls_data.py::_merge_knowledge_base() (os.path.join(dirname, "knowledge",
# filename)), just one directory up since this module lives in
# src/core/parsers/ rather than src/core/. Missing/invalid file is non-fatal:
# OEM matching just yields no results, same "fail open, never crash" contract
# as _merge_knowledge_base().

_OEM_READINESS_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)), "knowledge", "pqc_oem_readiness.json"
)


def _load_oem_readiness_table() -> dict:
    try:
        with open(_OEM_READINESS_PATH, "r", encoding="utf-8") as fh:
            return json.load(fh)
    except (FileNotFoundError, json.JSONDecodeError, OSError):
        return {}


_OEM_READINESS_TABLE = _load_oem_readiness_table()


def _build_oem_matchers() -> List[Tuple["re.Pattern", str, str, str]]:
    """Builds (compiled_regex, product_name, vendor, status) tuples for every
    product-name key and vendor name in the readiness table. Each pattern
    matches its literal name on real word boundaries (custom lookaround, not
    bare substring containment) -- e.g. "F5" must not match inside "UTF5".
    Sorted longest-name-first so a more specific name (e.g. "Palo Alto
    PAN-OS") is tried before a shorter, less specific one (e.g. a bare vendor
    name like "Oracle") that might also appear in the same text."""
    matchers = []
    for product_name, info in _OEM_READINESS_TABLE.items():
        if not isinstance(info, dict):
            continue
        vendor = info.get("vendor", "")
        status = info.get("status", "")
        names = {product_name}
        if vendor:
            names.add(vendor)
        for name in names:
            if not name:
                continue
            pattern = re.compile(
                r'(?<![A-Za-z0-9])' + re.escape(name) + r'(?![A-Za-z0-9])', re.IGNORECASE
            )
            # Sort key is the literal matched name's own length (not the
            # product_name key's length) -- a short vendor name like "Oracle"
            # must not be tried before a longer, more specific product name
            # just because it happens to belong to a long product_name key.
            matchers.append((len(name), pattern, product_name, vendor, status))
    matchers.sort(key=lambda t: -t[0])
    return [(pattern, product_name, vendor, status) for _len, pattern, product_name, vendor, status in matchers]


_OEM_MATCHERS = _build_oem_matchers()


def _match_oem_readiness(content: str, filename: str) -> Tuple[str, str, str]:
    """Best-effort OEM/vendor PQC-readiness lookup against the static
    pqc_oem_readiness.json reference table. Checks both evidence content and
    filename. Returns (product_name, vendor, status), or ("", "", "") if
    nothing matched -- never fabricated."""
    if not _OEM_MATCHERS:
        return "", "", ""
    haystack = f"{content or ''} {filename or ''}"
    for pattern, product_name, vendor, status in _OEM_MATCHERS:
        if pattern.search(haystack):
            return product_name, vendor, status
    return "", "", ""


# ══════════════════════════════════════════════════════════════════════════════
# BLOCK-HEADING CONTEXT  (best-effort, supports Enhancements B & D)
# ══════════════════════════════════════════════════════════════════════════════

_DASH_LINE_RE = re.compile(r'^[-_=]{3,}$')


def _find_block_heading_context(content: str, match_start: int) -> str:
    """Best-effort broader context for a match, used ONLY to enrich
    finding.asset_name as a fallback when the stricter Phase-1
    _find_asset_context() found no 'Host:'/'Target:'-style heading (i.e. fell
    back to the bare filename) -- never overwrites a real Phase-1 match, and
    never touches _find_asset_context() itself.

    Narrative evidence (an architecture diagram's extracted text, a free-form
    device write-up) often labels a block with plain heading lines instead of
    'Label: value' pairs, e.g.:
        Internet Banking App - External Production System
        Palo Alto PAN-OS Firewall Configuration
        IPSec VPN Profile
        -----------------------------
        Certificate : RSA2048
    This walks back from the match to the start of the current blank-line
    delimited paragraph and collects the heading lines at its top -- stopping
    at the first dashed separator or 'label: value' line. Returns "" if
    nothing usable is found.
    """
    para_start = content.rfind("\n\n", 0, match_start)
    para_start = 0 if para_start == -1 else para_start + 2
    para_text = content[para_start:match_start]

    heading_lines = []
    for line in para_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith(("#", "//", ";", "--", "/*", "*/", "http", "server", "{", "}")):
            break
        if _DASH_LINE_RE.match(stripped):
            break
        if ":" in stripped or "=" in stripped:
            break
        heading_lines.append(stripped)
    return " ".join(heading_lines)


# ══════════════════════════════════════════════════════════════════════════════
# DEPENDENCY CHAIN PARSING  (evidence-stated only, Enhancement D)
# ══════════════════════════════════════════════════════════════════════════════

_CHAIN_LINE_RE = re.compile(r'^.+(?:->|→|\|).+$')
_CHAIN_SPLIT_RE = re.compile(r'->|→|\|')
_VERTICAL_CONNECTOR_TOKENS = ("|", "│", "▼")


def _parse_dependency_chains(content: str) -> List[List[str]]:
    """Best-effort extraction of EXPLICITLY-stated dependency/architecture
    chains from evidence text. Only surfaces a chain when the evidence text
    itself states one -- never infers/guesses topology from nothing.

    Supports two formats:
      1. Single-line arrow/pipe chain: 'A -> B -> C', 'A | B | C', 'A -> B -> C'
         (arrow variants), comma-free simple separator-delimited chains.
      2. A simple vertical chain mirroring an architecture diagram's own
         layout -- one node name per line, delimited by connector-only lines
         ('|', a vertical bar glyph, or a downward-triangle glyph). Requires
         at least one such connector line to be present, so ordinary
         multi-line paragraph text is never mistaken for a chain.

    Returns [] if nothing matches.
    """
    chains: List[List[str]] = []
    lines = content.splitlines()

    # ── Format 1: single-line arrow/pipe chains ──
    for line in lines:
        stripped = line.strip()
        if not stripped or not _CHAIN_LINE_RE.match(stripped):
            continue
        segments = [seg.strip() for seg in _CHAIN_SPLIT_RE.split(stripped)]
        segments = [seg for seg in segments if seg]
        if len(segments) >= 2:
            chains.append(segments)

    # ── Format 2: simple vertical chain (node / connector / node / ...) ──
    current: List[str] = []
    saw_connector = False
    for line in lines + [""]:  # sentinel blank line flushes the last run
        stripped = line.strip()
        if not stripped:
            if saw_connector and len(current) >= 2:
                chains.append(current)
            current, saw_connector = [], False
            continue
        if stripped in _VERTICAL_CONNECTOR_TOKENS:
            saw_connector = True
            continue
        if _CHAIN_LINE_RE.match(stripped):
            # A single-line arrow chain, not a bare vertical node name --
            # already captured by Format 1 above.
            if saw_connector and len(current) >= 2:
                chains.append(current)
            current, saw_connector = [], False
            continue
        current.append(stripped)

    return chains


# ══════════════════════════════════════════════════════════════════════════════
# PARSER
# ══════════════════════════════════════════════════════════════════════════════

class PQCParser(BaseParser):
    """Deterministic Post-Quantum Cryptography readiness scanner. Parses plain
    text / HTML / XML / PDF-extracted evidence (TLS, SSH, IPSec, PKI, DB
    encryption, HSM/KMS, code-signing config exports) and flags cryptographic
    algorithms as Quantum-Vulnerable, Classically-Weak, or Quantum-Safe.
    """

    def can_parse(self, filename: str, content: str) -> bool:
        # ── Binary document formats (PDF, DOCX, images) ─────────────────────
        # These require text extraction before keyword matching. We accept them
        # by extension alone -- text will be extracted inside parse().
        ext_lower = os.path.splitext(filename.lower())[1]
        if ext_lower in _PQC_BINARY_EXTENSIONS:
            # Binary formats need text extraction before keywords can be checked,
            # so an unextracted binary is accepted for parse() to OCR/extract.
            # But claiming EVERY pdf/docx/image on extension alone contradicts the
            # rule every other parser follows -- dispatch is by content signature,
            # never by filename -- and would let any ISO policy PDF be answered by
            # the PQC parser. When the caller has already extracted text, gate on
            # it exactly as the plain-text path below does.
            if content and content.strip():
                return _count_pqc_signals(content.lower()) >= 2
            return True

        # ── Plain-text / config formats ──────────────────────────────────────
        if not content:
            return False
        # Guard: reject image files that somehow slipped through (content garbage)
        if is_image_file(filename):
            return False

        sample = content.lower()
        weak_signal_count = _count_pqc_signals(sample)
        if weak_signal_count >= 2:
            return True

        # Recognizable config-export extension -- still content-gated (a lower
        # bar of 1 keyword hit, not filename-only), since the extension alone
        # is already meaningful signal for this file type.
        if filename.lower().endswith(_PQC_CONFIG_EXTENSIONS) and weak_signal_count >= 1:
            return True

        return False

    def parse(self, filename: str, content: str) -> Tuple[List[Finding], List[Finding]]:
        # ── Binary document: extract text first ──────────────────────────────
        ext_lower = os.path.splitext(filename.lower())[1]
        if ext_lower in _PQC_BINARY_EXTENSIONS:
            # content may be empty (binary was passed as bytes from bg_worker
            # then decoded as garbage, or it's genuinely empty)
            raw_bytes: Optional[bytes] = None
            if isinstance(content, (bytes, bytearray)):
                raw_bytes = bytes(content)
            elif not content or len(content.strip()) < 20:
                # Nothing useful in content string -- signal caller to provide bytes
                print(
                    f"[PQC PARSER] '{filename}' is a binary document but no usable "
                    f"text was supplied. Call pqc_extract_text(filename, raw_bytes) "
                    f"before passing to parse().",
                    flush=True,
                )
                return [], []
            elif content and content.strip().startswith("-----BEGIN"):
                # PEM content passed as text -- re-extract to get structured algorithm text
                raw_bytes = content.encode("utf-8", errors="ignore")
            else:
                # content is already extracted text (bg_worker did it upstream)
                pass
            if raw_bytes is not None:
                content = pqc_extract_text(filename, raw_bytes)
            if not content or not content.strip():
                print(
                    f"[PQC PARSER] Could not extract text from '{filename}' "
                    f"(empty after extraction). No PQC findings generated.",
                    flush=True,
                )
                return [], []
            print(
                f"[BINARY TEXT EXTRACTOR] Extracted {len(content)} chars from binary '{filename}' "
                f"(ext: {ext_lower}).",
                flush=True,
            )

        if not content:
            return [], []

        actionable_findings: List[Finding] = []
        info_findings: List[Finding] = []
        accepted_spans: List[Tuple[int, int]] = []

        # Per-line dedup state. A cipher list names the same primitive several
        # times on one line (ECDHE-RSA-...:ECDHE-RSA-...:ECDHE-RSA-...), and
        # emitting one finding per occurrence produced 3x duplicates carrying an
        # identical evidence quote -- which inflated the P1/P2 counts the
        # dashboard and executive summary are built from.
        _dedup_seen: Dict[Tuple[str, str, str], Finding] = {}
        _dedup_counts: Dict[Tuple[str, str, str], int] = {}

        def _overlaps(start: int, end: int) -> bool:
            for s, e in accepted_spans:
                if start < e and end > s:
                    return True
            return False

        for rule_id, pattern, namer, quantum_status, crypto_category, severity_rule in ALGORITHM_RULES:
            for m in pattern.finditer(content):
                start, end = m.start(), m.end()
                if _overlaps(start, end):
                    continue
                accepted_spans.append((start, end))

                algo_name = _resolve(namer, m)
                severity = _resolve(severity_rule, m)
                evidence_line = _line_containing(content, start, end)

                # A commented-out directive is not a deployed algorithm. Scanning the
                # whole file meant a config that had already migrated -- classical
                # ciphers left in place as commented reference, only X25519MLKEM768
                # active -- still produced CRITICAL/HIGH findings for RSA, ECC P-384
                # and ECDHE. Reporting quantum-vulnerable crypto from lines that
                # explicitly say it is no longer in use is the same class of false
                # positive as reading remediation prose as a live vulnerability.
                #
                # Only a line that BEGINS with a comment marker is skipped, so a real
                # directive carrying a trailing comment ("ssl_ciphers ...;  # legacy")
                # is still assessed.
                if _is_comment_line(evidence_line):
                    continue

                # Collapse repeat hits of the same algorithm on the same line to
                # a single finding, recording how many times it appeared. Keyed
                # on the evidence line (not the whole document) so the same
                # algorithm configured on two different directives stays two
                # findings -- those are genuinely separate facts.
                dedup_key = (rule_id, algo_name, evidence_line)
                if dedup_key in _dedup_seen:
                    _dedup_counts[dedup_key] += 1
                    continue
                _dedup_counts[dedup_key] = 1

                asset_ctx = _find_asset_context(content, start, filename)

                # Enhancement 2: exposure-based severity escalation (EXTERNAL only).
                exposure_context = _classify_exposure(content, start, asset_ctx, filename)

                # Second-pass inference: if keyword detection couldn't determine exposure
                # (the config file doesn't literally say "external" / "internal"), infer
                # it from structural signals — same approach as Nessus/Qualys which assign
                # AV:N (Network/External) based on service type, not document vocabulary.
                if not exposure_context:
                    _cat = getattr(finding, "asset_category", "") if False else ""
                    # Use the already-set category if available, else classify now.
                    _cat_window_exp = content[max(0, start - 400): min(len(content), end + 400)]
                    _inferred_cat = _classify_asset_category(asset_ctx, filename, _cat_window_exp)

                    # Definitionally external-facing asset categories:
                    _EXTERNAL_CATEGORIES = {"Load Balancer", "Firewall", "VPN", "Web / App"}
                    # Definitionally internal/server-side categories:
                    _INTERNAL_CATEGORIES = {"Database", "Server", "SSH / Remote Access"}

                    # Structural TLS signals — a config that terminates TLS on 443 is
                    # internet-facing by function even if it never says "external".
                    _content_lower = content.lower()
                    _tls_external_signals = (
                        "listen 443" in _content_lower or
                        "ssl_certificate" in _content_lower or
                        "server_name" in _content_lower or
                        "ssl on" in _content_lower
                    )

                    # Definitionally external-facing categories:
                    # Cloud KMS/HSM = public cloud APIs (HNDL threat applies)
                    # PKI/HSM = public CA endpoints (OCSP/CRL are internet-facing)
                    _EXTERNAL_CATEGORIES = {
                        "Load Balancer", "Firewall", "VPN", "Web / App",
                        "Cloud", "PKI / HSM"
                    }
                    # Definitionally internal/server-side categories:
                    _INTERNAL_CATEGORIES = {"Database", "Server", "SSH / Remote Access"}

                    if _inferred_cat in _EXTERNAL_CATEGORIES or _tls_external_signals:
                        exposure_context = "EXTERNAL"
                    elif _inferred_cat in _INTERNAL_CATEGORIES:
                        exposure_context = "INTERNAL"

                # Track whether escalation actually fired, so post-pass 2 explains
                # only the findings whose severity really did change. A finding
                # that was already CRITICAL by rule, or an INFO-level SAFE one,
                # is EXTERNAL too but was not escalated.
                _was_escalated = False
                if exposure_context == "EXTERNAL":
                    if severity == "MEDIUM":
                        severity = "HIGH"
                        _was_escalated = True
                    elif severity == "HIGH":
                        severity = "CRITICAL"
                        _was_escalated = True

                if rule_id in _CONFIG_PRESENCE_RULES:
                    # Config-posture rule: the evidence is a directive or a file
                    # path, not an algorithm, so do not claim an algorithm was
                    # detected. quantum_status stays VULNERABLE so the HNDL/QV
                    # risk scoring in control_mapper.py is unaffected.
                    title = f"PQC Readiness Gap: {algo_name}"
                    description = (
                        f"Evidence in '{filename}' shows {algo_name} ({crypto_category}) is "
                        f"configured, but the configuration specifies no post-quantum key "
                        f"exchange or signature algorithm. Traffic protected by this service "
                        f"therefore relies entirely on classical cryptography and is exposed "
                        f"to 'harvest now, decrypt later' capture. This finding reports a "
                        f"missing post-quantum capability, not a specific weak algorithm."
                    )
                    remediation = _get_remediation_vulnerable(algo_name, crypto_category)
                elif quantum_status == "VULNERABLE":
                    title = f"Quantum-Vulnerable Algorithm Detected: {algo_name}"
                    description = (
                        f"Evidence in '{filename}' shows use of {algo_name} ({crypto_category}). "
                        f"This algorithm is quantum-vulnerable: a sufficiently large quantum computer "
                        f"running Shor's algorithm can efficiently break it, undermining the "
                        f"confidentiality/integrity of anything protected by it once such hardware "
                        f"exists (including data captured today and decrypted later -- "
                        f"'harvest now, decrypt later')."
                    )
                    remediation = _get_remediation_vulnerable(algo_name, crypto_category)
                elif quantum_status == "WEAK":
                    title = f"Classically Weak / Deprecated Algorithm Detected: {algo_name}"
                    description = (
                        f"Evidence in '{filename}' shows use of {algo_name} ({crypto_category}). "
                        f"This is a classically weak or deprecated algorithm/protocol version -- not "
                        f"specifically a quantum-computing concern, but already considered broken or "
                        f"unsafe against today's classical attacks and should be retired regardless of "
                        f"the organization's PQC migration timeline."
                    )
                    remediation = _REMEDIATION_WEAK
                else:  # SAFE
                    title = f"Quantum-Safe Algorithm Confirmed: {algo_name}"
                    description = (
                        f"Evidence in '{filename}' shows use of {algo_name} ({crypto_category}). "
                        f"This algorithm is considered quantum-resistant against currently known "
                        f"quantum attacks (Grover's algorithm only provides a quadratic speed-up "
                        f"against symmetric/hash primitives of this size, and NIST-selected PQC "
                        f"algorithms are designed to resist Shor's algorithm entirely)."
                    )
                    remediation = _REMEDIATION_SAFE

                # ── Name the ONE standard this finding bears on ──────────────
                # Only for findings that call for migration. A SAFE finding has
                # nothing to migrate to, and a classically-WEAK symmetric cipher
                # or hash is not replaced by FIPS 203-206 at all -- so neither
                # gets a PQC standard appended, which is the whole point of
                # being specific rather than listing all four.
                _cves = _cves_for(rule_id, m.group(0))
                if quantum_status == "VULNERABLE":
                    _fips = _fips_impacted(rule_id, crypto_category)
                    if _fips:
                        description += (
                            f" The applicable post-quantum replacement standard is {_fips}."
                        )
                if _cves:
                    description += (
                        f" This algorithm additionally carries a published classical weakness: "
                        f"{', '.join(_cves)}."
                    )

                finding = Finding(
                    title=title,
                    severity=severity,
                    target=asset_ctx,
                    description=description,
                    remediation=remediation,
                    evidence=evidence_line,
                    plugin_id=f"PQC-{rule_id}",
                    source_tool="PQC-Scan",
                    cve_list=_cves,
                )
                finding.nist_80053_controls = _nist_80053_for(crypto_category)
                finding.asset_name = asset_ctx
                # Auto-classify the asset category from filename + context window.
                _cat_window = content[max(0, start - 400): min(len(content), end + 400)]
                finding.asset_category = _classify_asset_category(asset_ctx, filename, _cat_window)
                finding.quantum_status = quantum_status
                # Which rule produced this, for the post-loop suppression pass.
                # Underscore-prefixed so it stays out of to_dict()/serialisation.
                finding._pqc_rule_id = rule_id
                finding._pqc_severity_escalated = _was_escalated
                _dedup_seen[dedup_key] = finding

                # Enhancement B/D infrastructure: best-effort asset_name
                # enrichment, ONLY when Phase-1's stricter _find_asset_context()
                # found no real "Host:"/"Target:"-style heading and fell back
                # to the bare filename -- never overrides a genuine Phase-1
                # match. Narrative/architecture-diagram evidence (this is the
                # exact style Enhancements B and D are meant to read) labels
                # blocks with plain heading lines instead, which the broader
                # _find_block_heading_context() heuristic picks up.
                if asset_ctx == filename:
                    block_heading = _find_block_heading_context(content, start)
                    if block_heading:
                        finding.asset_name = block_heading

                # Enhancement 1: CA / Key / Protocol layer classification.
                _assign_crypto_layers(finding, algo_name, crypto_category, content=content)
                crypto_layer = _classify_crypto_layer(content, start, end)
                if crypto_layer == "CA" and not finding.ca_algorithm:
                    finding.ca_algorithm = algo_name
                elif crypto_layer == "KEY" and not finding.key_algorithm:
                    finding.key_algorithm = algo_name
                elif crypto_layer == "PROTOCOL" and not finding.protocol_version:
                    finding.protocol_version = algo_name

                # Enhancement 2: exposure context (severity already escalated above).
                finding.exposure_context = exposure_context

                # Enhancement 3: nearby port reference (best-effort, never fabricated).
                finding.port = _find_port(content, start, end)

                # Enhancement 4: prod/non-prod environment tag (informational only).
                finding.environment = _classify_environment(content, start, asset_ctx, filename)

                # Enhancement C: OEM/vendor PQC readiness matrix lookup (best-effort).
                oem_product, _oem_vendor, oem_status = _match_oem_readiness(content, filename)
                finding.oem_product = oem_product
                finding.oem_readiness_status = oem_status

                if finding.severity in ("CRITICAL", "HIGH", "MEDIUM", "LOW"):
                    actionable_findings.append(finding)
                else:
                    info_findings.append(finding)

        # ── POST-PASS 1: record repeat occurrences collapsed by the dedup above ──
        for _key, _count in _dedup_counts.items():
            if _count > 1:
                _f = _dedup_seen.get(_key)
                if _f is not None:
                    _f.description = (
                        f"{_f.description} This algorithm appears {_count} times on the "
                        f"cited line (for example, repeated across a cipher suite list); "
                        f"they are reported once as a single configuration fact."
                    )

        # ── POST-PASS 2: explain an exposure-driven severity escalation ─────────
        # Severity is escalated one step for internet-facing assets, which is why
        # the same algorithm can be CRITICAL in an edge config and HIGH in an
        # internal one. Without saying so, that reads to an auditor as the tool
        # contradicting itself, so state the reason on the finding.
        for _f in actionable_findings:
            if getattr(_f, "_pqc_severity_escalated", False):
                _f.description = (
                    f"{_f.description} Severity is escalated one level because this asset "
                    f"is internet-facing, which makes 'harvest now, decrypt later' capture "
                    f"of the traffic practical for a remote adversary."
                )

        # ── EXTENDED SCAN: OID / IANA Cipher Suite / liboqs algorithms ────────
        # Runs the offline pqc_crypto_db lookup on the same text to catch
        # algorithm references expressed as:
        #   - X.509 dotted-decimal OID strings (e.g. 1.2.840.10045.4.3.2 = ECDSA)
        #   - IANA TLS cipher suite hex codes (e.g. 0xC02B = ECDHE-ECDSA-AES128-GCM-SHA256)
        #   - liboqs/OQS algorithm keywords (FrodoKEM, HQC, BIKE, Kyber768, etc.)
        # Each hit that doesn't overlap an already-accepted regex span becomes
        # a new Finding, enriched with CWE ID, NIST OID reference, and IANA
        # cipher suite key-exchange details.

        def _make_db_finding(name: str, meta: dict, evidence_hint: str, source_tag: str) -> Finding:
            qs = meta.get("quantum_status", "VULNERABLE")
            # A quantum-SAFE hit must never inherit the VULNERABLE default of
            # HIGH. liboqs_algorithms.json carries no 'severity' key at all, so
            # all 46 of its post-quantum entries were landing as HIGH-severity
            # *actionable* findings -- i.e. a correctly deployed FrodoKEM or
            # Kyber768 was reported as a high-severity problem. SAFE hits are
            # informational, exactly as the regex table's own PQC rules are.
            sev = meta.get("severity") or (_SEV_INFO if qs == "SAFE" else "HIGH")
            category = meta.get("category", "Cryptographic Algorithm")
            nist_ref = meta.get("nist_ref", "")
            cwe = meta.get("cwe", "")
            kex = meta.get("kex", "")
            if qs == "VULNERABLE":
                f_title = f"Quantum-Vulnerable Algorithm Detected: {name}"
                f_desc = (
                    f"Evidence in '{filename}' shows use of {name} ({category}). "
                    f"This algorithm is quantum-vulnerable and broken by Shor's algorithm "
                    f"on a sufficiently large quantum computer. "
                    + (f"IANA Key Exchange: {kex}. " if kex and kex not in ("N/A", "") else "")
                    + (f"NIST Reference: {nist_ref}. " if nist_ref else "")
                    + (f"CWE: {cwe}." if cwe else "")
                )
                f_remed = _get_remediation_vulnerable(name, category)
            elif qs == "WEAK":
                f_title = f"Classically Weak / Deprecated Algorithm Detected: {name}"
                f_desc = (
                    f"Evidence in '{filename}' shows use of {name} ({category}). "
                    f"This is a classically weak or deprecated algorithm -- already breakable "
                    f"with classical computing and should be retired regardless of PQC timeline. "
                    + (f"NIST Reference: {nist_ref}. " if nist_ref else "")
                    + (f"CWE: {cwe}." if cwe else "")
                )
                f_remed = _REMEDIATION_WEAK
            else:  # SAFE
                f_title = f"Quantum-Safe Algorithm Confirmed: {name}"
                f_desc = (
                    f"Evidence in '{filename}' shows use of {name} ({category}). "
                    f"This algorithm is quantum-resistant (NIST-selected PQC). "
                    + (f"NIST Reference: {nist_ref}." if nist_ref else "")
                )
                f_remed = _REMEDIATION_SAFE

            db_finding = Finding(
                title=f_title,
                severity=sev,
                target=filename,
                description=f_desc,
                remediation=f_remed,
                evidence=evidence_hint,
                plugin_id=f"PQC-db-{source_tag}-{name[:20].replace(' ', '_')}",
                source_tool="PQC-Scan",
            )
            db_finding.quantum_status = qs
            # Tag PQC-safe hits from this scan with the same marker the regex
            # rules use, so the readiness-gap suppression below counts a
            # post-quantum algorithm expressed only as an X.509 OID, an IANA
            # suite code, or a liboqs keyword as real evidence of PQC.
            # liboqs (source_tag 'oqs') is a post-quantum algorithm database in
            # its entirety, but its entries carry no 'category', so match on the
            # source there and on the category for the OID database (whose SAFE
            # entries are all labelled 'PQC ...'). A SAFE hit that is merely
            # quantum-resistant symmetric crypto (AES-256) must NOT qualify --
            # it does not make a "no post-quantum suite configured" gap false.
            db_finding._pqc_rule_id = (
                "pqc-extended-scan"
                if qs == "SAFE" and (
                    source_tag == "oqs" or "PQC" in (category or "").upper()
                )
                else ""
            )
            db_finding.asset_name = filename
            db_finding.asset_category = _classify_asset_category("", filename, content[:500])
            _assign_crypto_layers(db_finding, name, category, meta=meta, content=content)
            oem_product, _, oem_status = _match_oem_readiness(content, filename)
            db_finding.oem_product = oem_product
            db_finding.oem_readiness_status = oem_status
            return db_finding

        # -- OID scan ---------------------------------------------------------
        for oid_str, meta in scan_oids_in_text(content):
            name = meta["name"]
            # Skip if already covered by a regex rule match (avoids double-reporting)
            oid_already_covered = any(
                name.split()[0].upper() in (f.title or "").upper()
                for f in actionable_findings + info_findings
            )
            if oid_already_covered:
                continue
            hint = f"OID detected in text: {oid_str} = {name}"
            db_f = _make_db_finding(name, meta, hint, f"oid")
            if db_f.severity in ("CRITICAL", "HIGH", "MEDIUM", "LOW"):
                actionable_findings.append(db_f)
            else:
                info_findings.append(db_f)

        # -- IANA cipher suite hex scan ---------------------------------------
        for hex_code, meta in scan_iana_hex_in_text(content):
            suite_name = meta["name"]
            kex = meta.get("kex", "")
            # Only flag if not already reported
            suite_already_covered = any(
                meta.get("kex", "").split("-")[0].upper() in (f.title or "").upper()
                for f in actionable_findings + info_findings
            )
            if suite_already_covered:
                continue
            hint = f"IANA cipher suite hex 0x{hex_code} = {suite_name} (KEX: {kex})"
            db_f = _make_db_finding(suite_name, meta, hint, f"iana")
            if db_f.severity in ("CRITICAL", "HIGH", "MEDIUM", "LOW"):
                actionable_findings.append(db_f)
            else:
                info_findings.append(db_f)

        # -- liboqs / OQS algorithm keyword scan ------------------------------
        for kw, meta in scan_liboqs_in_text(content):
            algo_name = meta["name"]
            qs = meta.get("quantum_status", "SAFE")
            # Skip if already covered by ALGORITHM_RULES regex (e.g. kyber / dilithium)
            already_covered = any(
                kw.lower() in (f.title or "").lower()
                for f in actionable_findings + info_findings
            )
            if already_covered:
                continue
            hint = f"liboqs/OQS algorithm keyword detected: {kw} = {algo_name}"
            db_f = _make_db_finding(algo_name, meta, hint, f"oqs")
            if db_f.severity in ("CRITICAL", "HIGH", "MEDIUM", "LOW"):
                actionable_findings.append(db_f)
            else:
                info_findings.append(db_f)

        if not actionable_findings and not info_findings:
            return [], []

        # Enhancement D: dependency-chain mapping (evidence-stated only, never
        # inferred). Only VULNERABLE findings participate -- SAFE/WEAK assets
        # aren't part of the "migration dependency" concept this models.
        chains = _parse_dependency_chains(content)
        if chains:
            vulnerable_findings = [f for f in actionable_findings if f.quantum_status == "VULNERABLE"]
            for f in vulnerable_findings:
                asset_lower = (f.asset_name or "").lower()
                if not asset_lower:
                    continue

                matched_chain = None
                matched_index = -1
                for chain in chains:
                    for idx, node in enumerate(chain):
                        node_lower = node.lower()
                        if node_lower in asset_lower or asset_lower in node_lower:
                            matched_chain, matched_index = chain, idx
                            break
                    if matched_chain:
                        break
                if matched_chain is None:
                    continue

                f.dependency_chain = " -> ".join(matched_chain)

                # A "migration dependency" exists when this asset shares an
                # explicit chain with at least one OTHER asset that also has
                # its own VULNERABLE finding elsewhere in this file --
                # direction-agnostic (checked against every other node in the
                # chain, not just strictly-downstream ones), since fixing one
                # crypto asset in a stated dependency chain has migration-
                # sequencing implications for every other vulnerable asset in
                # that same chain regardless of traffic direction.
                for other_idx, other_node in enumerate(matched_chain):
                    if other_idx == matched_index:
                        continue
                    other_node_lower = other_node.lower()
                    for other_f in vulnerable_findings:
                        if other_f is f:
                            continue
                        other_asset_lower = (other_f.asset_name or "").lower()
                        if other_asset_lower and (
                            other_node_lower in other_asset_lower or other_asset_lower in other_node_lower
                        ):
                            f.migration_dependency_flag = True
                            break
                    if f.migration_dependency_flag:
                        break

        # ── POST-PASS 3: drop PQC-readiness gaps that a PQC algorithm disproves ──
        # A config that already negotiates ML-KEM / ML-DSA is not "TLS enabled
        # with no post-quantum suite"; emitting both would have the report assert
        # and deny the same fact on one page. Runs here, after the extended
        # OID/IANA/liboqs scan, so a post-quantum algorithm expressed only as an
        # X.509 OID or an IANA suite code still counts as evidence of PQC.
        _has_pqc_algorithm = any(
            getattr(_f, "_pqc_rule_id", "") in _PQC_ALGORITHM_RULES
            for _f in actionable_findings + info_findings
        )
        if _has_pqc_algorithm:
            _before = len(actionable_findings) + len(info_findings)
            actionable_findings = [
                _f for _f in actionable_findings
                if getattr(_f, "_pqc_rule_id", "") not in _CONFIG_PRESENCE_RULES
            ]
            info_findings = [
                _f for _f in info_findings
                if getattr(_f, "_pqc_rule_id", "") not in _CONFIG_PRESENCE_RULES
            ]
            _dropped = _before - (len(actionable_findings) + len(info_findings))
            if _dropped:
                print(
                    f"[PQC PARSER] '{filename}': suppressed {_dropped} PQC-readiness gap "
                    f"finding(s) -- a NIST post-quantum algorithm is present in this file.",
                    flush=True,
                )

        map_pqc_findings_list(actionable_findings)
        map_pqc_findings_list(info_findings)

        print(
            f"[PQC PARSER] Extracted {len(actionable_findings)} actionable + "
            f"{len(info_findings)} informational finding(s) from '{filename}'.",
            flush=True
        )
        return actionable_findings, info_findings
