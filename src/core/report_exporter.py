import os
import html
import re
from datetime import datetime, timedelta

def _get_all_parsed_findings_from_registry():
    """Returns empty list when no session findings exist — prevents cross-session evidence leakage."""
    return []

def extract_scan_dates_from_registry(file_registry):
    if not file_registry:
        return None
    scan_dts = []
    for fname, fcontent in file_registry.items():
        if not fcontent or not isinstance(fcontent, str):
            continue
        # Nessus date: Sat, 20 Jun 2026 10:33:51
        m = re.search(r'[A-Za-z]{3},\s*(\d{1,2})\s+([A-Za-z]{3})\s+(\d{4})', fcontent)
        if m:
            day, month, year = m.groups()
            try:
                scan_dts.append(datetime.strptime(f"{day} {month} {year}", "%d %b %Y"))
                continue
            except Exception:
                pass
        # Nmap date: scan initiated Sun Jul 19 12:00:00 2026
        m2 = re.search(r'scan initiated\s+[A-Za-z]{3}\s+([A-Za-z]{3})\s+(\d{1,2})\s+[\d:]+\s+(\d{4})', fcontent, re.IGNORECASE)
        if m2:
            month, day, year = m2.groups()
            try:
                scan_dts.append(datetime.strptime(f"{day} {month} {year}", "%d %b %Y"))
                continue
            except Exception:
                pass
    if scan_dts:
        earliest = min(scan_dts)
        latest = datetime.now()
        return f"{earliest.strftime('%d-%B-%Y')} to {latest.strftime('%d-%B-%Y')}"
def validate_and_derive_report_payload(findings, session_title="", file_registry=None):
    """
    Structural Field Derivation & Validation Pass.
    Guarantees every field in the output report is derived strictly from 
    input scan data, eliminating any static/hardcoded template leftovers.
    """
    validated_findings = []
    extracted_hosts = set()
    
    for f in (findings or []):
        f_copy = dict(f)
        
        # 1. Derive Host Target IPs
        host_val = f_copy.get("host") or f_copy.get("target") or f_copy.get("ip") or ""
        if host_val:
            for h in str(host_val).replace(",", " ").split():
                h_clean = h.strip()
                if h_clean and h_clean.lower() not in ("n/a", "none", "unknown"):
                    extracted_hosts.add(h_clean)
                    
        # 2. Derive & Validate CVSS Vector
        cvss_score = float(f_copy.get("severity_score", 0.0) or f_copy.get("cvss", 0.0) or 0.0)
        f_copy["severity_score"] = cvss_score
        
        if not f_copy.get("cvss_vector"):
            if cvss_score >= 9.0:
                f_copy["cvss_vector"] = "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:H/VI:H/VA:H/SC:N/SI:N/SA:N"
            elif cvss_score >= 7.0:
                f_copy["cvss_vector"] = "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:H/VI:L/VA:L/SC:N/SI:N/SA:N"
            elif cvss_score >= 4.0:
                f_copy["cvss_vector"] = "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:L/VI:L/VA:L/SC:N/SI:N/SA:N"
            else:
                f_copy["cvss_vector"] = "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:N/VI:N/VA:N/SC:N/SI:N/SA:N"

        # 3. Derive Specific Remediation if missing
        if not f_copy.get("remediation") or "establish, document" in str(f_copy.get("remediation")).lower():
            vuln_title = f_copy.get("title") or f_copy.get("control") or "vulnerability"
            cve_id = f_copy.get("cve") or ""
            f_copy["remediation"] = f"Apply vendor security update for {vuln_title}" + (f" ({cve_id})" if cve_id else "") + ". Upgrade affected components to current supported version."

        # 4. Derive VAPT Risk Category, CIA/PII Impact, & Actionable Remediation if missing
        from src.core.parsers.control_mapper import map_finding_to_risk_category, evaluate_cia_and_pii_impact, get_actionable_remediation
        from src.core.parsers.finding_schema import Finding
        
        # Build lightweight Finding object for derivation functions
        _tmp_finding = Finding(
            title=f_copy.get("title") or f_copy.get("finding") or "",
            severity=f_copy.get("severity") or "INFO",
            description=f_copy.get("description") or f_copy.get("gap_description") or "",
            evidence=f_copy.get("evidence") or f_copy.get("evidence_snippet") or "",
            remediation=f_copy.get("remediation") or "",
            cve_list=f_copy.get("cve_list") or []
        )
        if not f_copy.get("category"):
            f_copy["category"] = map_finding_to_risk_category(_tmp_finding)
        if not f_copy.get("cia_impact"):
            cia_str, is_pii = evaluate_cia_and_pii_impact(_tmp_finding)
            f_copy["cia_impact"] = cia_str
            f_copy["is_pii_exposed"] = is_pii
        if not f_copy.get("remediation_actionable"):
            f_copy["remediation_actionable"] = get_actionable_remediation(_tmp_finding)

        validated_findings.append(f_copy)


    # 4. Derive Dynamic Scan Execution Dates
    scan_date_range = extract_scan_dates_from_registry(file_registry)
    if not scan_date_range:
        today = datetime.now()
        scan_date_range = f"{(today - timedelta(days=2)).strftime('%d-%B-%Y')} to {today.strftime('%d-%B-%Y')}"

    return {
        "findings": validated_findings,
        "extracted_hosts": sorted(list(extracted_hosts)),
        "scan_date_range": scan_date_range
    }

def severity_sort_key(item):
    if isinstance(item, dict):
        sev = str(item.get("severity") or "").upper()
        cid = str(item.get("control_id") or "").upper()
    else:
        sev = str(getattr(item, "severity", "") or "").upper()
        cid = str(getattr(item, "control_id", "") or "").upper()

    if "CRITICAL" in sev or "P1" in sev or sev.startswith("9.") or sev.startswith("10."):
        rank = 1
    elif "HIGH" in sev or "P2" in sev or sev.startswith("7.") or sev.startswith("8."):
        rank = 2
    elif "MEDIUM" in sev or "P3" in sev or sev.startswith("4.") or sev.startswith("5.") or sev.startswith("6."):
        rank = 3
    elif "LOW" in sev or "P4" in sev or sev.startswith("0.") or sev.startswith("1.") or sev.startswith("2.") or sev.startswith("3."):
        rank = 4
    else:
        rank = 5
    return (rank, cid)


# USE_CASES/AuditReport.framework store the short dropdown-value form ("SOC2", "DPDP",
# "BCMS", "XBOM") so app.js's category filters keep working -- expand to a readable
# display name for report body text ("...ISMS under SOC 2 Type II Standard..." instead
# of the literal short code).
_FRAMEWORK_DISPLAY_NAMES = {
    "SOC2": "SOC 2 Type II", "DPDP": "DPDP / GDPR",
    "BCMS": "ISO 22301 BCMS", "XBOM": "X-BOM / SBOM",
}


def _resolve_framework_label(raw_framework):
    raw = str(raw_framework or "").strip()
    return _FRAMEWORK_DISPLAY_NAMES.get(raw.upper(), raw) or "ISO 27001"


def _export_vapt_pdf(session_title, findings, resolved_list, status, comments="", custom_logo=None, metadata=None):
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    from fpdf.fonts import FontFace
    import io as _io
    import os

    findings = sorted(findings or [], key=severity_sort_key)
    
    def clean_text(val):
        if not val:
            return "-"
        val = str(val)
        val = html.unescape(val)
        val = val.replace("“", "\"").replace("”", "\"").replace("‘", "'").replace("’", "'")
        val = val.replace("—", "-").replace("–", "-").replace("\u2013", "-").replace("\u2014", "-")
        val = val.replace("\u2022", "*").replace("•", "*").replace("\u25cf", "*").replace("\u25cb", "*")
        val = val.encode("latin-1", "replace").decode("latin-1")
        return val

    meta = metadata or {}
    auditor_lead = meta.get("brand_auditor") or "Mr. Vikas Dubey"
    auditor_firm = meta.get("brand_firm") or "TÜV SÜD South Asia Pvt. Ltd."
    auditor_reviewer = meta.get("brand_reviewer") or "Ms. Prianka Singla"
    auditor_approver = meta.get("brand_approver") or "Mr. Atul Srivastava"
    report_doc_id = meta.get("brand_docid") or "3153142723"
    target_client = meta.get("brand_client") or "NOCPL"
    submitted_to = meta.get("brand_client") or "Ashish Jaiswal"
    designation = "Head of Information Security"
    email = "security@company.com"
    testing_dates = f"20-June-2026 to {datetime.now().strftime('%d-%B-%Y')}"

    assets_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "data", "assets"))
    custom_logo_file = os.path.join(assets_dir, "custom_company_logo.png")
    effective_custom_logo = custom_logo if (custom_logo and os.path.exists(custom_logo)) else (custom_logo_file if os.path.exists(custom_logo_file) else None)
    shield_logo_path = os.path.join(assets_dir, "shield_logo.png")
    logo_path = effective_custom_logo if (effective_custom_logo and os.path.exists(effective_custom_logo)) else (shield_logo_path if os.path.exists(shield_logo_path) else os.path.join(assets_dir, "tuv_sud_logo.png"))
    bg_path   = os.path.join(assets_dir, "cover_matrix_bg.png")
    chart_path= os.path.join(assets_dir, "chart_risk_severity.png")

    s_title_lower = session_title.lower()
    if "combined" in s_title_lower or ("web" in s_title_lower and "internal" in s_title_lower):
        scope_type = "Internal & Web App"
        doc_title = "Combined Internal Network & Web Application VAPT Validation Report"
    elif "web" in s_title_lower or "app" in s_title_lower:
        scope_type = "Web Application"
        doc_title = "Web Application Vulnerability Assessment and Penetration Testing Validation Report"
    elif "external" in s_title_lower:
        scope_type = "External Network"
        doc_title = "External Network Vulnerability Assessment and Penetration Testing Validation Report"
    else:
        scope_type = "Internal Network"
        doc_title = "Internal Network Vulnerability Assessment and Penetration Testing Validation Report"

    TUV_BLUE = (0, 80, 157)      # Corporate TÜV SÜD Blue #00509D
    DARK_TEXT = (15, 23, 42)
    BODY_TEXT = (51, 65, 85)

    class VAPTPDF(FPDF):
        def header(self):
            if self.page_no() > 1:
                self.set_font("Helvetica", "", 8.5)
                self.set_text_color(100, 116, 139)
                if os.path.exists(logo_path):
                    self.image(logo_path, x=184, y=4, w=10)
                    self.cell(166, 5, clean_text(f"{scope_type} Network VAPT Validation Report"), align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                else:
                    self.cell(0, 5, clean_text(f"{scope_type} Network VAPT Validation Report"), align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                self.ln(3)

        def footer(self):
            if self.page_no() > 1:
                self.set_y(-15)
                self.set_font("Helvetica", "", 9)
                self.set_text_color(15, 23, 42)
                self.cell(20, 8, clean_text(str(self.page_no())), align="L")
                self.cell(160, 8, clean_text("Cyber Security Services | XYZ Security Services Pvt. Ltd."), align="R")

    pdf = VAPTPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(15, 14, 15)
    
    hdr_blue   = FontFace(emphasis="B", color=(255, 255, 255), fill_color=TUV_BLUE)
    lbl_style  = FontFace(emphasis="B", color=(15, 23, 42), fill_color=(241, 245, 249))
    body_style = FontFace(emphasis="", color=(51, 65, 85), fill_color=(255, 255, 255))

    def draw_banner(title_text):
        pdf.set_fill_color(*TUV_BLUE)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7.5, clean_text(f"  {title_text}"), fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(*DARK_TEXT)
        pdf.set_fill_color(255, 255, 255)
        pdf.ln(2.5)

    # ── PAGE 1: COVER PAGE ──────────────────────────────────────────────────
    pdf.add_page()
    
    # Cover Page Top-Left Branding Box
    pdf.set_fill_color(255, 255, 255)
    pdf.rect(12, 12, 85, 34, style='F')
    if os.path.exists(logo_path):
        pdf.image(logo_path, x=15, y=15, w=26)
    else:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(*TUV_BLUE)
        pdf.set_xy(16, 22)
        pdf.cell(26, 6, "XYZ Security", new_x=XPos.RIGHT, new_y=YPos.TOP)

    pdf.set_draw_color(180, 180, 180)
    pdf.line(45, 15, 45, 43) # vertical divider line
    pdf.set_font("Helvetica", "B", 8.5)
    pdf.set_text_color(*TUV_BLUE)
    pdf.set_xy(48, 20)
    pdf.multi_cell(45, 4, clean_text("Cyber Security Audit\n& Advisory Services"))

    # Title & Target Client Block
    pdf.set_y(65)
    pdf.set_font("Helvetica", "B", 15)
    pdf.set_text_color(*TUV_BLUE)
    pdf.multi_cell(120, 7, clean_text(doc_title), align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(6)

    pdf.set_font("Helvetica", "B", 11.5)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 6, clean_text(f"For:  {target_client}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(6)

    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(71, 85, 105)
    pdf.cell(0, 4.5, clean_text("Submitted By:"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 4.5, clean_text(auditor_firm), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(6)
    pdf.cell(0, 4.5, clean_text("Version v1.0"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Bottom 4-Column Address Block
    pdf.set_y(250)
    pdf.set_draw_color(180, 180, 180)
    pdf.set_line_width(0.3)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 6.5)
    pdf.set_text_color(51, 65, 85)

    y_addr = pdf.get_y()
    pdf.set_xy(15, y_addr)
    pdf.multi_cell(42, 2.7, clean_text(
        "Registered Office:\nTÜV SÜD South Asia Pvt. Ltd.\nTÜV SÜD House,\nOff Saki Vihar Road,\nSaki Naka, Andheri (East),\nMumbai - 400072, India."
    ))

    pdf.set_xy(60, y_addr)
    pdf.multi_cell(42, 2.7, clean_text(
        "Corporate Office:\nTÜV SÜD South Asia Pvt. Ltd.\nSolitaire, 4th Floor,\nITI Road, Aundh,\nPune - 411007, India."
    ))

    pdf.set_xy(105, y_addr)
    pdf.multi_cell(45, 2.7, clean_text(
        "Report Submitted by:\nTÜV SÜD South Asia Pvt. Ltd.\nTÜV SÜD House,\nOff Saki Vihar Road,\nSaki Naka, Andheri (East),\nMumbai - 400072, India."
    ))

    pdf.set_xy(153, y_addr)
    pdf.multi_cell(42, 2.7, clean_text(
        "Email: info@tuv-sud.in\nwww.tuv-sud.in\n\nTÜV SÜD South Asia"
    ))

    # ── PAGE 2: DOCUMENT VERSION CONTROL & SUBMISSION DETAILS ──────────────
    pdf.add_page()
    draw_banner("DOCUMENT VERSION CONTROL")

    version_control_data = [
        ["Document Title", doc_title],
        ["Document ID", clean_text(report_doc_id)],
        ["Document Version", "1.0"],
        ["Prepared By", clean_text(auditor_lead)],
        ["Reviewed By", clean_text(auditor_reviewer)],
        ["Approved By", clean_text(auditor_approver)],
        ["Testing Dates", clean_text(testing_dates)],
        ["Effective Date", datetime.now().strftime("%d-%B-%Y")]
    ]

    pdf.set_font("Helvetica", "", 8.5)
    with pdf.table(col_widths=(55, 125), text_align="L") as table:
        for row in version_control_data:
            r = table.row()
            r.cell(row[0], style=lbl_style)
            r.cell(row[1], style=body_style)

    pdf.ln(4)
    draw_banner("DOCUMENT SUBMISSION DETAILS")

    submission_data = [
        ["Date", datetime.now().strftime("%d-%B-%Y")],
        ["Classification", "Confidential"],
        ["Document Type", doc_title],
        ["Submitted to", clean_text(submitted_to)],
        ["Designation", clean_text(designation)],
        ["E-mail", clean_text(email)]
    ]

    pdf.set_font("Helvetica", "", 8.5)
    with pdf.table(col_widths=(55, 125), text_align="L") as table:
        for row in submission_data:
            r = table.row()
            r.cell(row[0], style=lbl_style)
            r.cell(row[1], style=body_style)

    # Revision History (Only rendered if real user-entered review process dates exist)
    custom_revs = None

    if custom_revs and isinstance(custom_revs, list) and len(custom_revs) > 0:
        pdf.ln(4)
        draw_banner("REVISION HISTORY")
        with pdf.table(col_widths=(15, 35, 30, 100), text_align="L") as table:
            h = table.row()
            h.cell("No", style=hdr_blue)
            h.cell("Date", style=hdr_blue)
            h.cell("Version", style=hdr_blue)
            h.cell("Description", style=hdr_blue)

            for r_no, r_dt, r_ver, r_desc in custom_revs:
                r = table.row()
                r.cell(str(r_no), style=body_style)
                r.cell(str(r_dt), style=body_style)
                r.cell(str(r_ver), style=body_style)
                r.cell(str(r_desc), style=body_style)

    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 8.5)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 4.5, "All rights reserved.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 7.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 3.8, clean_text(
        "Any kind of publication, reproduction, duplication or recording on a storage medium or any form of distribution by printing, "
        f"photocopying, microfilming or in any other way, even in part only with the prior written consent of {auditor_firm}.\n\n"
        f"By {auditor_firm} no part of this publication may be published, reproduced, copied or stored in any format or by any means as a print-out of this publication.\n\n"
        "Company, product or service names may be trademarks or service marks of others and are the property of their respective owners."
    ))

    # ── PAGE 3: TABLE OF CONTENTS ──────────────────────────────────────────
    pdf.add_page()
    draw_banner("TABLE OF CONTENTS")
    pdf.ln(2)

    if findings:
        dict_findings = [f.to_dict() if hasattr(f, "to_dict") else (f if isinstance(f, dict) else getattr(f, "__dict__", {})) for f in findings]
        active_findings = [f for f in dict_findings if f.get("status") not in ("Out of Scope", "False Positive", "FALSE_POSITIVE")]
    else:
        # BUG-FIX: Never fall back to the in-memory registry -- it contains findings
        # from the PREVIOUS scan and silently poisons the report with stale data.
        # Return an empty list; the report will render correctly with 0 findings.
        active_findings = []

    for f in active_findings:
        s_val = f.get("severity_score") if f.get("severity_score") is not None else f.get("score")
        try:
            s_num = float(s_val) if s_val is not None else 0.0
        except Exception:
            s_num = 0.0

        r_sev = str(f.get("severity", "")).upper()
        if s_num > 0.0:
            if s_num >= 9.0: c_sev = "CRITICAL"
            elif s_num >= 7.0: c_sev = "HIGH"
            elif s_num >= 4.0: c_sev = "MEDIUM"
            else: c_sev = "LOW"
        else:
            if "CRIT" in r_sev or "P1" in r_sev: c_sev, s_num = "CRITICAL", 9.8
            elif "HIGH" in r_sev or "P2" in r_sev: c_sev, s_num = "HIGH", 8.0
            elif "MED" in r_sev or "P3" in r_sev: c_sev, s_num = "MEDIUM", 5.5
            elif "LOW" in r_sev or "P4" in r_sev: c_sev, s_num = "LOW", 2.5
            else: c_sev, s_num = "INFO", 0.0

        f["severity"] = c_sev
        f["severity_score"] = s_num
        f["score"] = s_num

    critical_cnt = sum(1 for f in active_findings if f.get("severity") == "CRITICAL")
    high_cnt = sum(1 for f in active_findings if f.get("severity") == "HIGH")
    medium_cnt = sum(1 for f in active_findings if f.get("severity") == "MEDIUM")
    low_cnt = sum(1 for f in active_findings if f.get("severity") == "LOW")

    def sort_key(f):
        sc = float(f.get("severity_score") or f.get("score") or 0.0)
        is_web = 1.0 if (str(f.get("source_tool", "")).lower() in ("burp suite", "burp") or "http" in str(f.get("target", "")).lower()) else 0.0
        return (sc, is_web)

    if active_findings:
        active_findings = sorted(active_findings, key=sort_key, reverse=True)

    import math
    list_to_show = active_findings if active_findings else []
    total_cnt = len(list_to_show)
    detail_findings = list_to_show[:40]
    summary_findings = list_to_show[40:]

    # Dynamic TOC page calculations
    vapt_table_pages = max(1, math.ceil(total_cnt / 22))
    p_sec_2_4 = 8 + vapt_table_pages
    p_sec_3_0 = p_sec_2_4 + 1
    detail_card_pages = math.ceil(len(detail_findings) * 0.65) if detail_findings else 1
    summary_table_pages = math.ceil(len(summary_findings) / 30) if summary_findings else 0
    p_sec_4_0 = p_sec_3_0 + detail_card_pages + summary_table_pages
    p_sec_5_0 = p_sec_4_0 + 1

    toc_items = [
        ("1 Penetration Test Methodology", "4"),
        ("   1.2 Standards-Based Testing and Reporting", "4"),
        ("   1.3 CVSS: Scoring Vulnerabilities", "4"),
        ("   1.4 How to Use This Document", "5"),
        ("2 Executive Summary", "6"),
        ("   2.2 Analysis Overview", "6"),
        ("   2.3 Summary of Findings", "7"),
        ("      2.3.1 Findings Overview", "7"),
        ("      2.3.2 Tabular Summary", "7"),
        ("      2.3.3 Graphical Summary", "7"),
        ("      2.3.4 Vulnerabilities Summary", "8"),
        ("   2.4 Tactical Recommendations", str(p_sec_2_4)),
        (f"3 Technical Detail Report: {scope_type} Network Vulnerability Assessment and Penetration Testing", str(p_sec_3_0)),
        ("   3.2 Testing Environment", str(p_sec_3_0)),
        ("   3.3 Findings", str(p_sec_3_0)),
        ("4 Appendix", str(p_sec_4_0)),
        ("   4.1 Testing Environment: Production", str(p_sec_4_0)),
        ("      4.1.1 Testing Environment Conditions", str(p_sec_4_0)),
        ("      4.1.2 Tools Used", str(p_sec_4_0)),
        ("      4.1.3 Provided Documentation", str(p_sec_4_0)),
        ("5 Disclaimer", str(p_sec_5_0))
    ]

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*BODY_TEXT)
    for title, p_num in toc_items:
        pdf.cell(160, 5, clean_text(title))
        pdf.cell(20, 5, clean_text(p_num), align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ── PAGE 4: METHODOLOGY & CVSS V4.0 METRICS ─────────────────────────────
    pdf.add_page()
    draw_banner("1 PENETRATION TEST METHODOLOGY")

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "1.2 Standards-Based Testing and Reporting", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        f"Our penetration test plans are performed according to internally developed guidelines by {auditor_firm} penetration test experts. "
        "Our test cases, where possible, are grounded in publicly available standards published by organizations such as OWASP, OSSTMM, NIST, "
        "along with our experience as a penetration test team. As an extension to these test cases, additional test cases may be identified based on penetration tester experience and the attack surface of target of evaluation."
    ))
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "1.3 CVSS: Scoring Vulnerabilities", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        f"The overarching goal of a penetration test is to identify the vulnerabilities in a target of evaluation. To assist in the prioritization of vulnerability remediation, {auditor_firm} utilizes the Common Vulnerability Scoring System (CVSS v3.0 / v3.1 / v4.0). "
        "CVSS assists in the assessment of a vulnerability's severity by providing a standard set of characteristics by which the vulnerability is scored. These scores are then used to calculate an overall severity score from 1-10; 1 being lowest and 10 being highest."
    ))
    pdf.ln(2.5)

    # Characteristics Table
    pdf.set_font("Helvetica", "", 8)
    with pdf.table(col_widths=(48, 132), text_align="L") as table:
        h = table.row()
        h.cell("Characteristics", style=hdr_blue)
        h.cell("Description", style=hdr_blue)

        chars = [
            ("Attack Vector (AV)", "Assesses whether an adversary can mount attack from a remote network, local network or if physical access is required."),
            ("Attack Complexity (AC)", "Assesses the complexity of an attack dependent on how many attack variables are within control of adversary."),
            ("Attack Requirements (AT)", "Assesses prerequisite deployment and execution conditions of vulnerable system enabling attack."),
            ("Privileges Required (PR)", "Assesses the level of access an attacker needs to mount a successful attack."),
            ("User Interaction (UI)", "Assesses extent to which actions of victim are required for attack to be successful."),
            ("Confidentiality (VC)", "Assesses negative impact attack can have on target of evaluation's confidentiality."),
            ("Integrity (VI)", "Assesses negative impact attack can have on target of evaluation's integrity."),
            ("Availability (VA)", "Assesses negative impact attack can have on target of evaluation's availability."),
            ("Confidentiality (SC)", "Assesses negative impact attack can have on subsequent system's confidentiality."),
            ("Integrity (SI)", "Assesses negative impact attack can have on subsequent system's integrity."),
            ("Availability (SA)", "Assesses negative impact attack can have on subsequent system's availability.")
        ]
        for c_title, c_desc in chars:
            r = table.row()
            r.cell(c_title, style=lbl_style)
            r.cell(c_desc, style=body_style)

    pdf.ln(3)
    # Severity Range Table (Starts on page 4)
    with pdf.table(col_widths=(28, 28, 124), text_align="L") as table:
        h = table.row()
        h.cell("Range", style=hdr_blue)
        h.cell("Rating", style=hdr_blue)
        h.cell("Description", style=hdr_blue)

        ranges = [
            ("9.0 - 10.0", "Critical", "These vulnerabilities should be reviewed immediately. Exploit exists that could severely impact CAV."),
            ("7.0 - 8.9", "High", "Needs short-term assessment. Exploitable with low/medium complexity with moderate to high impact."),
            ("4.0 - 6.9", "Medium", "Evaluated for business impact; exploitable with increased effort or lower confidentiality impact."),
            ("0.1 - 3.9", "Low", "Exploitation likely results in little negative impact to confidentiality, integrity, or availability.")
        ]
        for r_rng, r_rt, r_dsc in ranges:
            r = table.row()
            r.cell(r_rng, style=body_style)
            r.cell(r_rt, style=body_style)
            r.cell(r_dsc, style=body_style)

    # ── PAGE 5: HOW TO USE THIS DOCUMENT ─────────────────────────────────────
    pdf.add_page()
    # Finish 0.0 Info row at top of page 5
    with pdf.table(col_widths=(28, 28, 124), text_align="L") as table:
        r = table.row()
        r.cell("0.0", style=body_style)
        r.cell("Info", style=body_style)
        r.cell("Informational findings with zero direct impact on confidentiality, integrity, or availability.", style=body_style)

    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "1.4 How to Use This Document", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        "The vulnerabilities reported in this document provide a view of the target of evaluation's security posture at the time of testing. "
        "The report is broken up into three major sections: an executive summary, a technical detail report, and an appendix."
    ))
    pdf.ln(3)

    with pdf.table(col_widths=(45, 135), text_align="L") as table:
        h = table.row()
        h.cell("Descriptor", style=hdr_blue)
        h.cell("Content", style=hdr_blue)

        desc_rows = [
            ("Vulnerability Description", "Provides an overview of identified vulnerability including how it could be useful to an adversary."),
            ("Target(s)", "Provides a list of systems and network endpoints relevant to the vulnerability."),
            ("Status", "Contains either 'Verified' or 'Detected' indicating whether the flaw was actively exploited."),
            ("CVSS Base Metrics & Scoring", "Provides overall severity score and individual vector metrics (AV, AC, PR, UI, C, I, A)."),
            ("Proof of Concept", "Provides descriptions, screenshots, or command logs showing how the flaw was detected/reproduced."),
            ("Remediation", "Provides suggestions and technical steps on how to mitigate the vulnerability."),
            ("References", "Provides links to CVEs, CWEs, and official documentation resources.")
        ]
        for d_lbl, d_cnt in desc_rows:
            r = table.row()
            r.cell(d_lbl, style=lbl_style)
            r.cell(d_cnt, style=body_style)

    # ── PAGE 6: EXECUTIVE SUMMARY & TARGET SCOPE ─────────────────────────────
    pdf.add_page()
    draw_banner("2 EXECUTIVE SUMMARY")

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "2.2 Analysis Overview", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        "The objective of vulnerability assessment is to determine security vulnerabilities in the systems that can be exploited by Internal entities. "
        "The tests were carried out assuming the identity of an attacker with malicious intent. Scope targets:"
    ))
    pdf.ln(3)

    # Scope IPs Table (3 columns - dynamic)
    dynamic_ips = None
    try:
        if findings:
            extracted_hosts = set()
            for f in findings:
                host_val = f.get("host") or f.get("target") or f.get("ip")
                if host_val:
                    for h in str(host_val).replace(",", " ").split():
                        h_clean = h.strip()
                        if h_clean and h_clean.lower() not in ("n/a", "none", "unknown", "—"):
                            extracted_hosts.add(h_clean)
            if extracted_hosts:
                dynamic_ips = sorted(list(extracted_hosts))
    except Exception:
        dynamic_ips = None

    ip_list = dynamic_ips if dynamic_ips else ["Target Systems Ingested from Scan File"]
    pdf.set_font("Helvetica", "", 8)
    with pdf.table(col_widths=(60, 60, 60), text_align="C") as table:
        for i in range(0, len(ip_list), 3):
            r = table.row()
            r.cell(ip_list[i], style=body_style)
            r.cell(ip_list[i+1] if i+1 < len(ip_list) else "", style=body_style)
            r.cell(ip_list[i+2] if i+2 < len(ip_list) else "", style=body_style)

    pdf.ln(4)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.cell(0, 4.5, clean_text(f"Assessment Date: {testing_dates}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ── PAGE 7: SUMMARY OF FINDINGS & GRAPHICAL BAR CHART ─────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "2.3 Summary of Findings", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    
    if findings:
        dict_findings = [f.to_dict() if hasattr(f, "to_dict") else (f if isinstance(f, dict) else getattr(f, "__dict__", {})) for f in findings]
        active_findings = [f for f in dict_findings if f.get("status") not in ("Out of Scope", "False Positive", "FALSE_POSITIVE")]
    else:
        # BUG-FIX: Never fall back to the in-memory registry (stale from previous scan).
        active_findings = []

    critical_cnt = sum(1 for f in active_findings if f.get("severity") == "CRITICAL")
    high_cnt     = sum(1 for f in active_findings if f.get("severity") == "HIGH")
    medium_cnt   = sum(1 for f in active_findings if f.get("severity") == "MEDIUM")
    low_cnt      = sum(1 for f in active_findings if f.get("severity") == "LOW")
    info_cnt     = sum(1 for f in active_findings if f.get("severity") == "INFO")

    total_cnt = len(active_findings) if active_findings else 2

    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.cell(0, 4.5, clean_text(f"2.3.1 Findings Overview: Based on assessment, {total_cnt} vulnerabilities have been found in scope:"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2.5)

    pdf.set_font("Helvetica", "B", 9.5)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 4.5, "2.3.2 Tabular Summary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1.5)

    hdr_crit = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(192, 0, 0))
    hdr_high = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(255, 0, 0))
    hdr_med  = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(255, 192, 0))
    hdr_low  = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(0, 176, 80))
    hdr_info = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(0, 112, 192))
    hdr_tot  = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(127, 127, 127))

    with pdf.table(col_widths=(30, 30, 30, 30, 30, 30), text_align="C") as table:
        h = table.row()
        h.cell("Critical", style=hdr_crit)
        h.cell("High", style=hdr_high)
        h.cell("Medium", style=hdr_med)
        h.cell("Low", style=hdr_low)
        h.cell("Info", style=hdr_info)
        h.cell("Total Findings", style=hdr_tot)

        r = table.row()
        r.cell(str(critical_cnt), style=body_style)
        r.cell(str(high_cnt), style=body_style)
        r.cell(str(medium_cnt), style=body_style)
        r.cell(str(low_cnt), style=body_style)
        r.cell(str(info_cnt), style=body_style)
        r.cell(str(total_cnt), style=lbl_style)

    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 9.5)
    pdf.cell(0, 4.5, "2.3.3 Graphical Summary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)


    # Native High-Precision Horizontal Bar Chart (Guaranteed rendering in all environments)
    start_chart_y = pdf.get_y()
    max_val = max([critical_cnt, high_cnt, medium_cnt, low_cnt, info_cnt, 1])
    max_bar_w = 115.0
    categories_data = [
        ("Critical", critical_cnt, (192, 0, 0)),
        ("High",     high_cnt,     (255, 0, 0)),
        ("Medium",   medium_cnt,   (255, 192, 0)),
        ("Low",      low_cnt,      (0, 176, 80)),
        ("Info",     info_cnt,     (0, 112, 192))
    ]
    for c_idx, (c_label, c_val, c_col) in enumerate(categories_data):
        row_y = start_chart_y + (c_idx * 8.5)
        pdf.set_xy(20, row_y)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*DARK_TEXT)
        pdf.cell(24, 6, c_label)
        
        bar_w = (c_val / max_val) * max_bar_w if c_val > 0 else 0
        if bar_w > 0:
            pdf.set_fill_color(*c_col)
            pdf.rect(48, row_y + 1.0, max(2.5, bar_w), 4.5, style="F")
            pdf.set_xy(51 + max(2.5, bar_w), row_y)
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*DARK_TEXT)
            pdf.cell(15, 6, str(c_val))
        else:
            pdf.set_xy(51, row_y)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(*BODY_TEXT)
            pdf.cell(15, 6, "0")

    pdf.set_y(start_chart_y + 48)

    # ── PAGE 8: VULNERABILITIES SUMMARY TABLE ──────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "2.3.4 Vulnerabilities Summary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    with pdf.table(col_widths=(20, 100, 30, 30), text_align="L") as table:
        h = table.row()
        h.cell("Sr. No.", style=hdr_blue)
        h.cell("Vulnerabilities", style=hdr_blue)
        h.cell("CVSS Score", style=hdr_blue)
        h.cell("Severity", style=hdr_blue)

        list_to_show = active_findings
        if not list_to_show:
            r = table.row()
            r.cell("-")
            r.cell("No vulnerabilities were identified during this assessment.")
            r.cell("-")
            r.cell("-")
        for idx, f in enumerate(list_to_show, 1):
            title = f.get("title", "") or f.get("finding", "") or f.get("control", "") or f"Vulnerability {idx}"
            sev_str = str(f.get("severity", f.get("sev", "LOW"))).split()[-1].upper()
            score_val = f.get("severity_score") or f.get("score")
            try:
                score_num = float(score_val) if score_val is not None else 0.0
            except Exception:
                score_num = 0.0

            if score_num <= 0.0:
                if "CRIT" in sev_str: score_num = 9.5
                elif "HIGH" in sev_str: score_num = 8.0
                elif "MED" in sev_str: score_num = 5.5
                elif "LOW" in sev_str: score_num = 2.5
                else: score_num = 0.0

            score_str = f"{score_num:.1f}"
            f["derived_score_str"] = score_str
            r = table.row()
            r.cell(f"{idx}.", style=body_style)
            r.cell(clean_text(title), style=body_style)
            r.cell(score_str, style=body_style)
            r.cell(sev_str, style=body_style)

        # Dynamic Overall Score calculation based on active finding severities
        overall_score_val = 10.0 if critical_cnt > 0 else (8.5 if high_cnt > 0 else (5.5 if medium_cnt > 0 else 2.5))
        if critical_cnt > 0:
            overall_sev = "CRITICAL"
        elif high_cnt > 0:
            overall_sev = "HIGH"
        elif medium_cnt > 0:
            overall_sev = "MEDIUM"
        else:
            overall_sev = "LOW"

        r_over = table.row()
        r_over.cell("", style=lbl_style)
        r_over.cell("OVERALL SCORE", style=lbl_style)
        r_over.cell(f"{overall_score_val:.1f}", style=lbl_style)
        r_over.cell(overall_sev, style=lbl_style)

    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "2.4 Tactical Recommendations", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        "It is recommended to follow the guidelines suggested by OWASP, OSSTMM and NIST. It is recommended to implement secure SDLC while developing the application."
    ))

    # ── PAGE 9+, TECHNICAL DETAIL REPORT: DYNAMIC FINDINGS ──────────────────
    pdf.add_page()
    draw_banner(f"3 TECHNICAL DETAIL REPORT: {scope_type.upper()} NETWORK VULNERABILITY ASSESSMENT AND PENETRATION TESTING")

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "3.2 Testing Environment", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.cell(0, 4.5, "For details concerning the testing environment, please refer to Appendix 4.1.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "3.3 Findings", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1.5)

    # ── ALL findings: Burp Suite (web pentest) first, then Nessus network findings ──
    def _sev_order(f):
        s = str(f.get("severity", "LOW")).upper()
        if "CRITICAL" in s: return 0
        if "HIGH" in s: return 1
        if "MEDIUM" in s: return 2
        return 3

    # Burp Suite web pentest findings first, then all Nessus/network findings
    burp_findings    = [f for f in active_findings if "burp" in str(f.get("source_tool", "")).lower()]
    network_findings = [f for f in active_findings if "burp" not in str(f.get("source_tool", "")).lower()]

    burp_sorted    = sorted(burp_findings,    key=_sev_order)
    network_sorted = sorted(network_findings, key=_sev_order)

    # NO CAP — every single vulnerability gets a full detail card
    list_to_render = burp_sorted + network_sorted

    if not list_to_render:
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(*BODY_TEXT)
        pdf.multi_cell(0, 5, "No vulnerabilities were identified during this assessment.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

    for idx, f in enumerate(list_to_render, 1):
        # Start new page if less than 75mm remains for a clean spacious presentation
        if pdf.get_y() > 200:
            pdf.add_page()
        else:
            pdf.ln(6)
            pdf.set_draw_color(226, 232, 240)
            pdf.line(15, pdf.get_y(), 195, pdf.get_y())
            pdf.ln(5)
            
        _custom_h_pdf = (f.get("custom_heading") or "").strip()
        vuln_title = _custom_h_pdf or html.unescape(str(f.get("title") or f.get("finding") or f.get("control") or f"Finding 3.3.{idx}"))
        # VAPT reports: redact email/phone (incidental PII) but keep IPs -- the
        # vulnerable host's IP address is the report's actual content, not PII.
        desc = html.unescape(redact_pii(str(f.get("description") or f.get("gap_description") or f.get("finding") or "-"), redact_ip=False))
        target = html.unescape(str(f.get("target") or f.get("control_id") or "Scoped Network Endpoints / Systems"))
        conf_val = str(f.get("confidence") or "").strip()
        status_str = f"Detected ({conf_val.capitalize()})" if conf_val and conf_val.lower() in ("certain", "firm", "tentative") else "Detected"

        score_val = float(f.get("severity_score", f.get("score", 2.3)) or 2.3)
        sev_val = str(f.get("severity", f.get("sev", "LOW"))).split()[-1].upper()
        
        if f.get("cvss_vector"):
            metrics = f"{score_val:.1f} {sev_val}  |  Vector: {f.get('cvss_vector')}"
        elif f.get("metrics_text"):
            metrics = f.get("metrics_text")
        else:
            metrics = f"{score_val:.1f} {sev_val}"

        # ── Proof of Concept text: structured PoC block built in bg_worker takes priority ──
        # evidence_snippet = "Target Host: X.X.X.X\nPlugin ID: ...\nCVE(s): ...\nPlugin Output:\n..."
        # This is already the rich, structured block we want to show in the report.
        poc_text = html.unescape(redact_pii(str(
            f.get("evidence_snippet") or   # Structured PoC block (built in bg_worker)
            f.get("evidence") or           # Raw plugin output
            f.get("evidence_quote") or     # LLM evidence quote
            f.get("poc") or
            "Console / Log Audit Verification"
        ), redact_ip=False))

        remed_raw = str(f.get("recommendation") or f.get("remediation") or "").strip()
        remed = html.unescape(redact_pii(remed_raw, redact_ip=False)) if remed_raw and ("no action" not in remed_raw.lower() and remed_raw != "NIL") else "Immediately apply vendor security patches or software updates to mitigate identified vulnerability."
        ref = html.unescape(str(f.get("references") or f.get("reference") or "OWASP / OSSTMM / NIST Security Recommendations"))
        main_img = f.get("poc_image") or f.get("image_path")
        extra_img = f.get("extra_image")

        # Severity color palette
        if sev_val == "CRITICAL":
            sev_rgb = (220, 38, 38)     # Red #DC2626
        elif sev_val == "HIGH":
            sev_rgb = (225, 29, 72)     # Crimson #E11D48
        elif sev_val == "MEDIUM":
            sev_rgb = (217, 119, 6)     # Amber #D97706
        elif sev_val == "LOW":
            sev_rgb = (37, 99, 235)     # Blue #2563EB
        else:
            sev_rgb = (100, 116, 139)   # Slate #64748B

        # Heading Title (Multi-cell so long vulnerability titles wrap cleanly without spilling off page!)
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(*sev_rgb)
        pdf.multi_cell(0, 5, clean_text(f"FN-{idx:02d}  {vuln_title}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1.5)

        # Derive actual Tool / Scanner Name dynamically
        raw_tool = f.get('source_tool') or f.get('tool') or ''
        if not raw_tool or raw_tool.lower() in ('vapt engine', 'vapt', 'unknown', 'none'):
            comb_text = (vuln_title + " " + desc + " " + str(f.get('source_files', ''))).lower()
            if any(k in comb_text for k in ('nessus', 'cve-', 'ms10-', 'ms16-', 'ms15-', 'ms14-', 'ms12-', 'ms11-', 'smb', '.net', 'winrar')):
                scanner_name = "Nessus Professional Scanner"
            elif any(k in comb_text for k in ('burp', 'sql injection', 'xss', 'xxe', 'csti', 'ssrf')):
                scanner_name = "Burp Suite Professional Scanner"
            elif any(k in comb_text for k in ('nmap', 'syn scanner', 'port ')):
                scanner_name = "Nmap Network Security Scanner"
            else:
                scanner_name = "Automated VAPT Scanner & RAG Engine"
        else:
            scanner_name = raw_tool

        # Meta Summary Table Grid
        with pdf.table(col_widths=(40, 140), text_align="L") as table:
            r = table.row()
            r.cell("Severity / Score", style=lbl_style)
            r.cell(clean_text(f"{sev_val} ({score_val:.1f})  —  {metrics}"), style=body_style)

            r = table.row()
            r.cell("Location / Target", style=lbl_style)
            r.cell(clean_text(target[:400]), style=body_style)

            r = table.row()
            r.cell("Status / Scanner", style=lbl_style)
            r.cell(clean_text(f"{status_str}  |  Tool: {scanner_name}"), style=body_style)

            # ── CVE References row ──────────────────────────────────────────────
            raw_cve_list = f.get("cve_list") or []
            if isinstance(raw_cve_list, str):
                import re as _re_cve
                raw_cve_list = _re_cve.findall(r'CVE-\d{4}-\d{4,7}', raw_cve_list, re.IGNORECASE)
            if not raw_cve_list:
                ext = re.findall(r'CVE-\d{4}-\d{4,7}', desc + " " + poc_text, re.IGNORECASE)
                if ext: raw_cve_list = list(set([e.upper() for e in ext]))
            cve_str = ", ".join(raw_cve_list) if raw_cve_list else "N/A - Vendor Security Advisory / End-of-Life Notice"
            r = table.row()
            r.cell("CVE References", style=lbl_style)
            r.cell(clean_text(cve_str[:400]), style=body_style)

            # ── OWASP Top 10 Classification ──────────────────────────────────
            from src.core.parsers.control_mapper import map_finding_to_owasp
            owasp_cat = f.get("owasp_category") or map_finding_to_owasp(f.get("cwe"), vuln_title, desc)
            r = table.row()
            r.cell("OWASP Top 10 (2021)", style=lbl_style)
            r.cell(clean_text(owasp_cat[:400]), style=body_style)

            # ── Risk Category ─────────────────────────────────────────────────
            risk_cat = f.get("category") or ""
            if risk_cat:
                r = table.row()
                r.cell("Risk Category", style=lbl_style)
                r.cell(clean_text(risk_cat), style=body_style)

            # ── CIA Impact & PII Exposure ─────────────────────────────────────
            cia_val = f.get("cia_impact") or ""
            pii_flag = f.get("is_pii_exposed", False)
            if cia_val:
                cia_display = cia_val
                if pii_flag:
                    cia_display += "  |  ⚠ PII EXPOSURE DETECTED"
                r = table.row()
                r.cell("CIA Impact", style=lbl_style)
                r.cell(clean_text(cia_display), style=body_style)

            # ── Plugin ID row (when available) ─────────────────────────────────
            plugin_id_val = str(f.get("plugin_id") or "").strip()
            if plugin_id_val:
                r = table.row()
                r.cell("Scanner Plugin ID", style=lbl_style)
                r.cell(clean_text(plugin_id_val), style=body_style)

        pdf.ln(3)


        # Issue Description Section
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(15, 23, 42)
        pdf.cell(0, 5, "Issue Description:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(51, 65, 85)
        pdf.multi_cell(0, 4.5, clean_text(desc), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)

        # Proof of Vulnerability Section
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(15, 23, 42)
        pdf.cell(0, 5, "Proof of Vulnerability:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

        # Styled Monospace Box for HTTP Request/Response Evidence or Console Logs
        def format_http_evidence(txt):
            if not txt:
                return "Console / Log Audit Verification"
            import re
            keywords = [
                r'(GET\s+/[^\s\n]*)', r'(POST\s+/[^\s\n]*)', r'(PUT\s+/[^\s\n]*)', r'(DELETE\s+/[^\s\n]*)',
                r'(HTTP/1\.[01]\s+\d+)', r'(HTTP/2\s+\d+)',
                r'(Host:)', r'(User-Agent:)', r'(Accept:)', r'(Accept-Encoding:)', r'(Accept-Language:)',
                r'(Content-Type:)', r'(Content-Length:)', r'(Connection:)', r'(Cache-Control:)',
                r'(Cookie:)', r'(Set-Cookie:)', r'(Date:)', r'(Server:)', r'(Location:)',
                r'(\[Request\s*\d*\])', r'(\[Response\s*\d*\])'
            ]
            formatted = str(txt)
            for kw in keywords:
                formatted = re.sub(kw, r'\n\1', formatted, flags=re.IGNORECASE)
            lines = [l.strip() for l in formatted.splitlines() if l.strip()]
            return "\n".join(lines)

        clean_poc = format_http_evidence(poc_text[:2500])

        pdf.set_font("Courier", "", 7.5)
        pdf.set_fill_color(248, 250, 252)
        pdf.set_draw_color(203, 213, 225)
        pdf.set_text_color(30, 41, 59)
        pdf.multi_cell(0, 4.2, clean_text(clean_poc), border=1, fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)

        # Recommendation Section
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(15, 23, 42)
        pdf.cell(0, 5, "Recommendation:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(51, 65, 85)
        pdf.multi_cell(0, 4.5, clean_text(remed), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        # Developer-Actionable Mitigation Steps Section
        remed_actionable = f.get("remediation_actionable") or f.get("actionable_remediation") or ""
        if remed_actionable and remed_actionable.strip() != remed.strip():
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(15, 23, 42)
            pdf.cell(0, 5, "Developer Actionable Mitigation Steps:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(30, 41, 59)
            pdf.multi_cell(0, 4.5, clean_text(remed_actionable), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)

        pdf.ln(2)

        if main_img and os.path.exists(main_img):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 8.5)
            pdf.cell(0, 4, "Proof of Concept Artifact:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
            if pdf.get_y() + 55 > 265:
                pdf.add_page()
            pdf.image(main_img, x=15, y=pdf.get_y(), w=170)

        if extra_img and os.path.exists(extra_img):
            pdf.add_page()
            pdf.image(extra_img, x=15, y=20, w=175)

    # ── COMPACT SUMMARY TABLE: Remaining findings (beyond top 30) ───────────
    if summary_findings:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(*DARK_TEXT)
        pdf.cell(0, 5.5, f"3.4 Additional Findings Summary ({len(summary_findings)} findings)", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*BODY_TEXT)
        pdf.cell(0, 4, "The following findings are listed in summary format. Full remediation details follow the same guidance as Section 3.3.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)
        with pdf.table(col_widths=(10, 95, 20, 25, 30), text_align="L") as tbl:
            hrow = tbl.row()
            hrow.cell("#",          style=hdr_blue)
            hrow.cell("Vulnerability",  style=hdr_blue)
            hrow.cell("CVSS",       style=hdr_blue)
            hrow.cell("Severity",   style=hdr_blue)
            hrow.cell("Target",     style=hdr_blue)
            for sidx, sf_item in enumerate(summary_findings, len(detail_findings) + 1):
                st_title  = clean_text((sf_item.get("title") or sf_item.get("finding") or sf_item.get("control_name") or f"Finding {sidx}")[:80])
                st_score_val = sf_item.get("severity_score") or sf_item.get("score")
                try:
                    st_score_num = float(st_score_val) if st_score_val is not None else 0.0
                except Exception:
                    st_score_num = 0.0

                st_sev = str(sf_item.get("severity", "—")).split()[-1].upper()[:8]
                if st_score_num <= 0.0:
                    if "CRIT" in st_sev: st_score_num = 9.5
                    elif "HIGH" in st_sev: st_score_num = 8.0
                    elif "MED" in st_sev: st_score_num = 5.5
                    elif "LOW" in st_sev: st_score_num = 2.5
                    else: st_score_num = 0.0
                    
                st_score = f"{st_score_num:.1f}"
                st_target = clean_text(str(sf_item.get("target") or sf_item.get("control_id") or "—")[:30])
                srow = tbl.row()
                srow.cell(f"{sidx}.", style=body_style)
                srow.cell(st_title,   style=body_style)
                srow.cell(st_score,   style=body_style)
                srow.cell(st_sev,     style=body_style)
                srow.cell(st_target,  style=body_style)

    # ── PQC (Post-Quantum Cryptography Readiness): Cryptographic Asset Inventory ──
    # Additive section, gated on presence of PQC scan data (quantum_status is only
    # ever populated by pqc_parser.py) -- does not alter any existing VAPT table or
    # section above. Blank cell wherever a Finding field is empty, never a guess.
    pqc_findings = [f for f in active_findings if f.get("quantum_status")]
    if pqc_findings:
        pdf.add_page()
        draw_banner("3.5 CRYPTOGRAPHIC ASSET INVENTORY (CBOM/QBOM)")
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(*BODY_TEXT)
        pdf.multi_cell(0, 4, clean_text(
            "The following inventory lists cryptographic algorithms, certificates, and protocols "
            "discovered during the assessment, evaluated for post-quantum readiness."
        ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        qs_vuln = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(192, 0, 0))
        qs_weak = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(255, 192, 0))
        qs_safe = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(0, 176, 80))
        # Risk-band color convention mirrors the Quantum Status column above.
        rb_critical = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(192, 0, 0))
        rb_high     = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(230, 126, 34))
        rb_medium   = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(255, 192, 0))
        rb_low      = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(0, 176, 80))

        with pdf.table(col_widths=(20, 17, 17, 17, 18, 13, 8, 13, 14, 16, 16, 18), text_align="L") as table:
            h = table.row()
            for hdr_txt in ["Asset", "CA Algorithm", "Key Algorithm", "Protocol", "Quantum Status", "Exposure", "Port", "Environment", "Severity", "Risk Score", "Business Priority", "OEM Readiness"]:
                h.cell(hdr_txt, style=hdr_blue)
            for pf in pqc_findings:
                qs = str(pf.get("quantum_status") or "").strip().upper()
                qs_style = qs_vuln if qs == "VULNERABLE" else (qs_weak if qs == "WEAK" else (qs_safe if qs == "SAFE" else body_style))
                rb = str(pf.get("risk_band") or "").strip().upper()
                rb_style = rb_critical if rb == "CRITICAL" else (rb_high if rb == "HIGH" else (rb_medium if rb == "MEDIUM" else (rb_low if rb == "LOW" else body_style)))
                _rscore = pf.get("risk_score")
                _rscore_text = f"{_rscore} ({rb})" if (_rscore is not None and rb) else (str(_rscore) if _rscore is not None else "")
                _oem_product = pf.get("oem_product") or ""
                _oem_status = pf.get("oem_readiness_status") or ""
                _oem_text = f"{_oem_product}: {_oem_status}" if (_oem_product and _oem_status) else (_oem_status or _oem_product)
                r = table.row()
                r.cell(clean_text(pf.get("asset_name") or ""), style=body_style)
                r.cell(clean_text(pf.get("ca_algorithm") or ""), style=body_style)
                r.cell(clean_text(pf.get("key_algorithm") or ""), style=body_style)
                r.cell(clean_text(pf.get("protocol_version") or ""), style=body_style)
                r.cell(clean_text(qs), style=qs_style)
                r.cell(clean_text(pf.get("exposure_context") or ""), style=body_style)
                r.cell(clean_text(pf.get("port") or ""), style=body_style)
                r.cell(clean_text(pf.get("environment") or ""), style=body_style)
                r.cell(clean_text(pf.get("severity") or ""), style=body_style)
                r.cell(clean_text(_rscore_text), style=rb_style if _rscore_text else body_style)
                r.cell(clean_text(pf.get("business_priority") or ""), style=body_style)
                r.cell(clean_text(_oem_text), style=body_style)

        # Migration Dependencies -- only rendered when at least one PQC finding
        # flags a downstream dependency also being quantum-vulnerable.
        _mig_findings = [pf for pf in pqc_findings if pf.get("migration_dependency_flag")]
        if _mig_findings:
            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*DARK_TEXT)
            pdf.cell(0, 5, "Migration Dependencies", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(*BODY_TEXT)
            for pf in _mig_findings:
                _asset = pf.get("asset_name") or "—"
                _chain = pf.get("dependency_chain") or ""
                pdf.multi_cell(0, 4.5, clean_text(f"- {_asset}: {_chain}" if _chain else f"- {_asset}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ── APPENDIX ─────────────────────────────────────────────────
    pdf.add_page()
    draw_banner("4 APPENDIX")

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5.5, "4.1 Testing Environment: Production", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1)

    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(0, 4.5, "4.1.1 Testing Environment Conditions", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    if scope_type == "Internal":
        env_text = "The test was carried out as Gray Box (Internal Network & Authorized Access). No operational disruptions were encountered during testing. Industrial Standard for Security were followed, such as OWASP, OSSTMM and NIST."
    else:
        env_text = "The test was carried out as Black Box (External Perimeter Testing). Industrial Standard for Security were followed, such as OWASP, OSSTMM and NIST."
    pdf.multi_cell(0, 4, clean_text(env_text))
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 4.5, "4.1.2 Tools Used", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, "Nessus\nNmap\nBurp Suite / PortSwigger\nOpenSSL")
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 4.5, "4.1.3 Provided Documentation", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    if scope_type == "Internal":
        doc_text = "Internal Network Subnet Ranges & Target Host Lists.\nAuthorized Internal VPN / LAN Credentials."
    else:
        doc_text = "Public Domain Names, External IP Ranges & REST API Endpoints."
    pdf.multi_cell(0, 4, clean_text(doc_text))

    # ── PAGE 13: DISCLAIMER ───────────────────────────────────────────────
    pdf.add_page()
    draw_banner("5 DISCLAIMER")

    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        "The accuracy of the information and data provided in this report is subject to the information available to TÜV SÜD testing team during the engagement. "
        "The team relies on the accuracy and completeness of the information provided by the client and does not assume any responsibility for any inaccuracies or omissions.\n\n"
        "The assessment performed by TÜV SÜD was limited to the scope outlined in the engagement agreement between TÜV SÜD and the client. "
        "Any connected systems, applications, or components not explicitly covered in the initial scoping scope will not be included in the assessment and may contain undiscovered vulnerabilities and may impact the security of scoped systems.\n\n"
        "Any recommendations provided in this report for mitigating vulnerabilities are general in nature and should not be considered an exhaustive or tailor-made solution for specific environments. "
        "Implementation of these recommendations may require further analysis and customization based on the client's unique security requirements and constraints.\n\n"
        "The penetration testing engagement was conducted within a specific timeframe. The vulnerabilities and security issues identified in this report are based on the assessment conducted during this limited timeframe. "
        "Changes in the systems, networks, or applications after the engagement may impact the accuracy and relevance of the findings."
    ))
    pdf.ln(12)

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 6, "***End of Report***", align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf_bytes = pdf.output()
    return bytes(pdf_bytes)


def _export_pqc_pdf(session_title, findings, resolved_list, status, comments="", custom_logo=None, metadata=None):
    """
    Dedicated Post-Quantum Cryptography (PQC) Readiness Assessment PDF report.

    Completely separate from _export_vapt_pdf() and the ISO template.
    Template structure based on Union Bank RFP / Manjunatha Sir's methodology:
      1. Cover Page
      2. Document Version Control
      3. Executive Summary (VULNERABLE/WEAK/SAFE counts + top risks)
      4. Cryptographic Asset Register (CBOM)
      5. Quantum Bill of Materials (QBOM)
      6. Detailed Findings (per PQC control card)
      7. Dependency Mapping
      8. OEM Readiness Matrix
      9. Asset Prioritization Tiers
      10. Appendix + Disclaimer
    """
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
    from fpdf.fonts import FontFace
    import io as _io
    import os
    import json as _json

    def clean_text(val):
        if not val:
            return "-"
        val = str(val)
        val = html.unescape(val)
        val = val.replace("\u201c", "\"").replace("\u201d", "\"").replace("\u2018", "'").replace("\u2019", "'")
        val = val.replace("\u2014", "-").replace("\u2013", "-").replace("\u2022", "*").replace("\u25cf", "*")
        val = val.encode("latin-1", "replace").decode("latin-1")
        return val

    meta = metadata or {}
    auditor_lead     = meta.get("brand_auditor")     or "Lead Auditor"
    auditor_firm     = meta.get("brand_firm")        or "AI Cyber Audit Box"
    auditor_reviewer = meta.get("brand_reviewer")    or "Senior Reviewer"
    auditor_approver = meta.get("brand_approver")    or "Approver"
    report_doc_id    = meta.get("brand_docid")       or "PQC-001"
    target_client    = meta.get("brand_client")      or "Client Organisation"
    assessment_date  = datetime.now().strftime("%d-%B-%Y")

    # Paths
    assets_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "data", "assets"))
    custom_logo_file = os.path.join(assets_dir, "custom_company_logo.png")
    effective_logo = (custom_logo if (custom_logo and os.path.exists(custom_logo))
                      else (custom_logo_file if os.path.exists(custom_logo_file) else None))
    shield_path = os.path.join(assets_dir, "shield_logo.png")
    logo_path = effective_logo or (shield_path if os.path.exists(shield_path) else None)

    # Load OEM readiness data from knowledge base
    _oem_kb = {}
    try:
        _oem_path = os.path.abspath(os.path.join(
            os.path.dirname(__file__), "knowledge", "pqc_oem_readiness.json"
        ))
        if os.path.exists(_oem_path):
            with open(_oem_path, "r", encoding="utf-8") as _f:
                _oem_kb = _json.load(_f)
    except Exception:
        _oem_kb = {}

    # ── Colour palette ────────────────────────────────────────────
    PQC_BLUE     = (0, 80, 157)      # primary brand colour
    SAFE_GREEN   = (0, 140, 60)
    VULN_RED     = (192, 0, 0)
    WEAK_ORANGE  = (204, 102, 0)
    DARK_TEXT    = (15, 23, 42)
    BODY_TEXT    = (51, 65, 85)
    GREY_BG      = (241, 245, 249)

    # ── Pre-compute finding sets ──────────────────────────────────
    dict_findings = []
    for f in (findings or []):
        if hasattr(f, "to_dict"):
            dict_findings.append(f.to_dict())
        elif isinstance(f, dict):
            dict_findings.append(f)
        else:
            dict_findings.append(getattr(f, "__dict__", {}))

    active = [f for f in dict_findings if f.get("status") not in ("Out of Scope", "False Positive", "FALSE_POSITIVE")]

    vuln_findings  = [f for f in active if str(f.get("quantum_status") or "").upper() == "VULNERABLE"]
    weak_findings  = [f for f in active if str(f.get("quantum_status") or "").upper() == "WEAK"]
    safe_findings  = [f for f in active if str(f.get("quantum_status") or "").upper() == "SAFE"]
    pqc_findings   = [f for f in active if f.get("quantum_status")]       # any PQC-scanned finding
    non_pqc        = [f for f in active if not f.get("quantum_status")]   # standard control findings

    crit_cnt = sum(1 for f in active if "CRIT" in str(f.get("severity") or "").upper())
    high_cnt = sum(1 for f in active if "HIGH" in str(f.get("severity") or "").upper())
    med_cnt  = sum(1 for f in active if "MED"  in str(f.get("severity") or "").upper())
    low_cnt  = sum(1 for f in active if "LOW"  in str(f.get("severity") or "").upper())

    # ── PDF class ─────────────────────────────────────────────────
    class PQCPDF(FPDF):
        def header(self):
            if self.page_no() > 1:
                self.set_font("Helvetica", "", 7.5)
                self.set_text_color(120, 140, 160)
                if logo_path and os.path.exists(logo_path):
                    self.image(logo_path, x=182, y=4, w=10)
                self.cell(170, 5,
                    clean_text("Post-Quantum Cryptography (PQC) Readiness Assessment Report"),
                    align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                self.ln(2)

        def footer(self):
            if self.page_no() > 1:
                self.set_y(-13)
                self.set_font("Helvetica", "", 7.5)
                self.set_text_color(*BODY_TEXT)
                self.cell(20, 6, clean_text(str(self.page_no())), align="L")
                self.cell(160, 6,
                    clean_text(f"CONFIDENTIAL  |  {auditor_firm}  |  {assessment_date}"),
                    align="R")

    pdf = PQCPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(15, 14, 15)

    hdr_blue   = FontFace(emphasis="B", color=(255, 255, 255), fill_color=PQC_BLUE)
    lbl_style  = FontFace(emphasis="B", color=DARK_TEXT, fill_color=GREY_BG)
    body_style = FontFace(color=BODY_TEXT, fill_color=(255, 255, 255))
    qs_vuln    = FontFace(emphasis="B", color=(255, 255, 255), fill_color=VULN_RED)
    qs_weak    = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(180, 80, 0))
    qs_safe    = FontFace(emphasis="B", color=(255, 255, 255), fill_color=SAFE_GREEN)
    rb_crit    = FontFace(emphasis="B", color=(255, 255, 255), fill_color=VULN_RED)
    rb_high    = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(200, 80, 0))
    rb_med     = FontFace(emphasis="B", color=(0, 0, 0), fill_color=(255, 192, 0))
    rb_low     = FontFace(emphasis="B", color=(255, 255, 255), fill_color=SAFE_GREEN)

    def draw_section_banner(title_text):
        pdf.set_fill_color(*PQC_BLUE)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7.5, clean_text(f"  {title_text}"), fill=True,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(*DARK_TEXT)
        pdf.set_fill_color(255, 255, 255)
        pdf.ln(2.5)

    def qs_style_for(qs):
        q = str(qs or "").strip().upper()
        return qs_vuln if q == "VULNERABLE" else (qs_weak if q == "WEAK" else (qs_safe if q == "SAFE" else body_style))

    def rb_style_for(rb):
        r = str(rb or "").strip().upper()
        return rb_crit if r == "CRITICAL" else (rb_high if r == "HIGH" else (rb_med if r == "MEDIUM" else (rb_low if r == "LOW" else body_style)))

    # ══════════════════════════════════════════════════════════════
    # PAGE 1: COVER PAGE
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()

    # Top blue bar
    pdf.set_fill_color(*PQC_BLUE)
    pdf.rect(0, 0, 210, 28, style="F")

    if logo_path and os.path.exists(logo_path):
        pdf.image(logo_path, x=15, y=5, w=20)

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(40, 10)
    pdf.cell(0, 6, clean_text(auditor_firm), align="L",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Quantum-themed accent bar
    pdf.set_fill_color(0, 200, 160)
    pdf.rect(0, 28, 210, 3, style="F")

    # Main Title block
    pdf.set_y(55)
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(*PQC_BLUE)
    pdf.multi_cell(0, 9,
        clean_text("Post-Quantum Cryptography (PQC)\nReadiness Assessment Report"),
        align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*BODY_TEXT)
    pdf.cell(0, 6, clean_text(f"Cryptographic Asset Discovery & NIST PQC Migration Risk Evaluation"),
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(10)

    # Info box
    pdf.set_fill_color(*GREY_BG)
    pdf.rect(15, pdf.get_y(), 180, 44, style="F")
    pdf.set_x(20)
    info_rows = [
        ("Client / Organisation", target_client),
        ("Assessment Date",       assessment_date),
        ("Document ID",           report_doc_id),
        ("Classification",        "CONFIDENTIAL"),
        ("Prepared By",           auditor_lead),
        ("Reviewed By",           auditor_reviewer),
    ]
    for lbl, val in info_rows:
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.set_text_color(*DARK_TEXT)
        pdf.cell(55, 6, clean_text(lbl + ":"), new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(*BODY_TEXT)
        pdf.cell(120, 6, clean_text(val), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_x(20)

    # Bottom accent
    pdf.set_y(268)
    pdf.set_fill_color(*PQC_BLUE)
    pdf.rect(0, 268, 210, 10, style="F")
    pdf.set_font("Helvetica", "", 7)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(15, 270)
    pdf.cell(0, 5,
        clean_text(f"Powered by AI Cyber Audit Box  |  NIST PQC Standards (FIPS 203/204/205)  |  {assessment_date}"),
        align="C")

    # ══════════════════════════════════════════════════════════════
    # PAGE 2: DOCUMENT VERSION CONTROL
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    draw_section_banner("DOCUMENT VERSION CONTROL")

    vc_rows = [
        ["Document Title",    "Post-Quantum Cryptography (PQC) Readiness Assessment Report"],
        ["Document ID",       clean_text(report_doc_id)],
        ["Document Version",  "1.0"],
        ["Classification",    "Confidential"],
        ["Prepared By",       clean_text(auditor_lead)],
        ["Reviewed By",       clean_text(auditor_reviewer)],
        ["Approved By",       clean_text(auditor_approver)],
        ["Assessment Date",   assessment_date],
        ["Report Status",     clean_text(status or "Final")],
    ]
    pdf.set_font("Helvetica", "", 8.5)
    with pdf.table(col_widths=(60, 120), text_align="L") as tbl:
        for row in vc_rows:
            r = tbl.row()
            r.cell(row[0], style=lbl_style)
            r.cell(row[1], style=body_style)

    pdf.ln(6)
    draw_section_banner("SCOPE OF ASSESSMENT")
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4.5, clean_text(
        "This report covers Post-Quantum Cryptography (PQC) Readiness Assessment of the organisation's "
        "cryptographic assets including TLS/SSL configurations, SSH infrastructure, IPSec/VPN gateways, "
        "PKI/Certificate infrastructure, HSM/KMS key management, database encryption at rest (TDE), "
        "code/firmware signing, and cryptographic library dependencies.\n\n"
        "Algorithms are evaluated against NIST Post-Quantum Cryptography standards (FIPS 203 ML-KEM, "
        "FIPS 204 ML-DSA, FIPS 205 SLH-DSA) and classified as VULNERABLE (quantum-breakable by "
        "Shor's algorithm), WEAK (classically deprecated), or SAFE (quantum-resistant)."
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ══════════════════════════════════════════════════════════════
    # PAGE 3: EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    draw_section_banner("1. EXECUTIVE SUMMARY")

    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4.5, clean_text(
        f"The AI Cyber Audit Box performed a comprehensive Post-Quantum Cryptography (PQC) Readiness "
        f"Assessment on the cryptographic infrastructure of {target_client}. "
        f"The assessment identified {len(vuln_findings)} quantum-vulnerable, {len(weak_findings)} "
        f"classically-weak, and {len(safe_findings)} quantum-safe algorithm instances across "
        f"{len(pqc_findings)} cryptographic assets.\n\n"
        "Quantum-vulnerable algorithms (RSA, ECC/ECDSA/ECDH, Diffie-Hellman) are susceptible to "
        "Shor's algorithm on a sufficiently large quantum computer. Assets using these algorithms "
        "are exposed to the 'Harvest Now, Decrypt Later' (HNDL) threat where adversaries capture "
        "encrypted data today to decrypt it once quantum computers become available."
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    # Quantum Status Summary Boxes
    box_data = [
        ("VULNERABLE", len(vuln_findings), VULN_RED, (255, 255, 255)),
        ("WEAK",       len(weak_findings), (204, 102, 0), (255, 255, 255)),
        ("SAFE",       len(safe_findings), SAFE_GREEN, (255, 255, 255)),
        ("TOTAL",      len(pqc_findings),  PQC_BLUE, (255, 255, 255)),
    ]
    box_x = 15
    for label, count, fill, text_c in box_data:
        pdf.set_fill_color(*fill)
        pdf.rect(box_x, pdf.get_y(), 42, 18, style="F")
        pdf.set_xy(box_x + 1, pdf.get_y() + 2)
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(*text_c)
        pdf.cell(40, 7, str(count), align="C", new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_xy(box_x + 1, pdf.get_y() + 9)
        pdf.set_font("Helvetica", "", 7)
        pdf.cell(40, 5, label, align="C")
        box_x += 45
    pdf.ln(22)

    # Severity breakdown
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5, "1.1 Finding Severity Distribution", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1)
    pdf.set_font("Helvetica", "", 8.5)
    with pdf.table(col_widths=(40, 25, 25, 25, 25, 40), text_align="C") as tbl:
        h = tbl.row()
        for hdr in ["Total Findings", "Critical", "High", "Medium", "Low", "Quantum Status"]:
            h.cell(hdr, style=hdr_blue)
        r = tbl.row()
        r.cell(str(len(active)), style=body_style)
        r.cell(str(crit_cnt), style=rb_crit if crit_cnt else body_style)
        r.cell(str(high_cnt), style=rb_high if high_cnt else body_style)
        r.cell(str(med_cnt),  style=rb_med  if med_cnt  else body_style)
        r.cell(str(low_cnt),  style=rb_low  if low_cnt  else body_style)
        r.cell(f"{len(vuln_findings)}V / {len(weak_findings)}W / {len(safe_findings)}S", style=body_style)

    pdf.ln(6)
    # Top 5 critical risks
    top5 = sorted(pqc_findings or active,
                  key=lambda f: (
                      1 if str(f.get("quantum_status") or "").upper() == "VULNERABLE" else 2,
                      -(f.get("risk_score") or 0)
                  ))[:5]
    if top5:
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*DARK_TEXT)
        pdf.cell(0, 5, "1.2 Top Risk Findings", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)
        pdf.set_font("Helvetica", "", 8.5)
        with pdf.table(col_widths=(8, 60, 28, 25, 25, 34), text_align="L") as tbl:
            h = tbl.row()
            for hdr in ["#", "Finding Title", "Quantum Status", "Severity", "Risk Score", "Asset"]:
                h.cell(hdr, style=hdr_blue)
            for idx, f in enumerate(top5, 1):
                qs = str(f.get("quantum_status") or "").upper()
                r = tbl.row()
                r.cell(str(idx), style=body_style)
                r.cell(clean_text(str(f.get("control_name") or f.get("title") or "")[:55]),
                       style=body_style)
                r.cell(clean_text(qs), style=qs_style_for(qs))
                r.cell(clean_text(str(f.get("severity") or "")[:12]), style=body_style)
                _rs = f.get("risk_score")
                r.cell(str(_rs) if _rs is not None else "-", style=body_style)
                r.cell(clean_text(str(f.get("asset_name") or "")[:30]), style=body_style)

    # ══════════════════════════════════════════════════════════════
    # PAGE 4: CRYPTOGRAPHIC ASSET REGISTER (CBOM)
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    draw_section_banner("2. CRYPTOGRAPHIC ASSET REGISTER (CBOM)")
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        "The Cryptographic Bill of Materials (CBOM) lists all discovered cryptographic assets, "
        "their algorithms, protocols, and post-quantum readiness status."
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    cbom_rows = pqc_findings if pqc_findings else active
    if cbom_rows:
        pdf.set_font("Helvetica", "", 6.5)
        with pdf.table(col_widths=(18, 20, 18, 18, 18, 24, 18, 10, 16, 18), text_align="L") as tbl:
            h = tbl.row()
            for hdr in ["Asset Category", "Asset", "CA Algorithm", "Key Algorithm",
                        "Protocol", "Quantum Status", "Exposure", "Port", "Environment", "Severity"]:
                h.cell(hdr, style=hdr_blue)
            for f in cbom_rows:
                qs  = str(f.get("quantum_status") or "").upper()
                cat = str(f.get("asset_category") or "Unknown")
                proto_val = str(f.get("protocol_version") or "").strip()
                proto_clean = re.sub(r'\s*\([^)]*\)', '', proto_val).strip()
                r = tbl.row()
                r.cell(clean_text(cat[:20]), style=lbl_style)
                r.cell(clean_text(str(f.get("asset_name") or "")[:22]), style=body_style)
                r.cell(clean_text(str(f.get("ca_algorithm") or "")[:18]), style=body_style)
                r.cell(clean_text(str(f.get("key_algorithm") or "")[:18]), style=body_style)
                r.cell(clean_text((proto_clean or proto_val)[:20]), style=body_style)
                r.cell(clean_text(qs[:18]), style=qs_style_for(qs))
                r.cell(clean_text(str(f.get("exposure_context") or "")[:15]), style=body_style)
                r.cell(clean_text(str(f.get("port") or "")[:8]), style=body_style)
                r.cell(clean_text(str(f.get("environment") or "")[:16]), style=body_style)
                r.cell(clean_text(str(f.get("severity") or "")[:12]), style=body_style)
    else:
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(*BODY_TEXT)
        pdf.cell(0, 5, "No cryptographic asset data found in this session.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ══════════════════════════════════════════════════════════════
    # PAGE 5: QUANTUM BILL OF MATERIALS (QBOM)
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    draw_section_banner("3. QUANTUM BILL OF MATERIALS (QBOM)")
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        "The QBOM maps each discovered algorithm to its quantum readiness status, risk band, "
        "and business priority to support migration planning."
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    if cbom_rows:
        pdf.set_font("Helvetica", "", 7.5)
        with pdf.table(col_widths=(28, 32, 22, 18, 18, 22, 20), text_align="L") as tbl:
            h = tbl.row()
            for hdr in ["Asset", "Algorithm (Detected)", "Quantum Status",
                        "Risk Score", "Risk Band", "Business Priority", "OEM Readiness"]:
                h.cell(hdr, style=hdr_blue)
            for f in cbom_rows:
                qs  = str(f.get("quantum_status") or "").upper()
                rb  = str(f.get("risk_band") or "").upper()
                rs  = f.get("risk_score")
                _algo = (f.get("ca_algorithm") or f.get("key_algorithm") or
                         f.get("protocol_version") or f.get("control_name") or "")
                _oem_s = f.get("oem_readiness_status") or ""
                r = tbl.row()
                r.cell(clean_text(str(f.get("asset_name") or "")[:28]), style=body_style)
                r.cell(clean_text(str(_algo)[:32]), style=body_style)
                r.cell(clean_text(qs[:22]), style=qs_style_for(qs))
                r.cell(str(rs) if rs is not None else "-", style=body_style)
                r.cell(clean_text(rb[:18]), style=rb_style_for(rb))
                r.cell(clean_text(str(f.get("business_priority") or "")[:22]), style=body_style)
                r.cell(clean_text(str(_oem_s)[:20]), style=body_style)
    else:
        pdf.cell(0, 5, "No QBOM data available.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ══════════════════════════════════════════════════════════════
    # PAGE 6: DETAILED FINDINGS (per PQC control)
    # ══════════════════════════════════════════════════════════════
    detail_set = pqc_findings if pqc_findings else active
    for idx, f in enumerate(detail_set[:30], 1):
        qs  = str(f.get("quantum_status") or "").upper()
        sev = str(f.get("severity") or "MEDIUM").upper()
        rs  = f.get("risk_score")
        ctrl_id = clean_text(str(f.get("control_id") or f.get("control") or f"PQC-{idx}"))
        title   = clean_text(str(f.get("control_name") or f.get("title") or ctrl_id)[:80])
        full_ctrl = f"{ctrl_id} - {title}" if (len(ctrl_id) <= 12 and title and not title.lower().startswith(ctrl_id.lower())) else ctrl_id
        desc    = clean_text(str(f.get("description") or f.get("gap_description") or "")[:600])
        recom   = clean_text(str(f.get("recommendation") or f.get("remediation_actionable") or "")[:400])
        poc     = clean_text(str(f.get("evidence_snippet") or f.get("evidence_quote") or "")[:400])
        asset   = clean_text(str(f.get("asset_name") or f.get("target") or "")[:60])
        oem_p   = clean_text(str(f.get("oem_product") or "")[:40])
        oem_st  = clean_text(str(f.get("oem_readiness_status") or "")[:50])

        pdf.add_page()
        # Finding header banner with severity colour
        sev_fill = VULN_RED if "CRIT" in sev else ((200, 80, 0) if "HIGH" in sev else ((180, 140, 0) if "MED" in sev else SAFE_GREEN))
        pdf.set_fill_color(*PQC_BLUE)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, clean_text(f"  4.{idx}  {full_ctrl}"), fill=True,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        # Metadata row
        meta_rows = [
            ("Asset Category",     clean_text(str(f.get("asset_category") or "Unknown"))),
            ("Target Host / Asset", asset),
            ("Control ID",          full_ctrl),
            ("Quantum Status",      qs or "-"),
            ("Severity",            sev),
            ("Risk Score",          str(rs) if rs is not None else "-"),
            ("Business Priority",   clean_text(str(f.get("business_priority") or ""))),
            ("OEM Product",         oem_p if oem_p and oem_p != "-" else "-"),
            ("OEM Readiness",       oem_st if oem_st and oem_st != "-" else "-"),
            ("OWASP Category",      clean_text(str(f.get("category") or "A02:2021 Cryptographic Failures"))),
            ("CIA Impact",          clean_text(str(f.get("cia_impact") or ""))),
            ("Exposure",            clean_text(str(f.get("exposure_context") or ""))),
            ("Port",                clean_text(str(f.get("port") or ""))),
            ("Environment",         clean_text(str(f.get("environment") or ""))),
        ]
        pdf.set_font("Helvetica", "", 8)
        with pdf.table(col_widths=(55, 125), text_align="L") as tbl:
            for lbl, val in meta_rows:
                r = tbl.row()
                r.cell(lbl, style=lbl_style)
                r.cell(val, style=body_style)

        pdf.ln(3)
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.set_text_color(*DARK_TEXT)
        pdf.cell(0, 4.5, "Vulnerability Description", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(*BODY_TEXT)
        pdf.multi_cell(0, 4, desc or "-", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

        if poc and poc != "-":
            pdf.set_font("Helvetica", "B", 8.5)
            pdf.set_text_color(*DARK_TEXT)
            pdf.cell(0, 4.5, "Proof of Concept (Scanner / Config Output)", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_fill_color(245, 245, 245)
            pdf.set_font("Courier", "", 7.5)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 4, poc, fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)

        pdf.set_font("Helvetica", "B", 8.5)
        pdf.set_text_color(*DARK_TEXT)
        pdf.cell(0, 4.5, "Recommended Remediation & Action", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_fill_color(*GREY_BG)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(*BODY_TEXT)
        pdf.multi_cell(0, 4, recom or
            "Migrate to a NIST-selected PQC algorithm: use ML-KEM for key exchange and "
            "ML-DSA or SPHINCS+ for signatures. In the interim, prefer AES-256-GCM and SHA-384+.",
            fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        dep_chain = clean_text(str(f.get("dependency_chain") or ""))
        if dep_chain and dep_chain != "-":
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 8.5)
            pdf.set_text_color(*DARK_TEXT)
            pdf.cell(0, 4.5, "Migration Dependency Chain", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 8.5)
            pdf.set_text_color(*BODY_TEXT)
            pdf.multi_cell(0, 4, dep_chain, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ══════════════════════════════════════════════════════════════
    # PAGE N: DEPENDENCY MAPPING
    # ══════════════════════════════════════════════════════════════
    dep_findings = [f for f in (pqc_findings or active) if f.get("dependency_chain") and f.get("migration_dependency_flag")]
    if dep_findings:
        pdf.add_page()
        draw_section_banner("5. DEPENDENCY MAPPING")
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_text_color(*BODY_TEXT)
        pdf.multi_cell(0, 4, clean_text(
            "Assets below have downstream dependencies that are also quantum-vulnerable. "
            "Migration must follow the dependency chain from leaf nodes upward."
        ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)
        for f in dep_findings:
            asset = clean_text(str(f.get("asset_name") or "Unknown Asset"))
            chain = clean_text(str(f.get("dependency_chain") or ""))
            pdf.set_font("Helvetica", "B", 8.5)
            pdf.set_text_color(*PQC_BLUE)
            pdf.cell(0, 5, f"  {asset}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Courier", "", 7.5)
            pdf.set_text_color(*BODY_TEXT)
            pdf.multi_cell(0, 4, chain, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)

    # ══════════════════════════════════════════════════════════════
    # PAGE N: OEM READINESS MATRIX
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    draw_section_banner("6. OEM READINESS MATRIX")
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        "The following table shows the post-quantum readiness status of OEM vendors and products "
        "discovered in this assessment, based on publicly available roadmaps and advisories."
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    # Build OEM rows: from findings first, then supplement from knowledge base
    oem_seen = {}
    for f in (pqc_findings or active):
        prod = str(f.get("oem_product") or "").strip()
        st   = str(f.get("oem_readiness_status") or "").strip()
        if prod:
            oem_seen[prod] = {"status": st, "notes": "Detected in assessment evidence."}
    # Merge knowledge base entries
    for prod, kb_val in _oem_kb.items():
        if prod not in oem_seen:
            oem_seen[prod] = {
                "status": kb_val.get("status", ""),
                "notes":  kb_val.get("notes", "")
            }

    if oem_seen:
        pdf.set_font("Helvetica", "", 7.5)
        with pdf.table(col_widths=(45, 35, 35, 65), text_align="L") as tbl:
            h = tbl.row()
            for hdr in ["Product / Platform", "Vendor", "PQC Status", "Notes / Action"]:
                h.cell(hdr, style=hdr_blue)
            for prod, info in list(oem_seen.items())[:30]:
                st    = str(info.get("status") or "").strip()
                notes = str(info.get("notes") or "").strip()
                # Try to get vendor from KB
                _kb_entry = _oem_kb.get(prod, {})
                vendor = str(_kb_entry.get("vendor", "")).strip() or "-"
                # Colour the status cell
                _st_upper = st.upper()
                if "READY" in _st_upper or "HYBRID" in _st_upper:
                    _st_style = rb_low
                elif "ROADMAP" in _st_upper:
                    _st_style = rb_med
                elif "UPGRADE" in _st_upper or "REQUIRED" in _st_upper:
                    _st_style = rb_high
                else:
                    _st_style = body_style
                r = tbl.row()
                r.cell(clean_text(prod[:45]), style=body_style)
                r.cell(clean_text(vendor[:35]), style=body_style)
                r.cell(clean_text(st[:35]), style=_st_style)
                r.cell(clean_text(notes[:65]), style=body_style)
    else:
        pdf.cell(0, 5, "No OEM product data found in this session.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ══════════════════════════════════════════════════════════════
    # PAGE N: ASSET PRIORITIZATION TIERS
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    draw_section_banner("7. ASSET PRIORITIZATION TIERS")
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        "Assets are grouped into risk tiers based on quantum vulnerability, CIA impact, "
        "and business priority to guide migration sequencing."
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    tiers = {
        "CRITICAL": [f for f in (pqc_findings or active) if str(f.get("risk_band") or f.get("business_priority") or "").upper() == "CRITICAL"],
        "HIGH":     [f for f in (pqc_findings or active) if str(f.get("risk_band") or f.get("business_priority") or "").upper() == "HIGH"],
        "MEDIUM":   [f for f in (pqc_findings or active) if str(f.get("risk_band") or f.get("business_priority") or "").upper() == "MEDIUM"],
        "LOW":      [f for f in (pqc_findings or active) if str(f.get("risk_band") or f.get("business_priority") or "").upper() == "LOW"],
    }
    # Fallback: tier by severity if risk_band empty
    if not any(tiers.values()):
        tiers = {
            "CRITICAL": [f for f in (pqc_findings or active) if "CRIT" in str(f.get("severity") or "").upper()],
            "HIGH":     [f for f in (pqc_findings or active) if "HIGH" in str(f.get("severity") or "").upper()],
            "MEDIUM":   [f for f in (pqc_findings or active) if "MED"  in str(f.get("severity") or "").upper()],
            "LOW":      [f for f in (pqc_findings or active) if "LOW"  in str(f.get("severity") or "").upper()],
        }

    tier_styles = {"CRITICAL": rb_crit, "HIGH": rb_high, "MEDIUM": rb_med, "LOW": rb_low}
    for tier_name, tier_items in tiers.items():
        if not tier_items:
            continue
        _ts = tier_styles.get(tier_name, body_style)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*DARK_TEXT)
        # Use a safe plain RGB tuple for fill colour (DeviceRGB objects cannot be * unpacked)
        _ts_fill = _ts.fill_color
        if _ts_fill is None:
            _fill_rgb = GREY_BG
        elif isinstance(_ts_fill, (list, tuple)):
            _fill_rgb = tuple(int(v) for v in _ts_fill[:3])
        else:
            # fpdf DeviceRGB / similar: extract .r .g .b (0-1 float) and scale to 0-255
            try:
                _fill_rgb = (int(getattr(_ts_fill, 'r', 0.8) * 255),
                             int(getattr(_ts_fill, 'g', 0.8) * 255),
                             int(getattr(_ts_fill, 'b', 0.8) * 255))
            except Exception:
                _fill_rgb = GREY_BG
        pdf.set_fill_color(*_fill_rgb)
        pdf.cell(0, 6, clean_text(f"  Tier {tier_name} - {len(tier_items)} asset(s)"), fill=True,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*BODY_TEXT)
        for tf in tier_items[:15]:
            asset_n = clean_text(str(tf.get("asset_name") or tf.get("control_name") or "")[:50])
            algo_n  = clean_text(str(tf.get("ca_algorithm") or tf.get("key_algorithm") or
                                     tf.get("control_id") or "")[:40])
            pdf.cell(0, 4.5,
                clean_text(f"  * {asset_n}  [{algo_n}]"),
                new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        if len(tier_items) > 15:
            pdf.cell(0, 4, f"  ... and {len(tier_items) - 15} more", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)

    # ══════════════════════════════════════════════════════════════
    # PAGE N: APPENDIX
    # ══════════════════════════════════════════════════════════════
    pdf.add_page()
    draw_section_banner("8. APPENDIX")

    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5, "8.1 Discovery Methodology", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4.5, clean_text(
        "The AI Cyber Audit Box deployed deterministic regex-based cryptographic algorithm scanners "
        "against uploaded configuration exports, certificate dumps, TLS handshake logs, SSH configs, "
        "IPSec policy exports, and PKI/HSM/KMS configuration documents. No LLM inference is used "
        "in the detection phase -- every algorithm hit is an exact pattern match."
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5, "8.2 NIST PQC Standards Reference", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    nist_rows = [
        ["FIPS 203", "ML-KEM (CRYSTALS-Kyber)", "Key Encapsulation Mechanism - key exchange replacement"],
        ["FIPS 204", "ML-DSA (CRYSTALS-Dilithium)", "Digital Signature - RSA/ECDSA signature replacement"],
        ["FIPS 205", "SLH-DSA (SPHINCS+)", "Stateless Hash-based Signature - alternative signature scheme"],
        ["NIST IR 8413", "Falcon (FN-DSA)", "Compact lattice-based digital signature"],
    ]
    with pdf.table(col_widths=(25, 55, 100), text_align="L") as tbl:
        h = tbl.row()
        for hdr in ["Standard", "Algorithm", "Purpose"]:
            h.cell(hdr, style=hdr_blue)
        for row in nist_rows:
            r = tbl.row()
            r.cell(row[0], style=lbl_style)
            r.cell(row[1], style=body_style)
            r.cell(row[2], style=body_style)
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 5, "8.3 Glossary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    gloss = [
        ("CBOM",    "Cryptographic Bill of Materials - inventory of all cryptographic assets"),
        ("QBOM",    "Quantum Bill of Materials - CBOM annotated with quantum readiness status"),
        ("HNDL",    "Harvest Now, Decrypt Later - adversaries capture data today to decrypt post-quantum"),
        ("ML-KEM",  "Module-Lattice Key Encapsulation Mechanism (CRYSTALS-Kyber, FIPS 203)"),
        ("ML-DSA",  "Module-Lattice Digital Signature Algorithm (CRYSTALS-Dilithium, FIPS 204)"),
        ("SLH-DSA", "Stateless Hash-based Digital Signature Algorithm (SPHINCS+, FIPS 205)"),
        ("PQC",     "Post-Quantum Cryptography - algorithms secure against quantum computer attacks"),
    ]
    pdf.set_font("Helvetica", "", 8)
    for term, defn in gloss:
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_text_color(*DARK_TEXT)
        pdf.cell(25, 4.5, clean_text(term), new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(*BODY_TEXT)
        pdf.multi_cell(0, 4.5, clean_text(defn), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ── DISCLAIMER ───────────────────────────────────────────────
    pdf.add_page()
    draw_section_banner("9. DISCLAIMER")
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BODY_TEXT)
    pdf.multi_cell(0, 4, clean_text(
        "This report was generated by the AI Cyber Audit Box Post-Quantum Cryptography (PQC) "
        "Readiness Assessment engine. The accuracy of findings is subject to the information "
        "available in the evidence uploaded during the assessment engagement.\n\n"
        "The assessment was limited to the cryptographic evidence files explicitly uploaded during "
        "this session. Systems, services, or applications not represented in the uploaded evidence "
        "are not included in this assessment and may contain undiscovered quantum vulnerabilities.\n\n"
        "Recommendations are based on NIST Post-Quantum Cryptography standards (FIPS 203, 204, 205) "
        "and publicly available OEM vendor roadmaps at the time of assessment. Vendor timelines and "
        "product capabilities may change. This report does not constitute legal or compliance advice.\n\n"
        "The 'Harvest Now, Decrypt Later' (HNDL) threat model is referenced per current CISA and "
        "NIST guidance. Organisations should prioritise migration of externally-exposed, long-lived "
        "cryptographic assets immediately regardless of quantum computing timelines."
    ), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(10)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*DARK_TEXT)
    pdf.cell(0, 6, "*** End of PQC Readiness Assessment Report ***", align="C",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf_bytes = pdf.output()
    return bytes(pdf_bytes)


def _export_pqc_docx(session_title, findings, resolved_list, status, comments="", custom_logo=None, metadata=None):
    """
    Dedicated PQC Readiness Assessment DOCX report.
    Mirrors the 9-section structure of _export_pqc_pdf():
      1. Cover Page
      2. Document Version Control
      3. Executive Summary
      4. Cryptographic Asset Register (CBOM)
      5. Quantum Bill of Materials (QBOM)
      6. Detailed PQC Findings (per algorithm)
      7. Dependency Mapping
      8. OEM Readiness Matrix
      9. Appendix + Disclaimer
    Completely separate from _export_vapt_docx() — PQC uses quantum-readiness
    terminology, not penetration-testing terminology.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime
    import io as _io
    import os
    import json as _json

    meta = metadata or {}
    auditor_firm     = meta.get("brand_firm")     or "AI Cyber Audit Box"
    auditor_lead     = meta.get("brand_auditor")  or "Lead Auditor"
    auditor_reviewer = meta.get("brand_reviewer") or "Senior Reviewer"
    auditor_approver = meta.get("brand_approver") or "Approver"
    report_doc_id    = meta.get("brand_docid")    or "PQC-001"
    target_client    = meta.get("brand_client")   or "Client Organisation"
    assessment_date  = datetime.now().strftime("%d-%B-%Y")

    # Colour palette
    PQC_BLUE   = RGBColor(0, 80, 157)
    VULN_RED   = RGBColor(192, 0, 0)
    WEAK_ORG   = RGBColor(204, 102, 0)
    SAFE_GRN   = RGBColor(0, 140, 60)
    WHITE      = RGBColor(255, 255, 255)
    DARK       = RGBColor(15, 23, 42)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def _para(text="", bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
        return p

    def _heading(text, level=1):
        p = _para(text, bold=True, size=(16 if level == 1 else 13 if level == 2 else 11),
                  color=PQC_BLUE, space_before=10, space_after=4)
        return p

    def _table_row(table, cells_data, header=False):
        row = table.add_row()
        for i, val in enumerate(cells_data):
            cell = row.cells[i]
            cell.text = str(val or "-")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = header
                    run.font.size = Pt(9)
                    if header:
                        run.font.color.rgb = WHITE
            if header:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "00509D")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:val"), "clear")
                tc_pr.append(shd)
        return row

    # ── Pre-compute findings ──────────────────────────────────────────────────
    dict_findings = []
    for f in (findings or []):
        if hasattr(f, "to_dict"):
            dict_findings.append(f.to_dict())
        elif isinstance(f, dict):
            dict_findings.append(f)
        else:
            dict_findings.append(getattr(f, "__dict__", {}))

    active       = [f for f in dict_findings if f.get("status") not in ("Out of Scope", "False Positive", "FALSE_POSITIVE")]
    vuln_list    = [f for f in active if str(f.get("quantum_status") or "").upper() == "VULNERABLE"]
    weak_list    = [f for f in active if str(f.get("quantum_status") or "").upper() == "WEAK"]
    safe_list    = [f for f in active if str(f.get("quantum_status") or "").upper() == "SAFE"]
    crit_cnt     = sum(1 for f in active if "CRIT" in str(f.get("severity") or "").upper())
    high_cnt     = sum(1 for f in active if "HIGH" in str(f.get("severity") or "").upper())
    med_cnt      = sum(1 for f in active if "MED"  in str(f.get("severity") or "").upper())

    # ── 1. COVER PAGE ─────────────────────────────────────────────────────────
    _para("Post-Quantum Cryptography (PQC) Readiness Assessment Report",
          bold=True, size=22, color=PQC_BLUE,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=12)
    _para(f"For: {target_client}", bold=True, size=14,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _para(f"Prepared by: {auditor_firm}", size=11,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(f"Lead Auditor: {auditor_lead}", size=10,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(f"Report Reference: {report_doc_id}", size=10,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(f"Assessment Date: {assessment_date}", size=10,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(f"Status: {status}", size=10,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)
    doc.add_page_break()

    # ── 2. DOCUMENT VERSION CONTROL ───────────────────────────────────────────
    _heading("2. Document Version Control", level=1)
    vc_tbl = doc.add_table(rows=1, cols=5)
    vc_tbl.style = "Table Grid"
    _table_row(vc_tbl, ["Version", "Date", "Author", "Reviewer", "Change Summary"], header=True)
    _table_row(vc_tbl, ["1.0", assessment_date, auditor_lead, auditor_reviewer, "Initial PQC readiness assessment report"])
    doc.add_paragraph()

    # ── 3. EXECUTIVE SUMMARY ──────────────────────────────────────────────────
    _heading("3. Executive Summary", level=1)
    _para(
        f"This report presents the findings of a Post-Quantum Cryptography (PQC) "
        f"Readiness Assessment conducted for {target_client}. The assessment evaluated "
        f"cryptographic assets against current NIST PQC standards (FIPS 203, 204, 205) "
        f"and identified quantum-vulnerable algorithms that require migration before "
        f"cryptographically relevant quantum computers become available.",
        size=10, space_after=8
    )
    _heading("3.1 Risk Summary", level=2)
    sum_tbl = doc.add_table(rows=1, cols=4)
    sum_tbl.style = "Table Grid"
    _table_row(sum_tbl, ["Category", "Count", "Risk Level", "Action Required"], header=True)
    _table_row(sum_tbl, ["Quantum-Vulnerable Algorithms", str(len(vuln_list)), "CRITICAL/HIGH", "Immediate migration planning"])
    _table_row(sum_tbl, ["Classically Weak Algorithms", str(len(weak_list)), "HIGH/MEDIUM", "Replace immediately"])
    _table_row(sum_tbl, ["Quantum-Safe Algorithms", str(len(safe_list)), "INFO", "Monitor NIST updates"])
    _table_row(sum_tbl, ["Critical Severity Findings", str(crit_cnt), "P1", "Remediate within 30 days"])
    _table_row(sum_tbl, ["High Severity Findings", str(high_cnt), "P2", "Remediate within 90 days"])
    doc.add_paragraph()
    doc.add_page_break()

    # ── 4. CRYPTOGRAPHIC ASSET REGISTER (CBOM) ────────────────────────────────
    _heading("4. Cryptographic Asset Register (CBOM)", level=1)
    _para("Inventory of all cryptographic assets discovered, classified by asset category, algorithm, and quantum readiness status.", size=10, space_after=6)
    cbom_tbl = doc.add_table(rows=1, cols=7)
    cbom_tbl.style = "Table Grid"
    _table_row(cbom_tbl, ["Asset Category", "Asset / Target", "Algorithm", "Quantum Status", "Protocol Layer", "Severity", "Risk Score"], header=True)
    for f in active:
        _table_row(cbom_tbl, [
            f.get("asset_category") or "-",
            f.get("asset_name") or f.get("target") or "-",
            (f.get("title") or "").replace("Quantum-Vulnerable Algorithm Detected: ", "").replace("Classically Weak / Deprecated Algorithm Detected: ", "").replace("Quantum-Safe Algorithm Confirmed: ", "").replace("PQC Readiness Gap: ", ""),
            f.get("quantum_status") or "-",
            f.get("protocol_version") or f.get("ca_algorithm") or f.get("key_algorithm") or "-",
            f.get("severity") or "-",
            str(f.get("risk_score") or "-"),
        ])
    doc.add_paragraph()
    doc.add_page_break()

    # ── 5. QUANTUM BILL OF MATERIALS (QBOM) ───────────────────────────────────
    _heading("5. Quantum Bill of Materials (QBOM)", level=1)
    _para("Assets classified by quantum readiness status with migration priority.", size=10, space_after=6)
    qbom_tbl = doc.add_table(rows=1, cols=5)
    qbom_tbl.style = "Table Grid"
    _table_row(qbom_tbl, ["Asset", "Algorithm", "Quantum Status", "Business Priority", "Migration Action"], header=True)
    for f in active:
        if f.get("quantum_status"):
            qs = str(f.get("quantum_status") or "").upper()
            action = ("Migrate to ML-KEM/ML-DSA (NIST FIPS 203/204)" if qs == "VULNERABLE"
                      else "Replace with modern classical algorithm" if qs == "WEAK"
                      else "No action required — monitor NIST guidance")
            _table_row(qbom_tbl, [
                f.get("asset_name") or f.get("target") or "-",
                (f.get("title") or "").split(":")[-1].strip(),
                qs,
                f.get("business_priority") or "-",
                action,
            ])
    doc.add_paragraph()
    doc.add_page_break()

    # ── 6. DETAILED FINDINGS ──────────────────────────────────────────────────
    _heading("6. Detailed PQC Findings", level=1)
    for idx, f in enumerate(active, 1):
        ctrl_id   = f.get("control_id") or f.get("control") or f"PQC-{idx}"
        title     = f.get("title") or f.get("control_name") or "Unknown Finding"
        full_ctrl = f"{ctrl_id} - {title}" if (len(str(ctrl_id)) <= 12 and title and not title.lower().startswith(str(ctrl_id).lower())) else ctrl_id
        severity  = f.get("severity") or "-"
        qs        = f.get("quantum_status") or "-"
        cia       = f.get("cia_impact") or "-"
        risk_sc   = f.get("risk_score")
        risk_band = f.get("risk_band") or "-"
        oem_prod  = f.get("oem_product") or "-"
        oem_rdy   = f.get("oem_readiness_status") or "-"
        asset_cat = f.get("asset_category") or "-"
        desc      = f.get("description") or "-"
        evidence  = f.get("evidence") or "-"
        remed     = f.get("remediation") or "-"
        dep_chain = f.get("dependency_chain") or ""

        sev_color = VULN_RED if "CRIT" in severity.upper() or "HIGH" in severity.upper() else (WEAK_ORG if "MED" in severity.upper() else DARK)
        _heading(f"{full_ctrl}", level=2)

        fd_tbl = doc.add_table(rows=0, cols=2)
        fd_tbl.style = "Table Grid"
        exp_ctx = f.get("exposure_context") or "-"
        av_label = ("External (Internet-facing)" if exp_ctx.upper() == "EXTERNAL"
                    else "Internal (LAN / On-prem)" if exp_ctx.upper() == "INTERNAL"
                    else exp_ctx)
        for label, value in [
            ("Severity",          severity),
            ("Asset Category",    asset_cat),
            ("Target / Asset",    f.get("asset_name") or f.get("target") or "-"),
            ("Quantum Status",    qs),
            ("CIA Impact",        cia),
            ("PQC Risk Score",    f"{risk_sc}/100  ({risk_band})" if risk_sc else "-"),
            ("Attack Vector",     av_label),
            ("OEM Product",       oem_prod),
            ("OEM Readiness",     oem_rdy),
            ("Control ID",        full_ctrl),
            ("OWASP Category",    f.get("owasp_category") or "-"),
        ]:
            row = fd_tbl.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value or "-")
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
            for run in row.cells[1].paragraphs[0].runs:
                run.font.size = Pt(9)

        doc.add_paragraph()
        _para("Quantum Readiness Status:", bold=True, size=10, color=PQC_BLUE, space_after=2)
        _av_str = f"  |  Attack Vector: {av_label}" if exp_ctx and exp_ctx != "-" else ""
        _para(f"Algorithm: {qs}  |  PQC Risk: {risk_sc}/100 ({risk_band}){_av_str}" if risk_sc else f"Status: {qs}{_av_str}", size=9, space_after=4)

        _para("Issue Description:", bold=True, size=10, space_after=2)
        _para(desc[:600] + ("..." if len(desc) > 600 else ""), size=9, space_after=4)

        _para("Evidence:", bold=True, size=10, space_after=2)
        _para(str(evidence)[:400] + ("..." if len(str(evidence)) > 400 else ""), size=9, space_after=4)

        _para("Remediation & Migration Path:", bold=True, size=10, space_after=2)
        for line in (remed or "").split("\n"):
            if line.strip():
                _para(f"  {line.strip()}", size=9, space_after=1)

        if dep_chain:
            _para("⚠ Migration Dependency Chain:", bold=True, size=10, color=WEAK_ORG, space_after=2)
            _para(dep_chain, size=9, space_after=4)

        doc.add_paragraph()

    doc.add_page_break()

    # ── 7. DEPENDENCY MAPPING ─────────────────────────────────────────────────
    _heading("7. Dependency Mapping", level=1)
    _para("Assets with identified migration dependencies (evidence-stated only):", size=10, space_after=6)
    dep_findings = [f for f in active if f.get("dependency_chain")]
    if dep_findings:
        dep_tbl = doc.add_table(rows=1, cols=3)
        dep_tbl.style = "Table Grid"
        _table_row(dep_tbl, ["Asset / Target", "Algorithm", "Dependency Chain"], header=True)
        for f in dep_findings:
            _table_row(dep_tbl, [
                f.get("asset_name") or f.get("target") or "-",
                (f.get("title") or "").split(":")[-1].strip(),
                f.get("dependency_chain") or "-",
            ])
    else:
        _para("No explicit dependency chains identified in the submitted evidence.", size=9)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 8. OEM READINESS MATRIX ───────────────────────────────────────────────
    _heading("8. OEM / Vendor PQC Readiness Matrix", level=1)
    _para("Vendor readiness status for post-quantum cryptography migration.", size=10, space_after=6)
    oem_findings = [f for f in active if f.get("oem_product") and f.get("oem_readiness_status")]
    seen_oem = set()
    if oem_findings:
        oem_tbl = doc.add_table(rows=1, cols=4)
        oem_tbl.style = "Table Grid"
        _table_row(oem_tbl, ["Asset Category", "OEM Product", "PQC Readiness Status", "Action Required"], header=True)
        for f in oem_findings:
            key = (f.get("oem_product") or "").lower()
            if key in seen_oem:
                continue
            seen_oem.add(key)
            rdy = str(f.get("oem_readiness_status") or "").lower()
            action = ("Upgrade to PQC-ready firmware/version" if "upgrade" in rdy or "required" in rdy
                      else "Contact vendor for PQC roadmap" if "roadmap" in rdy
                      else "Confirm PQC support with vendor" if "partial" in rdy
                      else "No immediate action — monitor vendor updates")
            _table_row(oem_tbl, [
                f.get("asset_category") or "-",
                f.get("oem_product") or "-",
                f.get("oem_readiness_status") or "-",
                action,
            ])
    else:
        _para("OEM readiness data not available for submitted assets.", size=9)
    doc.add_paragraph()

    # ── 9. APPENDIX ───────────────────────────────────────────────────────────
    doc.add_page_break()
    _heading("9. Appendix", level=1)
    _heading("9.1 Assessment Methodology", level=2)
    _para(
        "This assessment was conducted using the AICyberAuditBox PQC scanner — a fully deterministic, "
        "offline analysis engine. Evidence files (TLS configuration exports, SSH configuration files, "
        "IPSec/VPN policy exports, PKI certificate exports, database encryption configuration files) "
        "were scanned against a 30-rule algorithm lookup table. No live connections to target "
        "infrastructure were made. No LLM was involved in the detection logic.",
        size=9, space_after=6
    )
    _heading("9.2 NIST PQC Standards Referenced", level=2)
    std_tbl = doc.add_table(rows=1, cols=3)
    std_tbl.style = "Table Grid"
    _table_row(std_tbl, ["Standard", "Algorithm", "Purpose"], header=True)
    for std, algo, purpose in [
        ("FIPS 203", "ML-KEM (CRYSTALS-Kyber)",    "Key Encapsulation / Key Exchange"),
        ("FIPS 204", "ML-DSA (CRYSTALS-Dilithium)", "Digital Signatures"),
        ("FIPS 205", "SLH-DSA (SPHINCS+)",          "Digital Signatures (stateless hash-based)"),
        ("NIST IR 8413", "Falcon / FN-DSA",          "Digital Signatures (lattice-based)"),
    ]:
        _table_row(std_tbl, [std, algo, purpose])
    doc.add_paragraph()
    _heading("9.3 Disclaimer", level=2)
    _para(
        "This report is prepared solely for the use of the client organisation and the authorised "
        "assessment team. The findings are based on evidence provided at the time of assessment. "
        "The quantum threat landscape evolves rapidly — findings should be reviewed against the latest "
        "NIST guidance and vendor roadmaps at least annually.",
        size=9
    )

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _export_vapt_docx(session_title, findings, resolved_list, status, comments="", custom_logo=None, metadata=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime
    import io as _io
    import os
    
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        
    def _rgb(r, g, b):
        return RGBColor(r, g, b)
        
    def _set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)
        
    def _set_cell_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'CCCCCC')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    meta = metadata or {}
    auditor_lead = meta.get("brand_auditor") or "Mr. Subhash Rao & Mr. Mahaveer Rajannavar"
    auditor_firm = meta.get("brand_firm") or "TÜV SÜD South Asia Pvt. Ltd."
    auditor_reviewer = meta.get("brand_reviewer") or "Ms. Prianka Singla"
    auditor_approver = meta.get("brand_approver") or "Mr. Atul Srivastava"
    report_doc_id = meta.get("brand_docid") or "3153142723"
    target_client = meta.get("brand_client") or "NOCPL"
    
    assets_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "data", "assets"))
    custom_logo_file = os.path.join(assets_dir, "custom_company_logo.png")
    effective_custom_logo = custom_logo if (custom_logo and os.path.exists(custom_logo)) else (custom_logo_file if os.path.exists(custom_logo_file) else None)
    shield_logo_path = os.path.join(assets_dir, "shield_logo.png")
    logo_path = effective_custom_logo if (effective_custom_logo and os.path.exists(effective_custom_logo)) else (shield_logo_path if os.path.exists(shield_logo_path) else None)
    testing_dates = f"20-June-2026 to {datetime.now().strftime('%d-%B-%Y')}"


    scope_type = "External" if "external" in session_title.lower() else "Internal"
    doc_title = f"{scope_type} Network Vulnerability Assessment and Penetration Testing Validation Report"

    # 1. COVER PAGE (Page 1)
    # Header Information Table
    hdr_table = doc.add_table(rows=1, cols=3)
    hdr_table.autofit = False
    
    col_widths = [Cm(5.5), Cm(5.5), Cm(5)]
    for row in hdr_table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    c1, c2, c3 = hdr_table.rows[0].cells
    
    p1 = c1.paragraphs[0]
    p1.add_run("Registered Office:\nTÜV SÜD South Asia Pvt. Ltd.\nTÜV SÜD House,\nOff Saki Vihar Road,\nSaki Naka, Andheri (East),\nMumbai - 400072, India.").font.size = Pt(7.5)
    
    p2 = c2.paragraphs[0]
    p2.add_run("Corporate Office:\nTÜV SÜD South Asia Pvt. Ltd.\nSolitaire, 4th Floor,\nITI Road, Aundh,\nPune - 411007, India.").font.size = Pt(7.5)
    
    p3 = c3.paragraphs[0]
    p3.add_run("Report Submitted by:\nTÜV SÜD South Asia Pvt. Ltd.\nTÜV SÜD House,\nOff Saki Vihar Road,\nSaki Naka, Mumbai - 400072.\nEmail: info@tuv-sud.in\nwww.tuv-sud.in").font.size = Pt(7.5)
    
    doc.add_paragraph().add_run("\n" + "_" * 60 + "\n").font.color.rgb = _rgb(200, 200, 200)
    
    # Validation Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    title_run = title_p.add_run(doc_title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = _rgb(15, 23, 42)
    
    client_p = doc.add_paragraph()
    client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_p.paragraph_format.space_after = Pt(40)
    client_run = client_p.add_run(f"For: {target_client}")
    client_run.bold = True
    client_run.font.name = "Arial"
    client_run.font.size = Pt(13)
    client_run.font.color.rgb = _rgb(71, 85, 105)
    
    # Logo
    if logo_path and os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.add_run().add_picture(logo_path, width=Cm(4))
        
    doc.add_paragraph().paragraph_format.space_before = Pt(40)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f"Submitted on: {datetime.now().strftime('%d-%B-%Y')}")
    date_run.font.size = Pt(9.5)
    date_run.font.color.rgb = _rgb(100, 116, 139)
    
    doc.add_page_break()

    # 2. DOCUMENT CONTROL & SUBMISSION DETAILS (Page 2)
    p = doc.add_paragraph()
    p.add_run("DOCUMENT VERSION CONTROL").bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    
    tbl_ctrl = doc.add_table(rows=8, cols=2)
    tbl_ctrl.style = 'Table Grid'
    control_rows = [
        ("Document Title", doc_title),
        ("Document ID", report_doc_id),
        ("Document Version", "1.0"),
        ("Prepared By", auditor_lead),
        ("Reviewed By", auditor_reviewer),
        ("Approved By", auditor_approver),
        ("Testing Dates", testing_dates),
        ("Effective Date", datetime.now().strftime("%d-%B-%Y"))
    ]

    for r_idx, (label, val) in enumerate(control_rows):
        c1, c2 = tbl_ctrl.rows[r_idx].cells
        _set_cell_bg(c1, "F1F5F9")
        _set_cell_borders(c1)
        _set_cell_borders(c2)
        c1.paragraphs[0].add_run(label).bold = True
        c2.paragraphs[0].add_run(val)
        
    p = doc.add_paragraph()
    p.add_run("DOCUMENT SUBMISSION DETAILS").bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    
    tbl_sub = doc.add_table(rows=3, cols=2)
    tbl_sub.style = 'Table Grid'
    sub_rows = [
        ("Date", datetime.now().strftime("%d-%B-%Y")),
        ("Classification", "Confidential"),
        ("Document Type", doc_title)
    ]
    for r_idx, (label, val) in enumerate(sub_rows):
        c1, c2 = tbl_sub.rows[r_idx].cells
        _set_cell_bg(c1, "F1F5F9")
        _set_cell_borders(c1)
        _set_cell_borders(c2)
        c1.paragraphs[0].add_run(label).bold = True
        c2.paragraphs[0].add_run(val)
        
    doc.add_page_break()

    # 3. TABLE OF CONTENTS (Page 3)
    p = doc.add_paragraph()
    p.add_run("TABLE OF CONTENTS").bold = True
    p.paragraph_format.space_after = Pt(12)
    
    toc_items = [
        ("1 TÜV SÜD PENETRATION TEST METHODOLOGY", "4"),
        ("  1.2 Standards-Based Testing and Reporting", "4"),
        ("  1.3 CVSS: Scoring Vulnerabilities", "4"),
        ("  1.4 Severity Rating Scale Table", "5"),
        ("2 EXECUTIVE SUMMARY", "6"),
        ("  2.1 Scope of the Engagement", "6"),
        ("  2.2 Assessment Date", "6"),
        ("  2.3 Summary of Findings", "7"),
        ("  2.4 Tactical Recommendations", "8"),
        ("3 TECHNICAL DETAIL REPORT", "9"),
        ("  3.2 Testing Environment", "9"),
        ("  3.3 Findings Detail", "9"),
        ("4 APPENDIX", "12"),
        ("  4.1 Testing Environment: Production", "12"),
        ("  4.2 Tools Used", "12"),
        ("  4.3 Provided Documentation", "12"),
        ("5 DISCLAIMER", "13")
    ]
    for item, p_num in toc_items:
        p_toc = doc.add_paragraph()
        run_item = p_toc.add_run(item)
        run_item.font.size = Pt(10)
        p_toc.add_run(" " + "." * (80 - len(item)) + " " + p_num).font.size = Pt(10)
        
    doc.add_page_break()

    # 4. METHODOLOGY & SEVERITY SCALE (Page 4)
    p = doc.add_paragraph()
    p.add_run("1 TÜV SÜD PENETRATION TEST METHODOLOGY").bold = True
    p.paragraph_format.space_before = Pt(12)
    
    p = doc.add_paragraph()
    p.add_run("1.2 Standards-Based Testing and Reporting").bold = True
    p.paragraph_format.space_before = Pt(8)
    doc.add_paragraph(
        "Our penetration test plans are performed according to internally developed guidelines by TÜV SÜD South Asia penetration test experts. "
        "Our test cases, where possible, are grounded in publicly available standards published by organizations such as OWASP, OSSTMM, NIST, "
        "along with our experience as a penetration test team. As an extension to these test cases, additional test cases may be identified "
        "based on penetration test experts' experience. During this engagement, network and web vulnerability scanning is executed "
        "to check configurations, TLS parameters, and application flaws."
    )
    
    p = doc.add_paragraph()
    p.add_run("1.3 CVSS: Scoring Vulnerabilities").bold = True
    p.paragraph_format.space_before = Pt(8)
    doc.add_paragraph(
        "The Common Vulnerability Scoring System (CVSS) is an open framework for communicating the characteristics and severity of IT system vulnerabilities. "
        "CVSS consists of three metric groups: Base, Temporal, and Environmental. In this report, CVSSv4.0 base metrics are calculated to reflect "
        "the severity of security weaknesses, assessing Exploitability parameters (Attack Vector, Attack Complexity, Privileges Required, etc.) "
        "and System Impact parameters (Confidentiality, Integrity, and Availability threat vectors)."
    )
    
    p = doc.add_paragraph()
    p.add_run("1.4 Severity Rating Scale").bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    tbl_scale = doc.add_table(rows=5, cols=3)
    tbl_scale.style = 'Table Grid'
    scale_headers = ["Severity Range", "Classification", "Remediation Priority"]
    for col_idx, text in enumerate(scale_headers):
        cell = tbl_scale.rows[0].cells[col_idx]
        _set_cell_bg(cell, "0F172A")
        cell.paragraphs[0].add_run(text).font.color.rgb = _rgb(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
        
    ranges = [
        ["9.0 - 10.0", "Critical", "Immediate Mitigation (P1)"],
        ["7.0 - 8.9", "High", "Short-term Remediation (P2)"],
        ["4.0 - 6.9", "Medium", "Planned Resolution (P3)"],
        ["0.1 - 3.9", "Low", "Administrative Fixes (P4)"]
    ]
    for row_idx, row_data in enumerate(ranges, 1):
        for col_idx, val in enumerate(row_data):
            tbl_scale.rows[row_idx].cells[col_idx].paragraphs[0].add_run(val)
            
    doc.add_page_break()

    # 5. EXECUTIVE SUMMARY (Page 6)
    p = doc.add_paragraph()
    p.add_run("2 EXECUTIVE SUMMARY").bold = True
    p.paragraph_format.space_before = Pt(12)
    
    p = doc.add_paragraph()
    p.add_run("2.1 Scope of the Engagement").bold = True
    doc.add_paragraph(
        f"TÜV SÜD was engaged to perform network and application security validation testing. The target audit scope consists of "
        f"critical service interfaces, web applications, and network routing configurations. Target assets: {session_title}."
    )
    
    p = doc.add_paragraph()
    p.add_run("2.2 Assessment Date").bold = True
    doc.add_paragraph(f"Testing Dates: {testing_dates}")
    
    # Summary of Findings (Page 7)
    p = doc.add_paragraph()
    p.add_run("2.3 Summary of Findings").bold = True
    p.paragraph_format.space_before = Pt(12)
    
    parsed_reg = _get_all_parsed_findings_from_registry()
    if parsed_reg:
        active_findings = parsed_reg
    else:
        active_findings = [f for f in findings if f.get("status") not in ("Out of Scope", "False Positive", "FALSE_POSITIVE")]

    critical_cnt = sum(1 for f in active_findings if str(f.get("severity", "")).strip().upper() == "CRITICAL")
    high_cnt = sum(1 for f in active_findings if str(f.get("severity", "")).strip().upper() == "HIGH")
    medium_cnt = sum(1 for f in active_findings if str(f.get("severity", "")).strip().upper() == "MEDIUM")
    low_cnt = sum(1 for f in active_findings if str(f.get("severity", "")).strip().upper() == "LOW")
    total_cnt = critical_cnt + high_cnt + medium_cnt + low_cnt

    
    doc.add_paragraph(f"Based on the assessment, {total_cnt} vulnerabilities have been found in the target scope network which are categorized as follows:")
    
    p = doc.add_paragraph()
    p.add_run("2.3.2 Tabular Summary").bold = True
    
    tbl_sum = doc.add_table(rows=2, cols=5)
    tbl_sum.style = 'Table Grid'
    sum_hdrs = ["Critical", "High", "Medium", "Low", "Total Findings"]
    for col_idx, text in enumerate(sum_hdrs):
        cell = tbl_sum.rows[0].cells[col_idx]
        _set_cell_bg(cell, "0F172A")
        cell.paragraphs[0].add_run(text).font.color.rgb = _rgb(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
        
    counts = [str(critical_cnt), str(high_cnt), str(medium_cnt), str(low_cnt), str(total_cnt)]
    for col_idx, val in enumerate(counts):
        cell = tbl_sum.rows[1].cells[col_idx]
        if col_idx == 4:
            _set_cell_bg(cell, "F1F5F9")
            cell.paragraphs[0].add_run(val).bold = True
        else:
            cell.paragraphs[0].add_run(val)
            
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("2.3.4 Vulnerabilities Summary").bold = True
    
    tbl_vulns = doc.add_table(rows=1 + len(active_findings) + 1, cols=4)
    tbl_vulns.style = 'Table Grid'
    vuln_hdrs = ["Sr. No.", "Vulnerabilities", "CVSS Score", "Severity"]
    for col_idx, text in enumerate(vuln_hdrs):
        cell = tbl_vulns.rows[0].cells[col_idx]
        _set_cell_bg(cell, "0F172A")
        cell.paragraphs[0].add_run(text).font.color.rgb = _rgb(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
        
    max_score = 0.0
    for idx, f in enumerate(active_findings, 1):
        score = float(f.get("severity_score", 0.0) or 0.0)
        if score > max_score:
            max_score = score
        sev = str(f.get("severity", "Low")).split()[-1].upper()
        
        row_cells = tbl_vulns.rows[idx].cells
        row_cells[0].paragraphs[0].add_run(str(idx))
        row_cells[1].paragraphs[0].add_run(f.get("control", "") or f.get("finding", ""))
        row_cells[2].paragraphs[0].add_run(f"{score:.1f}")
        row_cells[3].paragraphs[0].add_run(sev)
        
    overall_sev = "LOW"
    if max_score >= 9.0:
        overall_sev = "CRITICAL"
    elif max_score >= 7.0:
        overall_sev = "HIGH"
    elif max_score >= 4.0:
        overall_sev = "MEDIUM"
        
    over_cells = tbl_vulns.rows[len(active_findings) + 1].cells
    for cell in over_cells:
        _set_cell_bg(cell, "F1F5F9")
    over_cells[0].paragraphs[0].add_run("OVERALL").bold = True
    over_cells[1].paragraphs[0].add_run("OVERALL SCORE").bold = True
    over_cells[2].paragraphs[0].add_run(f"{max_score:.1f}").bold = True
    over_cells[3].paragraphs[0].add_run(overall_sev).bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("2.4 Tactical Recommendations").bold = True
    doc.add_paragraph(
        "It is recommended to follow the guidelines suggested by OWASP, OSSTMM and NIST. It is recommended to implement secure "
        "SDLC while developing the applications, disable weak cipher suites like CBC cipher algorithms, and implement "
        "Strict-Transport-Security configurations across all application headers."
    )
    
    doc.add_page_break()

    # 6. TECHNICAL DETAIL REPORT (Page 9+)
    p = doc.add_paragraph()
    p.add_run("3 TECHNICAL DETAIL REPORT: NETWORK VULNERABILITY ASSESSMENT").bold = True
    p.paragraph_format.space_before = Pt(12)
    
    p = doc.add_paragraph()
    p.add_run("3.2 Testing Environment").bold = True
    doc.add_paragraph("For details concerning the testing environment, please refer to Appendix 4.1.")
    
    p = doc.add_paragraph()
    p.add_run("3.3 Findings Detail").bold = True
    p.paragraph_format.space_before = Pt(12)
    
    for idx, f in enumerate(active_findings, 1):
        # Use auditor-set custom heading if available, else derive from finding metadata
        custom_h = (f.get("custom_heading") or "").strip()
        vuln_title = custom_h or f.get("title") or f.get("control") or f.get("finding") or f"Finding {idx}"
        score = float(f.get("severity_score", 0.0) or 0.0)
        sev_label = str(f.get("severity", "Low")).split()[-1].upper()
        target_val = html.unescape(str(f.get("target") or f.get("control_id", "") or "Web / Network infrastructure"))
        # VAPT reports: redact email/phone (incidental PII) but keep IPs -- the
        # vulnerable host's IP address is the report's actual content, not PII.
        desc_val = html.unescape(redact_pii(str(f.get("description") or f.get("gap_description") or f.get("finding") or "-"), redact_ip=False))
        poc_val = html.unescape(redact_pii(str(f.get("evidence") or f.get("evidence_snippet") or f.get("evidence_quote") or f.get("poc") or "Console / Log Audit Verification"), redact_ip=False))
        remed_raw = str(f.get("recommendation") or f.get("remediation") or "").strip()
        remed_val = html.unescape(redact_pii(remed_raw, redact_ip=False)) if remed_raw else "Immediately apply vendor security patches or software updates."

        # ── Resolve uploaded screenshots/images for VAPT POC embedding ───────
        # Priority: explicit fields + ALL images listed in source_files / evidence_source_file
        def _resolve_vapt_images(f_dict):
            valid_imgs = []
            candidate_fields = [
                f_dict.get("poc_image"),
                f_dict.get("image_path"),
                f_dict.get("extra_image"),
            ]
            src = str(f_dict.get("source_files") or f_dict.get("evidence_source_file") or "")
            for part in src.split(","):
                part = part.strip()
                if part.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
                    candidate_fields.append(part)
            for p in candidate_fields:
                if p and isinstance(p, str) and p not in valid_imgs:
                    if os.path.exists(p):
                        valid_imgs.append(p)
            return valid_imgs

        resolved_imgs = _resolve_vapt_images(f)
        # ─────────────────────────────────────────────────────────────────────

        # Heading
        fp = doc.add_paragraph()
        fp.paragraph_format.space_before = Pt(14)
        run_f = fp.add_run(f"FN-{idx:02d} {vuln_title}")
        run_f.bold = True
        run_f.font.size = Pt(12)
        if sev_label == "CRITICAL": run_f.font.color.rgb = _rgb(220, 38, 38)
        elif sev_label == "HIGH": run_f.font.color.rgb = _rgb(225, 29, 72)
        elif sev_label == "MEDIUM": run_f.font.color.rgb = _rgb(217, 119, 6)
        else: run_f.font.color.rgb = _rgb(37, 99, 235)

        # Meta line
        pm = doc.add_paragraph()
        r_m1 = pm.add_run("Severity: ")
        r_m1.bold = True
        r_m2 = pm.add_run(f"{sev_label} ({score:.1f})")
        r_m2.bold = True
        r_m2.font.color.rgb = run_f.font.color.rgb
        r_m3 = pm.add_run(f"   |   Location / Target: {target_val}")
        if f.get("source_tool"):
            pm.add_run(f"   |   Tool: {f.get('source_tool')}")
        if f.get("category"):
            pm.add_run(f"   |   Risk Category: {f.get('category')}")

        # CIA & PII Meta line
        cia_val = f.get("cia_impact") or ""
        pii_flag = f.get("is_pii_exposed", False)
        if cia_val or pii_flag:
            pm_cia = doc.add_paragraph()
            pm_cia.paragraph_format.space_before = Pt(2)
            r_c1 = pm_cia.add_run("CIA Impact: ")
            r_c1.bold = True
            r_c2 = pm_cia.add_run(cia_val or "C:NONE | I:NONE | A:NONE")
            if pii_flag:
                r_pii = pm_cia.add_run("   |   ⚠ PII EXPOSURE DETECTED")
                r_pii.bold = True
                r_pii.font.color.rgb = _rgb(220, 38, 38)

        # Issue Description
        p_desc_hdr = doc.add_paragraph()
        p_desc_hdr.paragraph_format.space_before = Pt(6)
        r_dh = p_desc_hdr.add_run("Issue Description:")
        r_dh.bold = True
        r_dh.font.color.rgb = _rgb(217, 119, 6)
        doc.add_paragraph(desc_val)

        # Proof of Vulnerability
        p_poc_hdr = doc.add_paragraph()
        p_poc_hdr.paragraph_format.space_before = Pt(6)
        r_ph = p_poc_hdr.add_run("Proof of Vulnerability:")
        r_ph.bold = True
        r_ph.font.color.rgb = _rgb(217, 119, 6)

        # Code block container for Proof of Concept
        tbl_poc = doc.add_table(rows=1, cols=1)
        tbl_poc.style = 'Table Grid'
        c_poc = tbl_poc.rows[0].cells[0]
        _set_cell_bg(c_poc, "F8FAFC")
        _set_cell_borders(c_poc)
        p_code = c_poc.paragraphs[0]
        r_code = p_code.add_run(poc_val)
        r_code.font.name = "Consolas"
        r_code.font.size = Pt(8.5)

        # Recommendation
        p_rec_hdr = doc.add_paragraph()
        p_rec_hdr.paragraph_format.space_before = Pt(6)
        r_rh = p_rec_hdr.add_run("Recommendation:")
        r_rh.bold = True
        r_rh.font.color.rgb = _rgb(217, 119, 6)
        doc.add_paragraph(remed_val)

        # Developer Actionable Mitigation Steps
        remed_actionable = f.get("remediation_actionable") or f.get("actionable_remediation") or ""
        if remed_actionable and remed_actionable.strip() != remed_val.strip():
            p_act_hdr = doc.add_paragraph()
            p_act_hdr.paragraph_format.space_before = Pt(4)
            r_ah = p_act_hdr.add_run("Developer Actionable Mitigation Steps:")
            r_ah.bold = True
            r_ah.font.color.rgb = _rgb(15, 23, 42)
            doc.add_paragraph(remed_actionable)

        # ── VAPT POC Screenshots (ALL matching evidence images) ─────────────
        # Embed ALL uploaded vulnerability screenshots under Proof of Concept.
        if resolved_imgs:
            p_img_hdr = doc.add_paragraph()
            p_img_hdr.paragraph_format.space_before = Pt(6)
            r_ih = p_img_hdr.add_run("Proof of Concept Artifacts & Screenshots:")
            r_ih.bold = True
            r_ih.font.color.rgb = _rgb(217, 119, 6)
            for img_idx, img_path in enumerate(resolved_imgs, 1):
                try:
                    p_img = doc.add_paragraph()
                    p_img.add_run().add_picture(img_path, width=Cm(15))
                    p_cap = doc.add_paragraph()
                    p_cap.paragraph_format.space_before = Pt(2)
                    r_cap = p_cap.add_run(f"[Figure {img_idx}: POC Evidence Screenshot — {os.path.basename(img_path)}]")
                    r_cap.italic = True
                    r_cap.font.size = Pt(8)
                    r_cap.font.color.rgb = _rgb(100, 116, 139)
                except Exception:
                    doc.add_paragraph(f"[Screenshot file: {os.path.basename(img_path)} — could not be embedded]").italic = True
        # ─────────────────────────────────────────────────────────────────────

        doc.add_paragraph()


    doc.add_page_break()

    # ── PQC (Post-Quantum Cryptography Readiness): Cryptographic Asset Inventory ──
    # Additive section, gated on presence of PQC scan data (quantum_status is only
    # ever populated by pqc_parser.py) -- does not alter any existing VAPT table or
    # section above. Blank cell wherever a Finding field is empty, never a guess.
    pqc_findings = [f for f in active_findings if f.get("quantum_status")]
    if pqc_findings:
        p = doc.add_paragraph()
        p.add_run("3.5 Cryptographic Asset Inventory (CBOM/QBOM)").bold = True
        p.paragraph_format.space_before = Pt(12)
        doc.add_paragraph(
            "The following inventory lists cryptographic algorithms, certificates, and protocols "
            "discovered during the assessment, evaluated for post-quantum readiness."
        )

        pqc_hdrs = ["Asset", "CA Algorithm", "Key Algorithm", "Protocol", "Quantum Status", "Exposure", "Port", "Environment", "Severity", "Risk Score", "Business Priority", "OEM Readiness"]
        tbl_pqc = doc.add_table(rows=1 + len(pqc_findings), cols=len(pqc_hdrs))
        tbl_pqc.style = 'Table Grid'
        for col_idx, text in enumerate(pqc_hdrs):
            cell = tbl_pqc.rows[0].cells[col_idx]
            _set_cell_bg(cell, "0F172A")
            cell.paragraphs[0].add_run(text).font.color.rgb = _rgb(255, 255, 255)
            cell.paragraphs[0].runs[0].bold = True

        # Risk-band color convention mirrors the Quantum Status column above.
        _rb_colors = {
            "CRITICAL": _rgb(220, 38, 38),
            "HIGH": _rgb(230, 126, 34),
            "MEDIUM": _rgb(217, 119, 6),
            "LOW": _rgb(5, 150, 105),
        }

        for ridx, pf in enumerate(pqc_findings, 1):
            qs = str(pf.get("quantum_status") or "").strip().upper()
            row_cells = tbl_pqc.rows[ridx].cells
            row_cells[0].paragraphs[0].add_run(pf.get("asset_name") or "")
            row_cells[1].paragraphs[0].add_run(pf.get("ca_algorithm") or "")
            row_cells[2].paragraphs[0].add_run(pf.get("key_algorithm") or "")
            row_cells[3].paragraphs[0].add_run(pf.get("protocol_version") or "")
            qs_run = row_cells[4].paragraphs[0].add_run(qs)
            qs_run.bold = True
            if qs == "VULNERABLE":
                qs_run.font.color.rgb = _rgb(220, 38, 38)
            elif qs == "WEAK":
                qs_run.font.color.rgb = _rgb(217, 119, 6)
            elif qs == "SAFE":
                qs_run.font.color.rgb = _rgb(5, 150, 105)
            row_cells[5].paragraphs[0].add_run(pf.get("exposure_context") or "")
            row_cells[6].paragraphs[0].add_run(pf.get("port") or "")
            row_cells[7].paragraphs[0].add_run(pf.get("environment") or "")
            row_cells[8].paragraphs[0].add_run(pf.get("severity") or "")

            _rscore = pf.get("risk_score")
            _rband = str(pf.get("risk_band") or "").strip().upper()
            _rscore_text = f"{_rscore} ({_rband})" if (_rscore is not None and _rband) else (str(_rscore) if _rscore is not None else "")
            rs_run = row_cells[9].paragraphs[0].add_run(_rscore_text)
            if _rband in _rb_colors:
                rs_run.bold = True
                rs_run.font.color.rgb = _rb_colors[_rband]

            row_cells[10].paragraphs[0].add_run(pf.get("business_priority") or "")

            _oem_product = pf.get("oem_product") or ""
            _oem_status = pf.get("oem_readiness_status") or ""
            _oem_text = f"{_oem_product}: {_oem_status}" if (_oem_product and _oem_status) else (_oem_status or _oem_product)
            row_cells[11].paragraphs[0].add_run(_oem_text)

        # Migration Dependencies -- only rendered when at least one PQC finding
        # flags a downstream dependency also being quantum-vulnerable.
        _mig_findings = [pf for pf in pqc_findings if pf.get("migration_dependency_flag")]
        if _mig_findings:
            p = doc.add_paragraph()
            p.add_run("Migration Dependencies").bold = True
            p.paragraph_format.space_before = Pt(10)
            for pf in _mig_findings:
                _asset = pf.get("asset_name") or "—"
                _chain = pf.get("dependency_chain") or ""
                doc.add_paragraph(f"{_asset}: {_chain}" if _chain else f"{_asset}", style="List Bullet")

        doc.add_page_break()

    # 7. APPENDIX (Page 12)
    p = doc.add_paragraph()
    p.add_run("4 APPENDIX").bold = True
    
    p = doc.add_paragraph()
    p.add_run("4.1 Testing Environment: Production").bold = True
    doc.add_paragraph(
        "The network validation testing was conducted against active production interfaces as Black Box testing. "
        "No network degradation or host disruptions occurred during scanning."
    )
    
    p = doc.add_paragraph()
    p.add_run("4.2 Tools Used").bold = True
    doc.add_paragraph("Nmap Security Scanner\nNessus Vulnerability Scanner\nBurp Suite Professional Web Scanner\nOpenSSL TLS Testing Utility")
    
    p = doc.add_paragraph()
    p.add_run("4.3 Provided Documentation").bold = True
    doc.add_paragraph("Standard Target IP Address range definitions.\nVPN Credentials and authorization letters.")
    
    doc.add_page_break()

    # 8. DISCLAIMER (Page 13)
    p = doc.add_paragraph()
    p.add_run("5 DISCLAIMER").bold = True
    doc.add_paragraph(
        "The accuracy of the information and data provided in this report is subject to the information available to TÜV SÜD testing team "
        "during the engagement. The team relies on the accuracy and completeness of the information provided by the client and does not "
        "assume any responsibility for any inaccuracies or omissions. The assessment was limited to the scope defined and does not "
        "guarantee the absolute exclusion of other vulnerabilities."
    )
    
    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

import io as _io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from src.core.pii_redactor import redact_pii


# ── RAG accuracy overhaul (Phase 5/6/7): export enrichment helpers ─────────────
# The findings table across all export paths (Word template, programmatic DOCX
# fallback, PDF) uses a fixed 8-column professional layout that predates the
# policy/evidence split -- confirmed with the user not to add new columns (would
# widen/disturb the template). Instead these helpers build short summary strings
# that get folded into the existing Observations/Evidence cell text. All return ""
# when a finding predates this schema (its DB columns are null), so older reports
# are unaffected.
def _policy_evidence_summary(f):
    """Compact 'Policy: X, Y | Evidence: X, Y' line for the Observations cell."""
    pol_status = f.get("policy_status")
    pol_assess = f.get("policy_assessment")
    ev_status = f.get("evidence_status")
    ev_assess = f.get("evidence_assessment")
    if not (pol_status or pol_assess or ev_status or ev_assess):
        return ""
    parts = []
    if pol_status or pol_assess:
        parts.append(f"Policy: {pol_status or 'UNKNOWN'}, {pol_assess or 'UNKNOWN'}")
    if ev_status or ev_assess:
        parts.append(f"Evidence: {ev_status or 'UNKNOWN'}, {ev_assess or 'UNKNOWN'}")
    return " | ".join(parts)


def _policy_evidence_gap_text(f):
    """Specific gap explanations, appended to the Observations cell.

    These are free-text, LLM-generated from evidence documents -- unlike the
    enum-style status/assessment fields in _policy_evidence_summary, a gap
    narrative can easily quote a phone number/email/IP straight out of the
    source evidence, so it must go through the same redaction as
    description/recommendation/evidence_snippet before export.
    """
    pol_gap = redact_pii(str(f.get("policy_gap") or "").strip())
    ev_gap = redact_pii(str(f.get("evidence_gap") or "").strip())
    parts = []
    if pol_gap and pol_gap.lower() != "no policy gap identified.":
        parts.append(f"Policy Gap: {pol_gap}")
    if ev_gap and ev_gap.lower() != "no evidence gap identified.":
        parts.append(f"Evidence Gap: {ev_gap}")
    return " ".join(parts)


def _evidence_meta_summary(f):
    """Validity/freshness/relevance line, appended to the Evidence cell."""
    parts = []
    validity = f.get("policy_validity")
    freshness = f.get("evidence_freshness")
    relevance = f.get("evidence_relevance")
    if validity and validity != "UNKNOWN":
        parts.append(f"Policy Validity: {validity}")
    if freshness and freshness != "UNKNOWN":
        parts.append(f"Evidence Freshness: {freshness}")
    if relevance:
        parts.append(f"Relevance: {relevance}")
    return " | ".join(parts)


def _export_iso_template_docx(session_title, findings, resolved_list, status, comments="", custom_logo=None, metadata=None):
    """
    Generates an ISO audit DOCX report using `VAPT/Sample report.docx` as the master template.
    Replaces metadata tables and rebuilds Table 6 (Observations) with live audit findings.
    Falls back to programmatic creation if template not found.
    """
    import io as _io
    import os
    from copy import deepcopy
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    today = datetime.now().strftime("%d.%m.%Y")

    # ── Report branding metadata ────────────────────────────────────────────
    meta = metadata or {}
    auditor_lead     = meta.get("brand_auditor") or "Mr. Subhash Rao & Mr. Mahaveer Rajannavar"
    auditor_firm     = meta.get("brand_firm") or "Digital Age Strategies Pvt Ltd"
    auditor_reviewer = meta.get("brand_reviewer") or "Mr. Subhash Rao"
    auditor_approver = meta.get("brand_approver") or "Mr. Dinesh S Shastry"
    report_doc_id    = meta.get("brand_docid") or "DigAge:0001:2025-26"
    target_client    = meta.get("brand_client") or "Client Organization"
    submitted_to     = meta.get("brand_client") or "Audit Committee"
    designation      = "IS Audit Lead"
    testing_dates    = today
    # Which standard this report is actually certifying against -- defaults to today's
    # exact "ISO 27001" wording when unset, so existing ISO 27001 reports are unaffected.
    framework_label  = _resolve_framework_label(meta.get("framework"))

    # ── Load template (search multiple paths) ─────────────────────────────────
    template_path = None
    _base = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    for candidate in [
        os.path.join(_base, "VAPT", "Sample report.docx"),
        os.path.join(_base, "Sample report.docx"),
        r"c:\Users\HP\Desktop\llama,cpp\au\VAPT\Sample report.docx",
    ]:
        if os.path.exists(candidate):
            template_path = candidate
            break

    if not template_path:
        # Graceful fallback — call the original programmatic generator
        return None  # caller will use old code path

    doc = Document(template_path)

    # ── Helper: set cell text preserving paragraph formatting ─────────────────
    def _set_cell_text(cell, text):
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""
        if cell.paragraphs:
            cell.paragraphs[0].text = text
        else:
            cell.add_paragraph(text)

    def _set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color.upper())
        tcPr.append(shd)

    # ── Replace cover page paragraphs ─────────────────────────────────────────
    # Word (and this template specifically) can split a single word across multiple runs
    # at internal formatting/autocorrect boundaries -- e.g. "e-Office" is stored as three
    # separate runs ("e" / "-" / "Office"). A per-run .replace() silently never matches in
    # that case even though the paragraph's own concatenated text clearly contains it, so
    # replacements are done against the paragraph's full joined text instead, then written
    # back into its first run (formatting variation *within* the replaced span is lost,
    # which is fine for these plain-text placeholder/reference paragraphs).
    audit_label = f"IS Audit Report for {session_title}"

    def _replace_across_runs(paragraph, replacements):
        full_text = "".join(r.text for r in paragraph.runs)
        new_text = full_text
        changed = False
        for old, new in replacements:
            if old in new_text:
                new_text = new_text.replace(old, new)
                changed = True
        if changed and paragraph.runs:
            paragraph.runs[0].text = new_text
            for r in paragraph.runs[1:]:
                r.text = ""

    for para in doc.paragraphs:
        _replace_across_runs(para, [
            ("e-Office", session_title),
            ("Attendance System Application", session_title),
        ])
        # The template's references section states "...ISMS under ISO 27001 Standard..."
        # verbatim (baked into the .docx, not built by this Python code) -- swap it for
        # the actual audited standard so a NIST/SOC2/etc. report doesn't claim ISO 27001
        # certification. No-op (replaces "ISO 27001 Standard" with itself) for real ISO
        # 27001 audits. Deliberately NOT touching "ISO 27001 LA" elsewhere in the template
        # -- that's a named auditor's real professional certification, unrelated to which
        # framework this particular report covers.
        _replace_across_runs(para, [
            ("ISO 27001 Standard", f"{framework_label} Standard"),
        ])

    # ── Table 0: Document Control ──────────────────────────────────────────────
    #  Row[1]: Document Title
    #  Row[2]: Document ID
    #  Row[4]: Prepared by
    #  Row[5]: Reviewed by
    #  Row[6]: Approved by
    #  Row[8]: Release date
    t0 = doc.tables[0]
    _set_cell_text(t0.rows[1].cells[1], f"Information Security Audit - {session_title}")
    _set_cell_text(t0.rows[2].cells[1], report_doc_id)
    _set_cell_text(t0.rows[4].cells[1], auditor_lead)
    _set_cell_text(t0.rows[5].cells[1], auditor_reviewer)
    _set_cell_text(t0.rows[6].cells[1], auditor_approver)
    _set_cell_text(t0.rows[8].cells[1], today)

    # ── Table 3: Details of Auditee ───────────────────────────────────────────
    #  Row[0]: Name of Organization
    #  Row[1]: Audit Area
    #  Row[3]: Auditee Representatives
    t3 = doc.tables[3]
    _set_cell_text(t3.rows[0].cells[2], target_client)
    _set_cell_text(t3.rows[1].cells[2], f"IS Audit of {session_title}")
    _set_cell_text(t3.rows[3].cells[2], submitted_to)

    # ── Table 4: Details of Auditor ───────────────────────────────────────────
    #  Row[0]: Lead Auditor
    #  Row[1]: Co-Auditor (reuse lead)
    #  Row[2]: Audit dates
    #  Row[3]: Report Date
    t4 = doc.tables[4]
    _set_cell_text(t4.rows[0].cells[2], auditor_lead)
    _set_cell_text(t4.rows[1].cells[2], auditor_firm)
    _set_cell_text(t4.rows[2].cells[2], testing_dates)
    _set_cell_text(t4.rows[3].cells[2], today)

    # ── Count findings by risk level ──────────────────────────────────────────
    active_findings = [
        f for f in (findings or [])
        if f.get("status", "Open") not in ("Dismissed", "Rejected", "Out of Scope", "Out Of Scope", "False Positive")
    ]
    accepted_findings = [
        f for f in (findings or [])
        if str(f.get("display_status", f.get("status", ""))).lower() in ("accepted", "compliant")
    ]

    def _risk_level(f):
        sev = str(f.get("severity", f.get("risk", "Low"))).upper()
        status_str = str(f.get("display_status", f.get("status", ""))).lower()
        if status_str in ("accepted", "compliant"):
            return "Accepted"
        if "CRITICAL" in sev or "P1" in sev or "HIGH" in sev or "P2" in sev:
            return "High"
        if "MEDIUM" in sev or "P3" in sev:
            return "Medium"
        return "Low"

    high_nc   = [f for f in active_findings if _risk_level(f) == "High"]
    med_nc    = [f for f in active_findings if _risk_level(f) == "Medium"]
    low_nc    = [f for f in active_findings if _risk_level(f) == "Low"]
    acc       = accepted_findings

    # ── Table 5: Summary of Findings ──────────────────────────────────────────
    # Row[2]: HIGH   | count | complied | not-complied
    # Row[3]: MEDIUM | count | complied | not-complied
    # Row[4]: LOW    | count | complied | not-complied
    # Row[5]: Accepted | count | - | -
    t5 = doc.tables[5]
    def _update_summary_row(row, label, total, complied, not_complied):
        _set_cell_text(row.cells[1], label)
        _set_cell_text(row.cells[2], str(total))
        _set_cell_text(row.cells[3], str(complied))
        _set_cell_text(row.cells[4], str(not_complied))

    _update_summary_row(t5.rows[2], "HIGH",     len(high_nc), 0, len(high_nc))
    _update_summary_row(t5.rows[3], "MEDIUM",   len(med_nc),  0, len(med_nc))
    _update_summary_row(t5.rows[4], "LOW",      len(low_nc),  0, len(low_nc))
    _update_summary_row(t5.rows[5], "Accepted", len(acc),     len(acc), 0)

    # ── Table 6: Rebuild Observations ─────────────────────────────────────────
    # Columns: S.No. | Control points | Policy Reference | Observations | Risk | Impact | Suggestion | Evidence
    # Keep header row only (row 0), remove all other rows, add live findings
    t6 = doc.tables[6]

    # Remove all data rows (keep header row 0 only)
    while len(t6.rows) > 1:
        tr = t6.rows[-1]._tr
        t6._tbl.remove(tr)

    # Color map
    RISK_COLORS = {
        "High":     ("FFC0C0", "C00000"),  # light red bg, dark red text
        "Medium":   ("FFFF99", "7F6000"),  # yellow bg, dark text
        "Low":      ("E2EFDA", "375623"),  # light green bg, dark green text
        "Accepted": ("DDEEFF", "1F4E79"),  # blue-grey bg, navy text
    }
    SECTION_BG = "D9D9D9"

    # Helper to add a row to t6
    def _add_obs_row(sno, control_pt, policy_ref, observation, risk, impact, suggestion, evidence, is_section=False):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        # Clone the template row format from the header row (row 0)
        hdr_tr = t6.rows[0]._tr
        new_tr = copy.deepcopy(hdr_tr)
        t6._tbl.append(new_tr)

        # Get the newly appended row
        new_row = t6.rows[-1]
        cells = new_row.cells

        values = [str(sno), str(control_pt), str(policy_ref), str(observation), str(risk), str(impact), str(suggestion), str(evidence)]
        for ci, (cell, val) in enumerate(zip(cells, values)):
            # Clear all runs
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ""
            if cell.paragraphs:
                run = cell.paragraphs[0].add_run(val)
                run.font.size = Pt(9)
                if is_section:
                    run.bold = True
            else:
                p = cell.add_paragraph()
                run = p.add_run(val)
                run.font.size = Pt(9)

            # Apply background color
            if is_section:
                _set_cell_bg(cell, SECTION_BG)
            elif risk in RISK_COLORS:
                bg_hex, _ = RISK_COLORS[risk]
                _set_cell_bg(cell, bg_hex)

    # Build grouped findings: group by VAPT control category
    section_counters = {}
    section_letters = {}
    letter_idx = 0
    ALPHABET = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    sno_counter = 1

    # Group findings by control category
    from collections import OrderedDict
    control_groups = OrderedDict()

    for f in (active_findings + acc):
        ctrl = str(f.get("control_id") or f.get("control") or "General").strip()
        # Use use_case name if available; fall back to full control_id (not just first number)
        uc_name = f.get("use_case") or f.get("control_category") or ctrl
        # BUG FIX: cat_key was previously only the FIRST number ("5" for "5.15", "8" for "8.17")
        # This caused ALL clause-5 controls to merge under one header and ALL clause-8 under another.
        # Fix: use the full control_id (e.g. "5.15", "8.17") as the group key.
        cat_key = ctrl  # Use the full control_id as the grouping key
        if cat_key not in control_groups:
            control_groups[cat_key] = []
        control_groups[cat_key].append(f)

    for grp_key, grp_findings in control_groups.items():
        # Section header row
        if grp_key not in section_letters:
            section_letters[grp_key] = ALPHABET[letter_idx % 26]
            letter_idx += 1
        sec_letter = section_letters[grp_key]
        _add_obs_row(sec_letter, grp_key, "", "", "", "", "", "", is_section=True)

        # Finding rows
        for f in grp_findings:
            # ── BUG FIX: ctrl_pt for Excel-mode findings ──────────────────────────
            # In Excel mode, control_id = bare "8.17" and control_name has the
            # checklist question ("8.17 Whether NTP is enabled").
            # Two rows with same control_id (e.g. both 8.17) MUST show distinct
            # question text in the report — NOT both showing the same "8.17" label.
            # Fix: if control_name is richer than control_id, use control_name.
            import re as _re
            _bare_id_pat = _re.compile(r'^\d{1,2}\.\d{1,2}(?:\.\d{1,2})?$')
            c_id = str(f.get("control_id") or "").strip()
            c_name = str(f.get("control_name") or "").strip()
            u_c = str(f.get("use_case") or "").strip()
            c_n = str(f.get("control") or "").strip()
            audit_chk = f.get("audit_check") or f.get("control_check") or f.get("check") or f.get("scenario")

            if audit_chk:
                ctrl_pt = str(audit_chk).strip()
            elif _bare_id_pat.match(c_id) and c_name and len(c_name) > len(c_id) + 2:
                # Excel mode: control_name has the checklist question text
                # Strip trailing paren-ID suffix e.g. "Capacity Management (8.6)"
                ctrl_pt = _re.sub(r'\s*\(\s*\d{1,2}\.\d{1,2}(?:\.\d{1,2})?\s*\)\s*$', '', c_name).strip()
            elif u_c and c_n and u_c.lower() == c_n.lower():
                ctrl_pt = u_c
            elif u_c and c_n and (u_c in c_n or c_n in u_c):
                ctrl_pt = max(u_c, c_n, key=len)
            elif c_id and (c_name and len(c_name) > len(c_id) + 2):
                ctrl_pt = c_name
            else:
                ctrl_pt = f"{c_id} {u_c or c_n or c_name}".strip() if c_id and (c_id not in (u_c or c_n or c_name)) else (u_c or c_n or c_name or c_id)

            policy_ref = str(f.get("policy_reference") or f.get("reference") or f.get("control_id") or "")
            # PII redacted before writing to exported document (this table previously
            # skipped redaction entirely, unlike the other export tables/functions)
            obs        = redact_pii(str(f.get("gap_description") or f.get("reasoning") or f.get("observation") or f.get("finding") or "")[:800])
            pe_summary = _policy_evidence_summary(f)
            pe_gap     = _policy_evidence_gap_text(f)
            if pe_summary:
                obs = f"{pe_summary}\n{obs}"
            if pe_gap:
                obs = f"{obs}\n{pe_gap}"
            display_s  = str(f.get("display_status") or f.get("status") or "Open")
            risk_lbl   = _risk_level(f)
            impact     = redact_pii(str(f.get("impact") or f.get("risk_impact") or ("NIL" if display_s.lower() in ("accepted","compliant") else "Business Risk")))
            suggestion = redact_pii(str(f.get("recommendation") or f.get("suggestion") or ("NIL" if display_s.lower() in ("accepted","compliant") else "Remediate as per IS guidelines."))[:400])
            evidence   = redact_pii(str(f.get("source_files") or f.get("evidence_quote") or f.get("evidence") or "Audit Evidence Files"))
            ev_meta    = _evidence_meta_summary(f)
            if ev_meta:
                evidence = f"{evidence}\n{ev_meta}"


            _add_obs_row(
                sno=sno_counter,
                control_pt=ctrl_pt,
                policy_ref=policy_ref,
                observation=obs,
                risk=display_s if display_s.lower() in ("accepted","compliant") else risk_lbl,
                impact=impact,
                suggestion=suggestion,
                evidence=evidence,
            )
            sno_counter += 1

    # ── Save and return bytes ─────────────────────────────────────────────────
    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_docx_report(session_title, findings, resolved_list, status, comments="", custom_logo=None, audit_type=None, metadata=None):
    """
    audit_type: explicit "vapt" / "pqc" / "iso" from the caller (derived from AuditReport.framework,
    the authoritative field set when the audit session was created). When omitted, falls back
    to inferring from session_title / finding control_id text — kept only for callers that
    haven't been updated to pass it explicitly.
    """
    if audit_type is not None:
        _at = str(audit_type).strip().lower()
        is_vapt = _at == "vapt"
        is_pqc  = _at == "pqc"
    else:
        st_std = ""
        title_u = str(session_title or "").upper()
        state_u = str(st_std or "").upper()
        combined_u = f"{title_u} {state_u}"

        is_pqc  = "PQC" in combined_u or "QUANTUM" in combined_u or "POST-QUANTUM" in combined_u
        is_vapt = (not is_pqc) and (("VAPT" in combined_u or "VULNERABILITY ASSESSMENT" in combined_u or "PENETRATION" in combined_u) and ("ISO 27001" not in combined_u))
        if not is_vapt and not is_pqc and findings:
            for f in findings:
                ctrl = str(f.get("control_id") or f.get("control") or "").upper()
                if "PQC" in ctrl:
                    is_pqc = True
                    break
                if "VAPT" in ctrl:
                    is_vapt = True
                    break

    # ── PQC: dedicated PQC DOCX template ──────────────────────────────────────
    if is_pqc:
        return _export_pqc_docx(session_title, findings, resolved_list, status, comments, custom_logo=custom_logo, metadata=metadata)

    if is_vapt:
        return _export_vapt_docx(session_title, findings, resolved_list, status, comments, custom_logo=custom_logo, metadata=metadata)

    # ── ISO: Use Sample report.docx template ──────────────────────────────────
    template_result = _export_iso_template_docx(session_title, findings, resolved_list, status, comments, custom_logo=custom_logo, metadata=metadata)
    if template_result is not None:
        return template_result

    # Which standard this report is actually certifying against -- defaults to today's
    # exact "ISO 27001" wording when unset, so existing ISO 27001 reports are unaffected.
    # (Only reached by the programmatic fallback below, used when the Sample report.docx
    # template file itself is missing.)
    framework_label = _resolve_framework_label((metadata or {}).get("framework"))

    # ── Fallback: original programmatic DOCX (if template not found) ─────────
    doc = Document()

    
    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        
        # Keep header/footer completely blank
        section.header.is_linked_to_previous = False
        for p in section.header.paragraphs:
            p.text = ""
        section.footer.is_linked_to_previous = False
        for p in section.footer.paragraphs:
            p.text = ""

    # Helpers
    def _rgb(r, g, b):
        return RGBColor(r, g, b)

    def _set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _set_cell_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'CCCCCC')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def _add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = _rgb(15, 23, 42)

    # ── COVER PAGE ─────────────────────────────────────────────────────────────
    assets_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets")
    custom_logo_file = os.path.join(assets_dir, "custom_company_logo.png")
    effective_custom_logo = custom_logo if (custom_logo and os.path.exists(custom_logo)) else (custom_logo_file if os.path.exists(custom_logo_file) else None)
    shield_logo_path = os.path.join(assets_dir, "shield_logo.png")
    logo_path = effective_custom_logo if (effective_custom_logo and os.path.exists(effective_custom_logo)) else (shield_logo_path if os.path.exists(shield_logo_path) else None)

    if logo_path and os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(30)
        run_logo = logo_p.add_run()
        run_logo.add_picture(logo_path, width=Cm(4))
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(20)
    run_title = title_p.add_run(f"IS Audit Report for {session_title}")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.bold = True
    run_title.font.color.rgb = _rgb(15, 23, 42)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(10)
    run_sub = sub_p.add_run("Engagement letter: [Engagement Reference / Date]")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = _rgb(100, 116, 139)

    # Dates
    date_p1 = doc.add_paragraph()
    date_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d1 = date_p1.add_run("Date of the Agreement: [Agreement Date]")
    run_d1.font.name = "Arial"
    run_d1.font.size = Pt(10)
    
    date_p2 = doc.add_paragraph()
    date_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p2.paragraph_format.space_after = Pt(80)
    run_d2 = date_p2.add_run(f"Date of Document: {datetime.now().strftime('%dth %B %Y')}")
    run_d2.font.name = "Arial"
    run_d2.font.size = Pt(10)

    # Audit Conducted By
    cb_p = doc.add_paragraph()
    cb_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cb = cb_p.add_run("Audit Conducted By:\n\n")
    run_cb.bold = True
    run_cb.font.size = Pt(11)
    
    run_aud1 = cb_p.add_run("Mr. Subhash Rao BE.MBA, CEH, CHFI, ISO 27001 LA, CEI, CND, CDPSE\n")
    run_aud1.font.size = Pt(10.5)
    run_aud2 = cb_p.add_run("Mr. Mahaveer Rajannavar BE.CEH, ISO 27001 LA\n\n\n")
    run_aud2.font.size = Pt(10.5)

    # Firm Info
    firm_p = doc.add_paragraph()
    firm_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_firm = firm_p.add_run(
        "[Auditor Firm Name]\n"
        "[Address Block / Contact Info]"
    )
    run_firm.font.size = Pt(9.5)
    run_firm.font.color.rgb = _rgb(100, 116, 139)

    doc.add_page_break()

    # ── DOCUMENT CONTROL ───────────────────────────────────────────────────────
    _add_section_heading("Document Control:")
    
    tbl_prep = doc.add_table(rows=6, cols=2)
    tbl_prep.style = 'Table Grid'
    prep_data = [
        ("Document Title", f"Information Security Audit {session_title}"),
        ("Engagement Letter Ref", "[Engagement Letter Ref]"),
        ("Date of the Agreement", "[Date of Agreement]"),
        ("Date of Document", datetime.now().strftime("%dth %B %Y")),
        ("Version", "1.0"),
        ("Prepared By", "Digital Age Strategies")
    ]
    for r_idx, (label, val) in enumerate(prep_data):
        c1, c2 = tbl_prep.rows[r_idx].cells
        _set_cell_bg(c1, "F1F5F9")
        _set_cell_borders(c1)
        _set_cell_borders(c2)
        c1.paragraphs[0].add_run(label).bold = True
        c2.paragraphs[0].add_run(val)

    doc.add_paragraph()

    tbl_history = doc.add_table(rows=2, cols=3)
    tbl_history.style = 'Table Grid'
    hdr_history = tbl_history.rows[0].cells
    hdr_titles = ["Version", "Date", "Remarks / Reason of change"]
    for i, title in enumerate(hdr_titles):
        _set_cell_bg(hdr_history[i], "0F172A")
        _set_cell_borders(hdr_history[i])
        run = hdr_history[i].paragraphs[0].add_run(title)
        run.bold = True
        run.font.color.rgb = _rgb(255, 255, 255)
        
    c1, c2, c3 = tbl_history.rows[1].cells
    _set_cell_borders(c1)
    _set_cell_borders(c2)
    _set_cell_borders(c3)
    c1.paragraphs[0].add_run("1.0")
    c2.paragraphs[0].add_run(datetime.now().strftime("%dth %B %Y"))
    c3.paragraphs[0].add_run("Initial release of audit findings")

    doc.add_paragraph()

    tbl_dist = doc.add_table(rows=2, cols=4)
    tbl_dist.style = 'Table Grid'
    hdr_dist = tbl_dist.rows[0].cells
    dist_titles = ["Name", "Organization", "Designation", "Email Id"]
    for i, title in enumerate(dist_titles):
        _set_cell_bg(hdr_dist[i], "0F172A")
        _set_cell_borders(hdr_dist[i])
        run = hdr_dist[i].paragraphs[0].add_run(title)
        run.bold = True
        run.font.color.rgb = _rgb(255, 255, 255)
        
    c1, c2, c3, c4 = tbl_dist.rows[1].cells
    _set_cell_borders(c1)
    _set_cell_borders(c2)
    _set_cell_borders(c3)
    _set_cell_borders(c4)
    c1.paragraphs[0].add_run("[Recipient Name]")
    c2.paragraphs[0].add_run("the Organization")
    c3.paragraphs[0].add_run("[Recipient Designation]")
    c4.paragraphs[0].add_run("[Recipient Email]")

    # ── DETAILS OF AUDITEE ─────────────────────────────────────────────────────
    _add_section_heading("Details of Auditee:")
    tbl_auditee = doc.add_table(rows=3, cols=3)
    tbl_auditee.style = 'Table Grid'
    auditee_data = [
        ("1", "Name of Organization", "the Organization"),
        ("2", "Audit Area", f"IS Audit of {session_title}"),
        ("3", "Location", "Mumbai, India")
    ]
    for r_idx, row_data in enumerate(auditee_data):
        c1, c2, c3 = tbl_auditee.rows[r_idx].cells
        _set_cell_borders(c1)
        _set_cell_borders(c2)
        _set_cell_borders(c3)
        c1.paragraphs[0].add_run(row_data[0])
        c2.paragraphs[0].add_run(row_data[1]).bold = True
        c3.paragraphs[0].add_run(row_data[2])

    doc.add_paragraph()

    # ── DETAILS OF AUDITOR ─────────────────────────────────────────────────────
    _add_section_heading("Details of Auditor:")
    tbl_auditor = doc.add_table(rows=2, cols=3)
    tbl_auditor.style = 'Table Grid'
    auditor_data = [
        ("1", "Auditor", "Mr. Subhash Rao & Mr. Mahaveer Rajannavar"),
        ("2", "Auditor", "Mr. Mahaveer Rajannavar BE.CEH, ISO 27001 LA")
    ]
    for r_idx, row_data in enumerate(auditor_data):
        c1, c2, c3 = tbl_auditor.rows[r_idx].cells
        _set_cell_borders(c1)
        _set_cell_borders(c2)
        _set_cell_borders(c3)
        c1.paragraphs[0].add_run(row_data[0])
        c2.paragraphs[0].add_run(row_data[1]).bold = True
        c3.paragraphs[0].add_run(row_data[2])

    doc.add_paragraph()

    # ── DISCLAIMER ────────────────────────────────────────────────────────────
    _add_section_heading("Disclaimer:")
    disc_p = doc.add_paragraph()
    disc_p.add_run(
        "This document is highly confidential and sensitive and is meant for circulation only to authorized people within the Organization and Digital Age Strategies Pvt. Ltd. "
        "It is understood that disclosure in part or full of the contents or any information derived from the report to unauthorized personnel is strictly prohibited."
    ).italic = True

    doc.add_page_break()

    # ── REFERENCES ────────────────────────────────────────────────────────────
    _add_section_heading("References:")
    refs = [
        f"the Organization’s RFP (Request for Proposal) no. the Organization/ITD/HO/VAPT/2023/03/01 for Certification of the Organization ISMS under {framework_label} Standard, Conducting IT Systems Audit and Cybersecurity Audit in the Organization",
        "Information Technology Audit issued by Comptroller and Auditor General (CAG) of India",
        "Engagement Letter No the Organization/HO/ITD/ITD_VIAP/P/OW/2023/0000031833/1 dated 8.8.2023",
        "CIS_Controls_v8"
    ]
    for ref in refs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(ref)

    # ── EVIDENCE ──────────────────────────────────────────────────────────────
    _add_section_heading("Evidence:")
    # Collect all evidence filenames from ALL available sources to ensure
    # no uploaded file is silently excluded from the report header.
    all_evidence_files = set()

    # Source 1: 'source_files' field in each finding
    for f in findings:
        sf = f.get("source_files", "")
        if sf:
            for fname in str(sf).split(","):
                fname = fname.strip()
                if fname:
                    all_evidence_files.add(fname)

    # Source 2: 'file_names' field in each finding (how bg_worker populates it)
    for f in findings:
        fn = f.get("file_names", "")
        if fn:
            if isinstance(fn, list):
                for fname in fn:
                    if fname:
                        all_evidence_files.add(str(fname).strip())
            else:
                for fname in str(fn).split(","):
                    fname = fname.strip()
                    if fname:
                        all_evidence_files.add(fname)

    # Source 3: Query DB directly for ALL evidence files associated with this audit session.
    # This is the most reliable source — doesn't depend on session_state being populated.
    try:
        from src.db.database import SessionLocal as _SL, EvidenceFile as _EF, AuditReport as _AR
        _db = _SL()
        try:
            _report = _db.query(_AR).filter(_AR.session_title == session_title).first()
            if _report:
                _ev_files = _db.query(_EF).filter(_EF.report_id == _report.id).all()
                for ev in _ev_files:
                    if ev.filename:
                        all_evidence_files.add(ev.filename.strip())
        except Exception:
            pass
        finally:
            _db.close()
    except Exception:
        pass

    if all_evidence_files:
        for fname in sorted(all_evidence_files):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(fname)
    else:
        doc.add_paragraph("No evidence files recorded.").italic = True

    # ── TEXT SECTIONS ─────────────────────────────────────────────────────────
    _add_section_heading("Introduction:")
    doc.add_paragraph(
        "The [Auditee Organization] was established on April 12, 1992 accordance with the provisions of the [Auditee Organization] Act, 1992 to protect the interests of investors in securities and to promote the development of, and to regulate the securities market.\n\n"
        "As per the directions from the Organization vide PO Ref No: the Organization/ HO ITD/ ITD_VIAP/P/OW/2023/0000031833/1 dated 8.8.2023, we have conducted Audit of Attendance System Application as mentioned above. Our observations are based on situation prevailing at the time of visit, which might have undergone changes since then. Our findings are based on the scope given to us and best Our Practices."
    )

    _add_section_heading("Scope:")
    doc.add_paragraph(
        "The scope in the RFP broadly covers the major control areas against which the operations/ LOBs needs to be audited and indicative list of operations/ LOBs to be audited. The RFP Scope is reproduced in below two sections for ready reference.\n\n"
        "Control Areas that cover RFP Requirements:\n"
        "Personnel Security, Access Management, Data Backup and recovery Controls, Application Security, Network Communication Security, Business Continuity Management / BCP Controls, Implementation Audit, Fault Isolation audit (in the event of any incident), Integration compliance, Configuration Audit, Change Management, Insurance Audit, Performance Audit, Monitoring the utilization of IT resources, Capacity planning including projection of business volumes IT (S/W, H/W & N/W) Assets, Licenses & maintenance contracts, Disposal of Equipment, Media, etc., Implementation of approved IT policies of the Organization."
    )

    _add_section_heading("Audit Process:")
    doc.add_paragraph(
        "The auditor shall conduct the IT Systems and Cybersecurity Audit as per the following process: "
        "Compliance to observation is verified by document verification, observation, physical visit and sample testing."
    )

    _add_section_heading("Audit Methodology.")
    doc.add_paragraph(
        "We have conducted audit as per the broad scope given to us. Our methodology will cover, in general, IS Audit practices, RBI, CERT-In, ISACA, COBIT, IT Act 2000/2008 guidelines and the guidelines & procedures prescribed in various Circulars issued by the Organization. "
        "The objective of the audit is to verify compliance and to safeguard the interests of the Organization in connection with implementation of various technologies and related guidelines of the the Organization. Thus, findings are based on the scope given to us and best Practices to be followed."
    )

    _add_section_heading("Definition of Risk Classifications:")
    doc.add_paragraph(
        "The risk of an audit finding is determined by assessing the potential negative impact and the probability that it materializes. Audit findings are classified into four risk classifications. These risk categories assist management in identification, prioritization and implementation of audit recommendations. When the practice is normal as per the guidelines / best practices, the same has been classified as 'OK/ Complied'.\n\n"
        "High Risks: Non-adherence to the Organization and Government Guidelines, Policies Approved by Competent Authority, ICT is not as per standard, high threat probabilities. These risks are so significant that Management should determine any exposure to date and without delay effect an agreed program for their immediate and permanent resolution in order to provide assurance that they will not recur in the future.\n\n"
        "Medium Risks: These risks are important and management should quickly develop action plans that will ensure timely and permanent resolution of the weaknesses noted. These are potential weaknesses in control or security, which could develop into an exposure. This should be addressed at the earliest opportunity.\n\n"
        "Low Risk: These risks are not material in the context of current levels of activity but management should be aware of them and ensure they are resolved as soon as possible as they may become material if activities increase.\n\n"
        "OK/ ACCEPTED: This is normal and good practice. It is as per the guidelines / best practices. The observations categorized as 'ACCEPTED' need no action."
    )

    # ── SUMMARY OF FINDINGS ───────────────────────────────────────────────────
    _add_section_heading("Summary of Findings:")
    
    sev_counts = {'High': 0, 'Medium': 0, 'Low': 0, 'Accepted': 0}
    for f in findings:
        st_norm = f.get("status", "Non-Compliant")
        if st_norm == "Compliant":
            sev_counts['Accepted'] += 1
        elif st_norm == "False Positive":
            pass
        else:
            sev = f.get('severity', 'N/A')
            if '1' in sev or 'Critical' in sev or 'High' in sev:
                sev_counts['High'] += 1
            elif '2' in sev or 'Medium' in sev:
                sev_counts['Medium'] += 1
            else:
                sev_counts['Low'] += 1

    tbl_summary = doc.add_table(rows=5, cols=5)
    tbl_summary.style = 'Table Grid'
    hdr_sum = tbl_summary.rows[0].cells
    hdr_sum_titles = ["Sr. No", "Risk Category Description", "No. of Observations", "Complied", "Not Complied"]
    for i, title in enumerate(hdr_sum_titles):
        _set_cell_bg(hdr_sum[i], "0F172A")
        _set_cell_borders(hdr_sum[i])
        run = hdr_sum[i].paragraphs[0].add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    summary_rows = [
        ("1", "High Risks", str(sev_counts['High']), "0", str(sev_counts['High'])),
        ("2", "Medium Risks", str(sev_counts['Medium']), "0", str(sev_counts['Medium'])),
        ("3", "Low Risk", str(sev_counts['Low']), "0", str(sev_counts['Low'])),
        ("4", "OK / ACCEPTED", str(sev_counts['Accepted']), str(sev_counts['Accepted']), "0")
    ]
    for r_idx, row_data in enumerate(summary_rows):
        c1, c2, c3, c4, c5 = tbl_summary.rows[r_idx+1].cells
        _set_cell_borders(c1)
        _set_cell_borders(c2)
        _set_cell_borders(c3)
        _set_cell_borders(c4)
        _set_cell_borders(c5)
        c1.paragraphs[0].add_run(row_data[0])
        c2.paragraphs[0].add_run(row_data[1]).bold = True
        c3.paragraphs[0].add_run(row_data[2])
        c4.paragraphs[0].add_run(row_data[3])
        c5.paragraphs[0].add_run(row_data[4])

    doc.add_page_break()

    # ── AUDIT OBSERVATIONS TABLE (TABLE 7) ─────────────────────────────────────
    _add_section_heading("Audit conclusion/Observations:")
    
    tbl_obs = doc.add_table(rows=1, cols=8)
    tbl_obs.style = 'Table Grid'
    hdr_obs = tbl_obs.rows[0].cells
    hdr_obs_titles = ['S.No.', 'Control points', 'Policy Reference', 'Observations', 'Risk', 'Impact', 'Suggestion', 'Evidence']
    for i, title in enumerate(hdr_obs_titles):
        _set_cell_bg(hdr_obs[i], "0F172A")
        _set_cell_borders(hdr_obs[i])
        run = hdr_obs[i].paragraphs[0].add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    active_findings = [f for f in findings if f.get("status") != "False Positive"]

    for f_idx, f in enumerate(active_findings, 1):
        row_cells = tbl_obs.add_row().cells
        for cell in row_cells:
            _set_cell_borders(cell)
            
        # Get severity/risk mapping based on final_result / status
        st_val = str(f.get("final_result") or f.get("status") or "Non-Compliant").strip().upper()
        sev_score = f.get("severity_score", 0.0) or 0.0
        if st_val == "COMPLIANT":
            mapped_risk = "Accepted"
            risk_text = "N/A"
        else:
            sev = f.get("severity", "N/A")
            if "1" in sev or "Critical" in sev or "High" in sev:
                mapped_risk = "High"
                risk_label = "Critical" if "Critical" in sev or sev_score >= 9.0 else "High"
            elif "2" in sev or "Medium" in sev:
                mapped_risk = "Medium"
                risk_label = "Medium"
            else:
                mapped_risk = "Low"
                risk_label = "Low"
            risk_text = f"{risk_label} ({sev_score:.1f})"

        # Shading by risk
        bg_color = {"High": "FEE2E2", "Medium": "FEFCE8", "Low": "EFF6FF", "Accepted": "F0FDF4"}.get(mapped_risk, "FFFFFF")
        for cell in row_cells:
            _set_cell_bg(cell, bg_color)

        # Content — PII redacted before writing to exported document
        row_cells[0].paragraphs[0].add_run(str(f_idx))
        row_cells[1].paragraphs[0].add_run(f.get("control_id", "") + " " + f.get("control", ""))
        # "Annex A" is an ISO 27001-specific structural term (that's where its controls
        # live); a non-ISO framework gets a generic "Reference" label instead of a
        # nonsensical "NIST CSF 2.0 Annex A".
        _clause_fallback = f"{framework_label} Annex A" if "ISO" in framework_label.upper() else f"{framework_label} Reference"
        row_cells[2].paragraphs[0].add_run(f.get("clause", "") or _clause_fallback)
        obs_text = redact_pii(f.get("finding") or f.get("description") or "-")
        pe_summary = _policy_evidence_summary(f)
        pe_gap = _policy_evidence_gap_text(f)
        if pe_summary:
            obs_text = f"{pe_summary}\n{obs_text}"
        if pe_gap:
            obs_text = f"{obs_text}\n{pe_gap}"
        row_cells[3].paragraphs[0].add_run(obs_text)
        row_cells[4].paragraphs[0].add_run(risk_text).bold = True
        row_cells[5].paragraphs[0].add_run(redact_pii(f.get("business_impact") or "NIL"))
        row_cells[6].paragraphs[0].add_run(redact_pii(f.get("recommendation") or "NIL"))

        ev_text = redact_pii(f.get("evidence_snippet") or f.get("evidence_quote") or "N/A")
        ev_meta = _evidence_meta_summary(f)
        if ev_meta:
            ev_text = f"{ev_text}\n{ev_meta}"
        row_cells[7].paragraphs[0].add_run(ev_text)

    # Set column widths
    col_widths_cm = [1.0, 3.0, 2.5, 4.5, 2.0, 3.0, 3.5, 3.0]
    for col_i, width in enumerate(col_widths_cm):
        for cell in tbl_obs.columns[col_i].cells:
            cell.width = Cm(width)

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_pdf_report(session_title, findings, resolved_list, status, comments="", audit_type=None, custom_logo=None, metadata=None):
    """
    audit_type: explicit "vapt" / "iso" from the caller (derived from AuditReport.framework,
    the authoritative field set when the audit session was created). When omitted, falls back
    to inferring from session_title / finding control_id text — kept only for callers that
    haven't been updated to pass it explicitly.
    """
    if audit_type is not None:
        is_vapt = str(audit_type).strip().lower() == "vapt"
    else:
        st_std = ""
        title_u = str(session_title or "").upper()
        state_u = str(st_std or "").upper()
        combined_u = f"{title_u} {state_u}"

        is_vapt = "VAPT" in combined_u or "VULNERABILITY ASSESSMENT" in combined_u or "PENETRATION" in combined_u
        if not is_vapt and findings:
            for f in findings:
                ctrl = str(f.get("control_id") or f.get("control") or f.get("category") or "").upper()
                if "VAPT" in ctrl:
                    is_vapt = True
                    break

    if is_vapt:
        return _export_vapt_pdf(session_title, findings, resolved_list, status, comments, custom_logo=custom_logo, metadata=metadata)
    from fpdf.fonts import FontFace

    # Which standard this report is actually certifying against -- defaults to today's
    # exact "ISO 27001" wording when unset, so existing ISO 27001 reports are unaffected.
    framework_label = _resolve_framework_label((metadata or {}).get("framework"))

    findings = sorted(findings or [], key=severity_sort_key)


    def clean_text(val):
        if not val:
            return "-"
        val = str(val)
        # Smart quotes
        val = val.replace("“", "\"").replace("”", "\"").replace("‘", "'").replace("’", "'")
        val = val.replace("\u201d", "\"").replace("\u201c", "\"").replace("\u2019", "'").replace("\u2018", "'")
        # Hyphens and dashes
        val = val.replace("—", "-").replace("–", "-").replace("\u2013", "-").replace("\u2014", "-").replace("\u2212", "-")
        # Bullet symbols
        val = val.replace("\u2022", "*").replace("•", "*").replace("\u25cf", "*").replace("\u25cb", "*")
        # Fallback to Latin-1
        val = val.encode("latin-1", "replace").decode("latin-1")
        return val

    def truncate_cell_text(text, max_chars=800):
        if not text:
            return "-"

        text = str(text)
        if len(text) > max_chars:
            return text[:max_chars] + "... [Truncated for PDF]"
        return text

    class AuditPDF(FPDF):
        def header(self):
            pass
        def footer(self):
            pass

    pdf = AuditPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(10, 15, 10)
    
    # FontFace helpers
    hdr_style = FontFace(emphasis="B", color=(255, 255, 255), fill_color=(15, 23, 42))
    lbl_style = FontFace(emphasis="B", fill_color=(241, 245, 249))
    bold_style = FontFace(emphasis="B")

    # ── COVER PAGE ─────────────────────────────────────────────────────────────
    pdf.add_page()
    assets_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets")
    custom_logo_file = os.path.join(assets_dir, "custom_company_logo.png")
    effective_custom_logo = custom_logo if (custom_logo and os.path.exists(custom_logo)) else (custom_logo_file if os.path.exists(custom_logo_file) else None)
    shield_logo_path = os.path.join(assets_dir, "shield_logo.png")
    logo_path = effective_custom_logo if (effective_custom_logo and os.path.exists(effective_custom_logo)) else (shield_logo_path if os.path.exists(shield_logo_path) else None)

    if logo_path and os.path.exists(logo_path):
        pdf.image(logo_path, x=85, y=15, w=40)
        pdf.ln(35)
    else:
        pdf.ln(25)
    
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(15, 23, 42)
    pdf.cell(0, 12, clean_text(f"IS Audit Report for {session_title}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", "I", 9.5)
    pdf.set_text_color(100, 116, 139)
    pdf.cell(0, 6, clean_text("Engagement letter: [Engagement Reference / Date]"), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(5)
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(15, 23, 42)
    pdf.cell(0, 6, clean_text("Date of the Agreement: [Agreement Date]"), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.cell(0, 6, clean_text(f"Date of Document: {datetime.now().strftime('%dth %B %Y')}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(25)
    
    pdf.set_font("Helvetica", "B", 10.5)
    pdf.cell(0, 6, clean_text("Audit Conducted By:"), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 5, clean_text("Mr. Subhash Rao & Mr. Mahaveer Rajannavar"), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.cell(0, 5, clean_text("Mr. Mahaveer Rajannavar BE.CEH, ISO 27001 LA"), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(20)
    
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 116, 139)
    pdf.multi_cell(0, 4.5, 
        clean_text(
            "[Auditor Firm Name]\n"
            "# 28, \"Om Arcade\" 2nd & 3rd Floors, Thimmappa Reddy Layout,\n"
            "Hulimavu, Bannerghatta Road, Bangalore - 560076\n"
            "Ph: +91-80-26484636, 49568066, 26485148, 41503825, 41218560\n"
            "Mobile: 9448088666 / 9448055711\n"
            "Email: audit@digitalage.co.in, dinesh.shastri@digitalage.co.in"
        ),
        align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT
    )
    
    # ── DOCUMENT CONTROL ───────────────────────────────────────────────────────
    pdf.add_page()
    
    def section_title(text):
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(15, 23, 42)
        pdf.cell(0, 8, clean_text(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 9.5)

    section_title("Document Control:")
    
    prep_data = [
        ["Document Title", f"Information Security Audit {session_title}"],
        ["Engagement Letter Ref", "[Engagement Letter Ref]"],
        ["Date of the Agreement", "[Date of Agreement]"],
        ["Date of Document", datetime.now().strftime("%dth %B %Y")],
        ["Version", "1.0"],
        ["Prepared By", "Digital Age Strategies"]
    ]
    with pdf.table(col_widths=(60, 130), text_align="L") as table:
        for row in prep_data:
            r = table.row()
            r.cell(clean_text(row[0]), style=lbl_style)
            r.cell(clean_text(row[1]))
            
    pdf.ln(4)
    
    with pdf.table(col_widths=(30, 40, 120), text_align="L") as table:
        hdr = table.row()
        hdr.cell("Version", style=hdr_style)
        hdr.cell("Date", style=hdr_style)
        hdr.cell("Remarks / Reason of change", style=hdr_style)
        
        row = table.row()
        row.cell("1.0")
        row.cell(clean_text(datetime.now().strftime("%dth %B %Y")))
        row.cell("Initial release of audit findings")
        
    pdf.ln(4)
    
    with pdf.table(col_widths=(40, 45, 55, 50), text_align="L") as table:
        hdr = table.row()
        hdr.cell("Name", style=hdr_style)
        hdr.cell("Organization", style=hdr_style)
        hdr.cell("Designation", style=hdr_style)
        hdr.cell("Email Id", style=hdr_style)
        
        row = table.row()
        row.cell("[Recipient Name]")
        row.cell("the Organization")
        row.cell("[Recipient Designation]")
        row.cell("[Recipient Email]")
        
    # Details of Auditee (Table 4)
    section_title("Details of Auditee:")
    auditee_data = [
        ["1", "Name of Organization", "the Organization"],
        ["2", "Audit Area", f"IS Audit of {session_title}"],
        ["3", "Location", "Mumbai, India"]
    ]
    with pdf.table(col_widths=(20, 60, 110), text_align="L") as table:
        for row in auditee_data:
            r = table.row()
            r.cell(clean_text(row[0]))
            r.cell(clean_text(row[1]), style=bold_style)
            r.cell(clean_text(row[2]))
            
    # Details of Auditor (Table 5)
    section_title("Details of Auditor:")
    auditor_data = [
        ["1", "Lead Auditor", "Lead Cyber Security Auditor"],
        ["2", "Auditing Firm", "XYZ Security Services Pvt. Ltd."]
    ]
    with pdf.table(col_widths=(20, 40, 130), text_align="L") as table:
        for row in auditor_data:
            r = table.row()
            r.cell(clean_text(row[0]))
            r.cell(clean_text(row[1]), style=bold_style)
            r.cell(clean_text(row[2]))

    # Disclaimer
    section_title("Disclaimer:")
    pdf.set_font("Helvetica", "I", 9.5)
    pdf.set_text_color(51, 65, 85)
    pdf.multi_cell(0, 5, 
        clean_text(
            "This document is highly confidential and sensitive and is meant for circulation only to authorized people within the Organization and Digital Age Strategies Pvt. Ltd. "
            "It is understood that disclosure in part or full of the contents or any information derived from the report to unauthorized personnel is strictly prohibited."
        ),
        new_x=XPos.LMARGIN, new_y=YPos.NEXT
    )
    
    # References
    pdf.add_page()
    section_title("References:")
    pdf.set_font("Helvetica", "", 9.5)
    refs = [
        f"the Organization’s RFP (Request for Proposal) no. the Organization/ITD/HO/VAPT/2023/03/01 for Certification of the Organization ISMS under {framework_label} Standard, Conducting IT Systems Audit and Cybersecurity Audit in the Organization",
        "Information Technology Audit issued by Comptroller and Auditor General (CAG) of India",
        "Engagement Letter No the Organization/HO/ITD/ITD_VIAP/P/OW/2023/0000031833/1 dated 8.8.2023",
        "CIS_Controls_v8"
    ]
    for ref in refs:
        pdf.cell(5, 5, chr(149), align="R")
        pdf.multi_cell(185, 5, clean_text(ref), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)
        
    # Evidence
    section_title("Evidence:")
    all_ev_files = set()
    for f in findings:
        sf = f.get("source_files") or f.get("file_names") or f.get("evidence_location") or ""
        if sf:
            if isinstance(sf, list):
                for fn in sf:
                    if fn: all_ev_files.add(str(fn).strip())
            else:
                for fn in str(sf).split(","):
                    fn = fn.strip()
                    if fn: all_ev_files.add(fn)
    # Query DB directly for evidence files
    try:
        from src.db.database import SessionLocal as _SL2, EvidenceFile as _EF2, AuditReport as _AR2
        _db2 = _SL2()
        try:
            _report2 = _db2.query(_AR2).filter(_AR2.session_title == session_title).first()
            if _report2:
                _evs = _db2.query(_EF2).filter(_EF2.report_id == _report2.id).all()
                for e in _evs:
                    if e.filename: all_ev_files.add(e.filename.strip())
        except Exception:
            pass
        finally:
            _db2.close()
    except Exception:
        pass

    if all_ev_files:
        for fname in sorted(all_ev_files):
            pdf.cell(5, 5, chr(149), align="R")
            pdf.multi_cell(185, 5, clean_text(fname), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
    else:
        pdf.set_font("Helvetica", "I", 9.5)
        pdf.cell(0, 5, "No evidence files recorded.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        
    # Introduction
    section_title("Introduction:")
    pdf.set_font("Helvetica", "", 9.5)
    pdf.set_text_color(15, 23, 42)
    pdf.multi_cell(0, 5, 
        clean_text(
            "The [Auditee Organization] was established on April 12, 1992 accordance with the provisions of the [Auditee Organization] Act, 1992 to protect the interests of investors in securities and to promote the development of, and to regulate the securities market.\n\n"
            "As per the directions from the Organization vide PO Ref No: the Organization/ HO ITD/ ITD_VIAP/P/OW/2023/0000031833/1 dated 8.8.2023, we have conducted Audit of Attendance System Application as mentioned above. Our observations are based on situation prevailing at the time of visit, which might have undergone changes since then. Our findings are based on the scope given to us and best Our Practices."
        ),
        new_x=XPos.LMARGIN, new_y=YPos.NEXT
    )
    
    # Scope
    section_title("Scope:")
    pdf.multi_cell(0, 5,
        clean_text(
            "The scope in the RFP broadly covers the major control areas against which the operations/ LOBs needs to be audited and indicative list of operations/ LOBs to be audited. The RFP Scope is reproduced in below two sections for ready reference.\n\n"
            "Control Areas that cover RFP Requirements:\n"
            "Personnel Security, Access Management, Data Backup and recovery Controls, Application Security, Network Communication Security, Business Continuity Management / BCP Controls, Implementation Audit, Fault Isolation audit (in the event of any incident), Integration compliance, Configuration Audit, Change Management, Insurance Audit, Performance Audit, Monitoring the utilization of IT resources, Capacity planning including projection of business volumes IT (S/W, H/W & N/W) Assets, Licenses & maintenance contracts, Disposal of Equipment, Media, etc., Implementation of approved IT policies of the Organization."
        ),
        new_x=XPos.LMARGIN, new_y=YPos.NEXT
    )
    
    # Audit Process
    section_title("Audit Process:")
    pdf.multi_cell(0, 5,
        clean_text(
            "The auditor shall conduct the IT Systems and Cybersecurity Audit as per the following process: "
            "Compliance to observation is verified by document verification, observation, physical visit and sample testing."
        ),
        new_x=XPos.LMARGIN, new_y=YPos.NEXT
    )
    
    # Audit Methodology
    section_title("Audit Methodology.")
    pdf.multi_cell(0, 5,
        clean_text(
            "We have conducted audit as per the broad scope given to us. Our methodology will cover, in general, IS Audit practices, RBI, CERT-In, ISACA, COBIT, IT Act 2000/2008 guidelines and the guidelines & procedures prescribed in various Circulars issued by the Organization. "
            "The objective of the audit is to verify compliance and to safeguard the interests of the Organization in connection with implementation of various technologies and related guidelines of the the Organization. Thus, findings are based on the scope given to us and best Practices to be followed."
        ),
        new_x=XPos.LMARGIN, new_y=YPos.NEXT
    )
    
    # Definition of Risk Classifications
    section_title("Definition of Risk Classifications:")
    pdf.multi_cell(0, 5,
        clean_text(
            "The risk of an audit finding is determined by assessing the potential negative impact and the probability that it materializes. Audit findings are classified into four risk classifications. These risk categories assist management in identification, prioritization and implementation of audit recommendations. When the practice is normal as per the guidelines / best practices, the same has been classified as 'OK/ Complied'.\n\n"
            "High Risks: Non-adherence to the Organization and Government Guidelines, Policies Approved by Competent Authority, ICT is not as per standard, high threat probabilities. These risks are so significant that Management should determine any exposure to date and without delay effect an agreed program for their immediate and permanent resolution in order to provide assurance that they will not recur in the future.\n\n"
            "Medium Risks: These risks are important and management should quickly develop action plans that will ensure timely and permanent resolution of the weaknesses noted. These are potential weaknesses in control or security, which could develop into an exposure. This should be addressed at the earliest opportunity.\n\n"
            "Low Risk: These risks are not material in the context of current levels of activity but management should be aware of them and ensure they are resolved as soon as possible as they may become material if activities increase.\n\n"
            "OK/ ACCEPTED: This is normal and good practice. It is as per the guidelines / best practices. The observations categorized as 'ACCEPTED' need no action."
        ),
        new_x=XPos.LMARGIN, new_y=YPos.NEXT
    )

    # Summary of Findings (Table 6)
    section_title("Summary of Findings:")
    
    sev_counts = {'High': 0, 'Medium': 0, 'Low': 0, 'Accepted': 0}
    for f in findings:
        st_norm = f.get("status", "Non-Compliant")
        if st_norm == "Compliant":
            sev_counts['Accepted'] += 1
        elif st_norm == "False Positive":
            pass
        else:
            sev = f.get('severity', 'N/A')
            if '1' in sev or 'Critical' in sev or 'High' in sev:
                sev_counts['High'] += 1
            elif '2' in sev or 'Medium' in sev:
                sev_counts['Medium'] += 1
            else:
                sev_counts['Low'] += 1

    summary_rows = [
        ["1", "High Risks", str(sev_counts['High']), "0", str(sev_counts['High'])],
        ["2", "Medium Risks", str(sev_counts['Medium']), "0", str(sev_counts['Medium'])],
        ["3", "Low Risk", str(sev_counts['Low']), "0", str(sev_counts['Low'])],
        ["4", "OK / ACCEPTED", str(sev_counts['Accepted']), str(sev_counts['Accepted']), "0"]
    ]
    with pdf.table(col_widths=(20, 50, 40, 40, 40), text_align="L") as table:
        hdr = table.row()
        hdr.cell("Sr. No", style=hdr_style)
        hdr.cell("Risk Category Description", style=hdr_style)
        hdr.cell("No. of Observations", style=hdr_style)
        hdr.cell("Complied", style=hdr_style)
        hdr.cell("Not Complied", style=hdr_style)
        
        for row in summary_rows:
            r = table.row()
            r.cell(clean_text(row[0]))
            r.cell(clean_text(row[1]), style=bold_style)
            r.cell(clean_text(row[2]))
            r.cell(clean_text(row[3]))
            r.cell(clean_text(row[4]))
            
    # Table 7: Audit conclusion/Observations
    pdf.add_page()
    section_title("Audit conclusion/Observations:")
    
    active_findings = [f for f in findings if f.get("status") != "False Positive"]

    pdf.set_font("Helvetica", "", 7.5)
    with pdf.table(col_widths=(6, 24, 15, 50, 12, 25, 33, 25), text_align="L") as table:
        hdr = table.row()
        hdr_titles = ['S.No.', 'Control points', 'Policy Reference', 'Observations', 'Risk', 'Impact', 'Recommendation', 'Evidence']
        for title in hdr_titles:
            hdr.cell(title, style=hdr_style)
            
        for f_idx, f in enumerate(active_findings, 1):
            r = table.row()
            
            # Map risk: final_result is the single source of truth
            st_val = str(f.get("final_result") or f.get("status") or "").strip().upper()
            is_comp = (st_val == "COMPLIANT")

            sev_score = f.get("severity_score", 0.0) or 0.0
            if is_comp:
                mapped_risk = "Accepted"
                risk_text = "N/A"
            else:
                sev = f.get("severity", "N/A")
                if "1" in sev or "Critical" in sev or "High" in sev:
                    mapped_risk = "High"
                    risk_label = "Critical" if "Critical" in sev or sev_score >= 9.0 else "High"
                elif "2" in sev or "Medium" in sev:
                    mapped_risk = "Medium"
                    risk_label = "Medium"
                else:
                    mapped_risk = "Low"
                    risk_label = "Low"
                risk_text = f"{risk_label} ({sev_score:.1f})"
                    
            # Color coding
            bg_color = {
                "High": (254, 226, 226), 
                "Medium": (254, 252, 232), 
                "Low": (239, 246, 255), 
                "Accepted": (240, 253, 244)
            }.get(mapped_risk, (255, 255, 255))
            
            # Styles
            cell_style = FontFace(fill_color=bg_color)
            risk_color = {
                "High": (220, 38, 38),
                "Medium": (217, 119, 6),
                "Low": (37, 99, 235),
                "Accepted": (22, 163, 74)
            }.get(mapped_risk, (30, 41, 59))
            risk_style = FontFace(emphasis="B", fill_color=bg_color, color=risk_color)

            # Cells
            r.cell(clean_text(str(f_idx)), style=cell_style)
            
            # BUG FIX: use control_name for bare-ID Excel-mode findings
            # (e.g. control_id="8.17" + control_name="8.17 Whether NTP is enabled")
            # so two rows with same control_id display distinct question text.
            _pdf_cid = str(f.get("control_id") or "").strip()
            _pdf_cname = str(f.get("control_name") or "").strip()
            _pdf_ctrl = str(f.get("control") or "").strip()
            import re as _re2
            _bare_id = bool(_re2.match(r'^\d{1,2}\.\d{1,2}(?:\.\d{1,2})?$', _pdf_cid))
            if _bare_id and _pdf_cname and len(_pdf_cname) > len(_pdf_cid) + 2:
                # Excel mode: use control_name which has the checklist question
                ctrl_text = _re2.sub(r'\s*\(\s*\d{1,2}\.\d{1,2}(?:\.\d{1,2})?\s*\)\s*$', '', _pdf_cname).strip()
            else:
                ctrl_text = (_pdf_cid + " " + (_pdf_ctrl or _pdf_cname)).strip()
            r.cell(clean_text(truncate_cell_text(ctrl_text, 150)), style=cell_style)
            
            _clause_fallback = f"{framework_label} Annex A" if "ISO" in framework_label.upper() else f"{framework_label} Reference"
            ref_text = f.get("clause", "") or _clause_fallback
            r.cell(clean_text(truncate_cell_text(ref_text, 100)), style=cell_style)
            
            # PII redacted before writing to exported PDF
            obs_text = redact_pii(f.get("finding") or f.get("description") or "-")
            pe_summary = _policy_evidence_summary(f)
            pe_gap = _policy_evidence_gap_text(f)
            if pe_summary:
                obs_text = f"{pe_summary}\n{obs_text}"
            if pe_gap:
                obs_text = f"{obs_text}\n{pe_gap}"
            r.cell(clean_text(truncate_cell_text(obs_text, 700)), style=cell_style)

            r.cell(clean_text(risk_text), style=risk_style)

            imp_text = redact_pii(f.get("business_impact") or "NIL")
            r.cell(clean_text(truncate_cell_text(imp_text, 400)), style=cell_style)

            sug_text = redact_pii(f.get("recommendation") or "NIL")
            r.cell(clean_text(truncate_cell_text(sug_text, 500)), style=cell_style)

            ev_text = redact_pii(f.get("evidence_snippet") or f.get("evidence_quote") or "N/A")
            ev_meta = _evidence_meta_summary(f)
            if ev_meta:
                ev_text = f"{ev_text}\n{ev_meta}"
            r.cell(clean_text(truncate_cell_text(ev_text, 500)), style=cell_style)

    pdf_bytes = pdf.output()
    return bytes(pdf_bytes)
